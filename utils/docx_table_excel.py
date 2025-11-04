#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 Word 文档中“嵌入的 Excel 工作表对象”**就地**转换为标准 Word 表格：
- 保留输入文件的所有其它内容与结构；
- 在对象原所在段落处插入等价的 Word 表格；
- 删除该对象对应的 run（避免出现多余占位图片/对象）；
- **不**插入任何“从 xxx 转换的Excel表格”之类说明性文字；

兼容性修复：
- 不使用 xpath(namespaces=...)；改为 Clark notation（"{uri}local"）。
- 解决 KeyError: 'o'（未注册 o 命名空间）。
- 解决 KeyError: "no style with name 'Table Grid'"：增加**安全表格样式设置**，在缺失时优雅降级为其它可用表格样式，最后允许无样式，并可选加网格边框。

新增改动（仅处理工作表，避免误处理图片/图表）：
- 仅允许 Excel Worksheet ProgID：'Excel.Sheet.*'（统一小写匹配 'excel.sheet'/'excel.sheet.12'/'excel.sheet.8' 等）
- 严格扩展名过滤：只接受 .xlsx / .xls / .bin
- 对于 .bin：若缺少 olefile 或未找到 Workbook 流 → 直接跳过（不回退成占位表）
"""

import os
import zipfile
import tempfile
import shutil
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except Exception:
    PANDAS_AVAILABLE = False

try:
    import olefile
    OLEFILE_AVAILABLE = True
except Exception:
    OLEFILE_AVAILABLE = False

# 命名空间常量（使用 Clark notation）
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
O_NS = 'urn:schemas-microsoft-com:office:office'
V_NS = 'urn:schemas-microsoft-com:vml'

TAG_OLEOBJECT = f'{{{O_NS}}}OLEObject'   # o:OLEObject
TAG_RUN       = f'{{{W_NS}}}r'           # w:r
ATTR_RID      = f'{{{R_NS}}}id'          # r:id

# ✅ 仅放行“Excel 工作表”类型（统一转小写后做包含判断）
ALLOWED_EXCEL_PROGIDS = {
    'excel.sheet', 'excel.sheet.12', 'excel.sheet.8', 'worksheet'
}

# ------------------------------
# 工具函数
# ------------------------------

def build_rid_to_target_map(docx_path: str) -> Dict[str, str]:
    """构建 rId → embeddings/xxx 的映射（只处理主文档）。"""
    from xml.etree import ElementTree as ET
    rid_map: Dict[str, str] = {}
    with zipfile.ZipFile(docx_path, 'r') as z:
        rels_path = 'word/_rels/document.xml.rels'
        if rels_path not in z.namelist():
            return rid_map
        xml = z.read(rels_path)
        root = ET.fromstring(xml)
        for rel in root.findall('{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
            rId = rel.get('Id')
            target = rel.get('Target')
            if rId and target and target.startswith('embeddings/'):
                rid_map[rId] = 'word/' + target if not target.startswith('word/') else target
    return rid_map


def extract_embeddings(docx_path: str, temp_dir: str) -> Dict[str, str]:
    """将 embeddings 文件提取到临时目录，返回 Target → 提取后的本地路径 的映射。"""
    out: Dict[str, str] = {}
    with zipfile.ZipFile(docx_path, 'r') as z:
        for info in z.infolist():
            nm = info.filename
            if not nm.startswith('word/embeddings/'):
                continue
            if info.is_dir():
                continue
            filename = os.path.basename(nm)
            if not filename:
                continue
            dst = os.path.join(temp_dir, filename)
            os.makedirs(os.path.dirname(dst), exist_ok=True)
            with z.open(info, 'r') as src, open(dst, 'wb') as f:
                shutil.copyfileobj(src, f)
            out[nm] = dst
    return out


def clean_excel_data(df) -> 'pd.DataFrame':
    """智能清理Excel数据，处理重复列名和空白列名"""
    try:
        columns = list(df.columns)
        seen_columns = {}
        new_columns = []

        for i, col in enumerate(columns):
            col_str = str(col)
            if col_str in seen_columns:
                seen_columns[col_str] += 1
                new_columns.append(f"{col_str}.{seen_columns[col_str]}")
            else:
                seen_columns[col_str] = 0
                new_columns.append(col_str)

        df.columns = new_columns

        for i, col in enumerate(df.columns):
            if 'Unnamed' in str(col):
                if i > 0 and len(df) > 0:
                    first_row_value = str(df.iloc[0, i])
                    if first_row_value and first_row_value != 'nan':
                        if not first_row_value.replace('.', '').isdigit():
                            df.columns = list(df.columns[:i]) + [first_row_value] + list(df.columns[i+1:])
                            continue
                df.columns = list(df.columns[:i]) + [f'列{i+1}'] + list(df.columns[i+1:])

        df = df.loc[:, df.notna().any()]
        df = df.loc[df.notna().any(axis=1)]
        df = df.reset_index(drop=True)

        print(f"    清理后表格形状: {df.shape}")
        print(f"    清理后列名: {list(df.columns)}")
        return df

    except Exception as e:
        print(f"    数据清理失败，使用原始数据: {e}")
        return df


def dataframe_to_matrix(df) -> List[List[str]]:
    df = df.fillna('')
    cleaned_columns = []
    for col in df.columns:
        col_str = str(col)
        if col_str.startswith('Unnamed:'):
            cleaned_columns.append('')
        elif '.' in col_str and col_str.split('.')[0].isdigit():
            base_name = col_str.split('.')[0]
            if base_name in cleaned_columns:
                cleaned_columns.append('')
            else:
                cleaned_columns.append(base_name)
        else:
            cleaned_columns.append(col_str)

    data = [cleaned_columns]
    data.extend([[str(x) for x in row] for row in df.to_numpy().tolist()])
    return data


def read_excel_to_matrix(excel_path: str) -> Tuple[List[List[str]], str]:
    """读取 excel 为二维字符串数组；返回 (matrix, 描述信息)。若失败返回占位表。"""
    title = ''
    basename = os.path.basename(excel_path)

    # 只处理真正的 .xlsx / .xls（.bin 的解析由主流程专门处理）
    if not (excel_path.lower().endswith('.xlsx') or excel_path.lower().endswith('.xls')):
        return [["A","B","C"],["-","-","-"]], title

    if not PANDAS_AVAILABLE:
        return [["A","B","C"],["-","-","-"]], title
    try:
        dfs = pd.read_excel(excel_path, sheet_name=None)
        if not dfs:
            return [["A","B"],["-","-"]], title
        sheet_name, df = next(iter(dfs.items()))
        title = f"{basename} - {sheet_name}"
        return dataframe_to_matrix(df), title
    except Exception:
        return [["A","B","C"],["-","-","-"]], title


def _add_grid_borders_if_needed(tbl):
    try:
        tblPr = tbl._tbl.tblPr or tbl._tbl.get_or_add_tblPr()
        if tblPr.find(qn('w:tblBorders')) is not None:
            return
        from docx.oxml import OxmlElement
        borders = OxmlElement('w:tblBorders')
        for edge in ('top','left','bottom','right','insideH','insideV'):
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), 'auto')
            borders.append(el)
        tblPr.append(borders)
    except Exception:
        pass


def _apply_table_style_best_effort(doc: Document, tbl):
    preferred_names = (
        'Table Grid', 'Normal Table', 'Table Normal', 'TableGrid',
        '网格型', '普通表格'
    )
    try:
        for name in preferred_names:
            try:
                st = doc.styles[name]
                if st.type == WD_STYLE_TYPE.TABLE:
                    tbl.style = st
                    return
            except KeyError:
                continue
        for st in doc.styles:
            try:
                if st.type == WD_STYLE_TYPE.TABLE:
                    tbl.style = st
                    return
            except Exception:
                continue
    except Exception:
        pass
    _add_grid_borders_if_needed(tbl)


def _apply_merged_cell_borders(tbl, merged_cells):
    try:
        from docx.oxml import OxmlElement
        merged_map = {}
        for merged_range in merged_cells:
            rlo, rhi, clo, chi = merged_range
            for row in range(rlo, rhi):
                for col in range(clo, chi):
                    is_master = (row == rlo and col == clo)
                    merge_info = {
                        'master': (rlo, clo),
                        'size': (rhi - rlo, chi - clo),
                        'is_top': row == rlo,
                        'is_bottom': row == rhi - 1,
                        'is_left': col == clo,
                        'is_right': col == chi - 1
                    }
                    merged_map[(row, col)] = (is_master, merge_info)

        print(f"    为 {len(merged_cells)} 个合并区域应用边框规则")

        for row_idx in range(len(tbl.rows)):
            for col_idx in range(len(tbl.columns)):
                cell = tbl.cell(row_idx, col_idx)
                cell_tcPr = cell._tc.get_or_add_tcPr()
                old_borders = cell_tcPr.find(qn('w:tcBorders'))
                if old_borders is not None:
                    cell_tcPr.remove(old_borders)
                tc_borders = OxmlElement('w:tcBorders')

                if (row_idx, col_idx) in merged_map:
                    is_master, merge_info = merged_map[(row_idx, col_idx)]
                    if is_master:
                        _add_border_to_element(tc_borders, 'top', 'single', '4', 'auto')
                        _add_border_to_element(tc_borders, 'left', 'single', '4', 'auto')
                        _add_border_to_element(tc_borders, 'bottom', 'single', '4', 'auto')
                        _add_border_to_element(tc_borders, 'right', 'single', '4', 'auto')
                        print(f"    主单元格({row_idx+1},{col_idx+1}): 完整边框")
                    else:
                        borders_to_add = []
                        if merge_info['is_top'] and not merge_info['is_left']:
                            borders_to_add.append('top')
                        if merge_info['is_bottom']:
                            borders_to_add.append('bottom')
                        if merge_info['is_left'] and not merge_info['is_top']:
                            borders_to_add.append('left')
                        if merge_info['is_right']:
                            borders_to_add.append('right')
                        for border in borders_to_add:
                            _add_border_to_element(tc_borders, border, 'single', '4', 'auto')
                        if borders_to_add:
                            print(f"    从单元格({row_idx+1},{col_idx+1}): {', '.join(borders_to_add)}边框")
                        else:
                            print(f"    从单元格({row_idx+1},{col_idx+1}): 无边框")
                else:
                    _add_border_to_element(tc_borders, 'top', 'single', '4', 'auto')
                    _add_border_to_element(tc_borders, 'left', 'single', '4', 'auto')
                    _add_border_to_element(tc_borders, 'bottom', 'single', '4', 'auto')
                    _add_border_to_element(tc_borders, 'right', 'single', '4', 'auto')

                cell_tcPr.append(tc_borders)

        print(f"    合并单元格模式下跳过表级内部边框设置")

    except Exception as e:
        print(f"    应用合并单元格边框失败: {e}")
        _apply_standard_table_borders(tbl)


def _apply_standard_table_borders(tbl):
    try:
        from docx.oxml import OxmlElement
        for row_idx in range(len(tbl.rows)):
            for col_idx in range(len(tbl.columns)):
                cell = tbl.cell(row_idx, col_idx)
                cell_tcPr = cell._tc.get_or_add_tcPr()
                old_borders = cell_tcPr.find(qn('w:tcBorders'))
                if old_borders is not None:
                    cell_tcPr.remove(old_borders)
                tc_borders = OxmlElement('w:tcBorders')
                _add_border_to_element(tc_borders, 'top', 'single', '4', 'auto')
                _add_border_to_element(tc_borders, 'left', 'single', '4', 'auto')
                _add_border_to_element(tc_borders, 'bottom', 'single', '4', 'auto')
                _add_border_to_element(tc_borders, 'right', 'single', '4', 'auto')
                cell_tcPr.append(tc_borders)
        print(f"    为普通表格应用标准边框完成")
    except Exception as e:
        print(f"    应用标准边框失败: {e}")


def _apply_openpyxl_merged_cell_borders(tbl, merged_ranges):
    try:
        from docx.oxml import OxmlElement
        merged_map = {}
        for merged_range in merged_ranges:
            min_row, min_col, max_row, max_col = merged_range.min_row, merged_range.min_col, merged_range.max_row, merged_range.max_col
            for row in range(min_row, max_row + 1):
                for col in range(min_col, max_col + 1):
                    is_master = (row == min_row and col == min_col)
                    merge_info = {
                        'master': (min_row - 1, min_col - 1),
                        'size': (max_row - min_row, max_col - min_col),
                        'is_top': row == min_row,
                        'is_bottom': row == max_row,
                        'is_left': col == min_col,
                        'is_right': col == max_col
                    }
                    merged_map[(row - 1, col - 1)] = (is_master, merge_info)

        print(f"    为 {len(merged_ranges)} 个openpyxl合并区域应用边框规则")

        for row_idx in range(len(tbl.rows)):
            for col_idx in range(len(tbl.columns)):
                cell = tbl.cell(row_idx, col_idx)
                cell_tcPr = cell._tc.get_or_add_tcPr()
                old_borders = cell_tcPr.find(qn('w:tcBorders'))
                if old_borders is not None:
                    cell_tcPr.remove(old_borders)
                tc_borders = OxmlElement('w:tcBorders')

                if (row_idx, col_idx) in merged_map:
                    is_master, merge_info = merged_map[(row_idx, col_idx)]
                    if is_master:
                        _add_border_to_element(tc_borders, 'top', 'single', '4', 'auto')
                        _add_border_to_element(tc_borders, 'left', 'single', '4', 'auto')
                        _add_border_to_element(tc_borders, 'bottom', 'single', '4', 'auto')
                        _add_border_to_element(tc_borders, 'right', 'single', '4', 'auto')
                        print(f"    主单元格({row_idx+1},{col_idx+1}): 完整边框")
                    else:
                        borders_to_add = []
                        if merge_info['is_top'] and not merge_info['is_left']:
                            borders_to_add.append('top')
                        if merge_info['is_bottom']:
                            borders_to_add.append('bottom')
                        if merge_info['is_left'] and not merge_info['is_top']:
                            borders_to_add.append('left')
                        if merge_info['is_right']:
                            borders_to_add.append('right')
                        for border in borders_to_add:
                            _add_border_to_element(tc_borders, border, 'single', '4', 'auto')
                        if borders_to_add:
                            print(f"    从单元格({row_idx+1},{col_idx+1}): {', '.join(borders_to_add)}边框")
                        else:
                            print(f"    从单元格({row_idx+1},{col_idx+1}): 无边框")
                else:
                    _add_border_to_element(tc_borders, 'top', 'single', '4', 'auto')
                    _add_border_to_element(tc_borders, 'left', 'single', '4', 'auto')
                    _add_border_to_element(tc_borders, 'bottom', 'single', '4', 'auto')
                    _add_border_to_element(tc_borders, 'right', 'single', '4', 'auto')

                cell_tcPr.append(tc_borders)

        print(f"    合并单元格模式下跳过表级内部边框设置")

    except Exception as e:
        print(f"    应用openpyxl合并单元格边框失败: {e}")
        _apply_standard_table_borders(tbl)


def format_cell_value(cell_value):
    import xlrd
    if isinstance(cell_value, float):
        if cell_value.is_integer():
            return str(int(cell_value))
        else:
            formatted = f"{cell_value}".rstrip('0').rstrip('.')
            return formatted
    elif isinstance(cell_value, int):
        return str(cell_value)
    elif isinstance(cell_value, str):
        return cell_value.strip()
    else:
        return str(cell_value)


def _normalize_row_spacing(tbl):
    try:
        from docx.shared import Pt
        from docx.enum.table import WD_ROW_HEIGHT_RULE

        for r, row in enumerate(tbl.rows):
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            if r == 0:
                row.height = Pt(20)
            else:
                row.height = Pt(16)

            for cell in row.cells:
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
                    pf.line_spacing = 1.0
                    pf.first_line_indent = Pt(0)

                if len(cell.paragraphs) > 1:
                    first_para = cell.paragraphs[0]
                    for i in range(len(cell.paragraphs) - 1, 0, -1):
                        para = cell.paragraphs[i]
                        if para.text.strip() == '':
                            p_element = para._p
                            p_element.getparent().remove(p_element)

                for p in cell.paragraphs:
                    for i in range(len(p.runs) - 1, -1, -1):
                        run = p.runs[i]
                        if run.text.strip() == '':
                            r_element = run._r
                            r_element.getparent().remove(r_element)

        print(f"    清理段落格式完成")
    except Exception as e:
        print(f"    清理段落格式失败: {e}")


def _set_fixed_table_layout(tbl):
    try:
        from docx.oxml import OxmlElement
        tblPr = tbl._tbl.tblPr if tbl._tbl.tblPr is not None else tbl._tbl.get_or_add_tblPr()
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        print(f"    设置固定表格布局完成")
    except Exception as e:
        print(f"    设置固定表格布局失败: {e}")


def _add_border_to_element(parent, border_name, border_type, size, color):
    from docx.oxml import OxmlElement
    border_el = OxmlElement(f'w:{border_name}')
    border_el.set(qn('w:val'), border_type)
    border_el.set(qn('w:sz'), size)
    border_el.set(qn('w:color'), color)
    parent.append(border_el)


def _apply_fixed_table_formatting(tbl):
    """应用固定的表格格式：仿宋_GB2312，小四，居中对齐"""
    try:
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        print("    应用固定表格格式：仿宋_GB2312，小四，居中")

        for row in tbl.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # 设置段落对齐方式为居中
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # 设置段落格式
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.space_before = Pt(0)
                    paragraph_format.space_after = Pt(0)
                    paragraph_format.line_spacing = 1.0

                    # 设置字体格式
                    for run in paragraph.runs:
                        run.font.name = '仿宋_GB2312'
                        run.font.size = Pt(12)  # 小四 = 12磅

                    # 如果段落没有run，创建一个run来设置格式
                    if len(paragraph.runs) == 0:
                        run = paragraph.add_run()
                        run.font.name = '仿宋_GB2312'
                        run.font.size = Pt(12)

        print("    固定表格格式应用完成")
    except Exception as e:
        print(f"    应用固定表格格式失败: {e}")


def insert_table_with_merged_cells(doc: Document, paragraph, excel_data: bytes, title: str):
    """直接从Excel数据创建带合并结构的Word表格，保留原始合并信息（xlrd 优先）"""
    try:
        import tempfile
        import xlrd

        temp_excel = tempfile.mktemp(suffix='.xls')
        with open(temp_excel, 'wb') as f:
            f.write(excel_data)

        wb = xlrd.open_workbook(temp_excel, formatting_info=True)
        ws = wb.sheet_by_index(0)

        max_row = ws.nrows
        max_col = ws.ncols

        print(f"    Excel工作表大小: {max_row}行 x {max_col}列")

        has_data_columns = []
        for col in range(max_col):
            column_has_data = False
            for row in range(max_row):
                cell_value = ws.cell_value(row, col)
                if cell_value is not None and cell_value != '':
                    column_has_data = True
                    break
            if column_has_data:
                has_data_columns.append(col)

        actual_col_count = len(has_data_columns)
        if actual_col_count < max_col:
            print(f"    实际有数据的列数: {actual_col_count} (总列数: {max_col})")
            print(f"    有数据的列: {[c+1 for c in has_data_columns]}")
            max_col = actual_col_count
        else:
            print(f"    所有列都有数据")

        print(f"    合并区域数量: {len(ws.merged_cells)}")
        if ws.merged_cells:
            print(f"    合并区域详情 (gridSpan/vMerge):")
            for i, merged_range in enumerate(ws.merged_cells):
                rlo, rhi, clo, chi = merged_range
                row_span = rhi - rlo
                col_span = chi - clo
                gridSpan = col_span if col_span > 1 else 1
                vMerge = row_span if row_span > 1 else 1
                print(f"      合并区域{i+1}: 位置({rlo+1},{clo+1})-{rhi},{chi}, 行跨度={row_span}, 列跨度={col_span}, gridSpan={gridSpan}, vMerge={vMerge}")
        else:
            print(f"    未检测到合并区域")

        tbl = doc.add_table(rows=max_row, cols=max_col)
        _apply_table_style_best_effort(doc, tbl)

        for row in range(max_row):
            for col_idx, original_col in enumerate(has_data_columns):
                cell_value = ws.cell_value(row, original_col)
                if cell_value is None:
                    cell_value = ''
                else:
                    cell_value = format_cell_value(cell_value)
                tbl.cell(row, col_idx).text = str(cell_value)

        merge_count = 0
        for merged_range in ws.merged_cells:
            rlo, rhi, clo, chi = merged_range
            new_clo = None
            new_chi = None
            for i, original_col in enumerate(has_data_columns):
                if original_col == clo:
                    new_clo = i
                if original_col == chi - 1:
                    new_chi = i
            if new_clo is not None and new_chi is not None:
                try:
                    tbl.cell(rlo, new_clo).merge(tbl.cell(rhi-1, new_chi))
                    merge_count += 1
                    print(f"    合并单元格: ({rlo+1},{clo+1}) -> ({rhi},{chi}) -> 新位置({rlo+1},{new_clo+1}) -> ({rhi},{new_chi+1})")
                except Exception as e:
                    print(f"    合并失败 ({rlo+1},{new_clo+1}) -> ({rhi},{new_chi+1}): {e}")
            else:
                print(f"    跳过合并区域 ({rlo+1},{clo+1}) -> ({rhi},{chi}): 部分列无数据")

        print(f"    成功还原 {merge_count} 个合并区域")

        if merge_count > 0:
            adjusted_merged_cells = []
            for merged_range in ws.merged_cells:
                rlo, rhi, clo, chi = merged_range
                new_clo = None
                new_chi = None
                for i, original_col in enumerate(has_data_columns):
                    if original_col == clo:
                        new_clo = i
                    if original_col == chi - 1:
                        new_chi = i
                if new_clo is not None and new_chi is not None:
                    adjusted_merged_cells.append((rlo, rhi, new_clo, new_chi + 1))
            _apply_merged_cell_borders(tbl, adjusted_merged_cells)
        else:
            _apply_standard_table_borders(tbl)

        # 应用固定格式
        _apply_fixed_table_formatting(tbl)  # 应用固定格式：仿宋_GB2312，小四，居中
        _normalize_row_spacing(tbl)
        _set_fixed_table_layout(tbl)

        paragraph._p.addnext(tbl._tbl)

        try:
            os.remove(temp_excel)
        except:
            pass

        return tbl

    except ImportError:
        print("    xlrd不可用，尝试使用openpyxl")
        return insert_table_with_merged_cells_openpyxl(doc, paragraph, excel_data, title)
    except Exception as e:
        print(f"    xlrd创建合并表格失败: {e}")
        return insert_table_with_merged_cells_openpyxl(doc, paragraph, excel_data, title)


def insert_table_with_merged_cells_openpyxl(doc: Document, paragraph, excel_data: bytes, title: str):
    try:
        import tempfile
        import openpyxl

        temp_excel = tempfile.mktemp(suffix='.xlsx')
        with open(temp_excel, 'wb') as f:
            f.write(excel_data)

        wb = openpyxl.load_workbook(temp_excel, read_only=True, data_only=True)
        ws = wb.active

        max_row = ws.max_row
        max_col = ws.max_column

        print(f"    openpyxl工作表大小: {max_row}行 x {max_col}列")
        print(f"    合并区域数量: {len(ws.merged_cells.ranges)}")

        tbl = doc.add_table(rows=max_row, cols=max_col)
        _apply_table_style_best_effort(doc, tbl)

        for row in range(1, max_row + 1):
            for col in range(1, max_col + 1):
                cell_value = ws.cell(row=row, column=col).value or ''
                if cell_value:
                    cell_value = format_cell_value(cell_value)
                tbl.cell(row-1, col-1).text = str(cell_value)

        merge_count = 0
        for merged_range in ws.merged_cells.ranges:
            min_row, min_col, max_row, max_col = merged_range.min_row, merged_range.min_col, merged_range.max_row, merged_range.max_col
            if min_row != max_row or min_col != max_col:
                try:
                    tbl.cell(min_row-1, min_col-1).merge(tbl.cell(max_row-1, max_col-1))
                    merge_count += 1
                    print(f"    合并单元格: ({min_row},{min_col}) -> ({max_row},{max_col})")
                except Exception as e:
                    print(f"    合并失败 ({min_row},{min_col}) -> ({max_row},{max_col}): {e}")

        print(f"    openpyxl成功还原 {merge_count} 个合并区域")

        if merge_count > 0:
            _apply_openpyxl_merged_cell_borders(tbl, ws.merged_cells.ranges)
        else:
            _apply_standard_table_borders(tbl)

        # 应用固定格式
        _apply_fixed_table_formatting(tbl)  # 应用固定格式：仿宋_GB2312，小四，居中
        _normalize_row_spacing(tbl)
        _set_fixed_table_layout(tbl)

        paragraph._p.addnext(tbl._tbl)

        try:
            os.remove(temp_excel)
        except:
            pass

        wb.close()
        return tbl

    except Exception as e:
        print(f"    openpyxl也失败: {e}")
        return None


def insert_table_after_paragraph(doc: Document, paragraph, matrix: List[List[str]]):
    if not matrix:
        return None
    rows = len(matrix)
    cols = max(len(r) for r in matrix)
    tbl = doc.add_table(rows=rows, cols=cols)
    _apply_table_style_best_effort(doc, tbl)
    for i, row in enumerate(matrix):
        for j, val in enumerate(row):
            tbl.cell(i, j).text = val

    # 应用边框和格式
    _apply_standard_table_borders(tbl)  # 确保为表格添加标准边框
    _apply_fixed_table_formatting(tbl)  # 应用固定格式：仿宋_GB2312，小四，居中
    _normalize_row_spacing(tbl)        # 统一行高和段落格式
    _set_fixed_table_layout(tbl)       # 设置为固定布局，避免格式错乱

    paragraph._p.addnext(tbl._tbl)
    return tbl


# ------------------------------
# 不依赖 qn('o:...') / namespaces 的 OLE 对象查找
# ------------------------------

def iter_ole_objects_in_paragraph(para) -> List[object]:
    """返回段落中所有 o:OLEObject 元素（使用 Clark notation 比对 tag）。"""
    return [el for el in para._p.iter() if el.tag == TAG_OLEOBJECT]


# ------------------------------
# 主流程：就地转换（已加防误处理图片/图表的逻辑）
# ------------------------------

def convert_embedded_excels_inplace(source_docx: str, output_docx: str, placeholder_when_no_pandas: bool = False):
    doc = Document(source_docx)

    rid_to_target = build_rid_to_target_map(source_docx)
    temp_dir = tempfile.mkdtemp()
    try:
        target_to_local = extract_embeddings(source_docx, temp_dir)

        converted = 0
        for para in list(doc.paragraphs):
            objs = iter_ole_objects_in_paragraph(para)
            if not objs:
                continue

            for obj in objs:
                # r:id 可能在 o:OLEObject 元素上
                rid = obj.get(ATTR_RID) or obj.get('r:id')
                if not rid:
                    continue
                target = rid_to_target.get(rid)
                if not target:
                    continue
                local = target_to_local.get(target)
                if not local:
                    continue

                # ✅ 只处理 Excel 工作表 ProgID
                prog = (obj.get('ProgID') or obj.get('ProgId') or obj.get(f'{{{O_NS}}}ProgID') or '').lower()
                if prog and not any(key in prog for key in ALLOWED_EXCEL_PROGIDS):
                    # 例如 Excel.Chart.* / Package / 其它 OLE 一律跳过
                    continue

                # ✅ 严格扩展名过滤
                lower = local.lower()
                is_bin = lower.endswith('.bin')
                is_excel_file = lower.endswith('.xlsx') or lower.endswith('.xls')
                if not (is_bin or is_excel_file):
                    continue

                # ❗ .bin 只能在 olefile 可用时解析；否则直接跳过（避免把图片/图表误转占位表）
                if is_bin:
                    if not OLEFILE_AVAILABLE:
                        # 没有 olefile，跳过 .bin
                        print(f"    跳过 {os.path.basename(local)}：缺少 olefile")
                        continue

                    # 尝试读取 OLE Workbook 流
                    try:
                        ole = olefile.OleFileIO(local)
                        workbook_data = None
                        for entry in ole.listdir():
                            if len(entry) == 1 and entry[0] == 'Workbook':
                                workbook_data = ole.openstream('Workbook').read()
                                break
                        ole.close()
                    except Exception as e:
                        print(f"    解析 OLE 失败，跳过 {os.path.basename(local)}: {e}")
                        continue

                    if workbook_data:
                        tbl = insert_table_with_merged_cells(doc, para, workbook_data, f"OLE Excel - {os.path.basename(local)}")
                        if tbl:
                            print("    使用合并单元格方法成功创建表格")
                        else:
                            print("    合并方法失败（openpyxl/xlrd 均不可用），跳过")
                            continue
                    else:
                        # 没有 Workbook 流 → 不是工作表，跳过
                        print(f"    未找到 Workbook 流，跳过 {os.path.basename(local)}")
                        continue

                else:
                    # 处理 .xlsx / .xls
                    if not PANDAS_AVAILABLE and not placeholder_when_no_pandas:
                        # 不允许占位表时，缺少 pandas 就跳过
                        print("    缺少 pandas，且未允许占位表，跳过一个 Excel 文件")
                        continue
                    matrix, _title = read_excel_to_matrix(local)
                    insert_table_after_paragraph(doc, para, matrix)

                # 移除承载对象的 run：向上找 w:r
                r = obj.getparent()
                while r is not None and r.tag != TAG_RUN:
                    r = r.getparent()
                if r is not None and r.getparent() is not None:
                    r.getparent().remove(r)
                converted += 1

        Path(output_docx).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_docx)
        print(f"就地转换完成：{converted} 个 Excel 对象已替换为 Word 表格")
        return converted
    finally:
        try:
            shutil.rmtree(temp_dir)
        except Exception:
            pass


def main():
    source = 'input/test.docx'
    output = 'output/test_converted_inplace.docx'
    if not Path(source).exists():
        print(f'源文件不存在: {source}')
        return
    # 如需允许在无 pandas 时放置占位表，把 placeholder_when_no_pandas=True
    convert_embedded_excels_inplace(source, output, placeholder_when_no_pandas=True)
    print(f'输出: {output}')


if __name__ == '__main__':
    main()
