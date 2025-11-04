from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import re

# =========================
# 直接指定文件路径（自行修改）
# =========================
input_path = r"C:\Users\yanha\Desktop\原始文档.docx"
output_path = r"C:\Users\yanha\Desktop\格式化后.docx"

# 匹配模式：库号 + 冒号(中或英) + 可选空格 + 编号（字母数字）
pattern = re.compile(r'(库号[:：]\s*)([A-Za-z0-9]+)')

# 字体设置
LABEL_FONT_NAME = 'SimHei'            # 黑体
LABEL_FONT_EAST_ASIA = '黑体'
CODE_FONT_NAME = 'Times New Roman'    # 英文字体
FONT_SIZE_PT = 16                     # 三号 ≈ 16pt


def style_run(run, font_name, east_asia_name=None, size_pt=16, bold=True):
    """设置 run 的字体样式"""
    font = run.font
    font.name = font_name
    if east_asia_name:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia_name)
    font.size = Pt(size_pt)
    font.bold = bold


def process_paragraph(paragraph):
    """查找并格式化库号"""
    text = paragraph.text
    matches = list(pattern.finditer(text))
    if not matches:
        return

    pieces = []
    last_idx = 0
    for m in matches:
        if m.start() > last_idx:
            pieces.append(('text', text[last_idx:m.start()]))
        pieces.append(('label', m.group(1)))
        pieces.append(('code', m.group(2)))
        last_idx = m.end()
    if last_idx < len(text):
        pieces.append(('text', text[last_idx:]))

    # 清空原有 runs
    for r in paragraph.runs:
        r.text = ''

    # 按 pieces 重新生成 runs
    for kind, content in pieces:
        run = paragraph.add_run(content)
        if kind == 'label':
            style_run(run, LABEL_FONT_NAME, LABEL_FONT_EAST_ASIA, FONT_SIZE_PT, True)
        elif kind == 'code':
            style_run(run, CODE_FONT_NAME, None, FONT_SIZE_PT, True)
        # 其他文本保持默认格式


def iterate_cells(doc):
    """遍历所有表格单元格中的段落"""
    for table in doc.tables:
        for row in table.rows:
            try:
                cells = row.cells
            except ValueError:
                # python-docx 在处理缺失 vMerge 起点的表格时会抛错，这里直接跳过该行
                continue
            for cell in cells:
                for p in cell.paragraphs:
                    yield p


def iterate_headers_and_footers(doc):
    """遍历页眉页脚中的段落"""
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for p in section.footer.paragraphs:
            yield p


def process_document(input_path, output_path):
    """处理整个文档"""
    doc = Document(input_path)
    for p in doc.paragraphs:
        process_paragraph(p)
    for p in iterate_cells(doc):
        process_paragraph(p)
    for p in iterate_headers_and_footers(doc):
        process_paragraph(p)
    doc.save(output_path)
    print(f"格式化完成：{output_path}")


if __name__ == "__main__":
    process_document(input_path, output_path)
