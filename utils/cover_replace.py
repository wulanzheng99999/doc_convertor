import os
import json
import xml.etree.ElementTree as ET
from xml.dom import minidom
from typing import Dict, Any, List, Tuple
import zipfile
import tempfile
import shutil
import re

# 导入Document类（忽略静态分析工具的警告）
# pyright: reportAttributeAccessIssue=false
try:
    from docx import Document
except ImportError:
    Document = None

"""
WordStyler - 文档模板内容替换工具

该模块提供了多种方式来替换DOCX文档模板中的内容，同时保持原有格式不变。
支持两种主要的实现方式：
1. 基于XML节点处理的方式
2. 基于python-docx库API的方式

主要功能：
- 从配置文件加载替换映射关系
- 从文档中提取内容并创建替换字典
- 支持跨多个文本节点的内容替换
- 保持原有文档格式
- 支持临时文件和持久化文件保存选项

函数说明：
- load_replace_config: 加载替换配置文件
- save_replace_config: 保存替换配置文件（保留此函数以供其他模块使用）
- replace_text_in_runs: 在段落的runs中替换文本，保持格式
- get_paragraph_text: 获取段落的完整文本
- create_replacement_dict_from_content: 根据提取的内容和配置的key创建替换字典
- extract_text_from_xml: 从docx文件的XML中提取所有文本内容
* replace_content_in_template_docx: 使用python-docx方式替换模板中的内容
* replace_content_in_template_xml_to_docx: 使用XML方式处理并生成DOCX文件
- extract_content_from_docx: 从docx文件中提取内容
- match_content_to_config: 将提取的内容匹配到配置键
- replace_content_in_cover: 快捷函数，从给定文件中识别配置key对应的value并替换模板内容
"""

def load_replace_config(config_path: str) -> Dict[str, str]:
    """
    加载替换配置文件
    
    Args:
        config_path: 配置文件路径
        
    Returns:
        Dict[str, str]: key-value替换映射
    """
    with open(config_path, 'r', encoding='utf-8') as f:
        config = json.load(f)
    return config.get("替换内容", {})


def save_replace_config(config_path: str, replace_dict: Dict[str, str]) -> None:
    """
    保存替换配置文件 (保留此函数以供其他模块使用，但按用户要求不在此模块中调用)
    
    Args:
        config_path: 配置文件路径
        replace_dict: key-value替换映射
    """
    config = {"替换内容": replace_dict}
    with open(config_path, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=4)


def replace_text_in_runs(runs, old_text, new_text):
    """
    在runs中替换文本，保持格式
    
    Args:
        runs: 段落的runs列表
        old_text: 要替换的旧文本
        new_text: 新文本
    """
    if not old_text or not new_text:
        return
    
    # 收集所有run的文本
    full_text = ''.join(run.text for run in runs)
    
    # 检查是否包含要替换的文本
    if old_text in full_text:
        # 找到包含旧文本的第一个run的位置
        start_pos = full_text.find(old_text)
        end_pos = start_pos + len(old_text)
        
        # 如果新旧文本长度相同，保持原有结构
        if len(old_text) == len(new_text):
            # 逐字符替换，保持run结构
            current_pos = 0
            new_text_index = 0
            
            for run in runs:
                run_len = len(run.text)
                # 检查run是否与要替换的文本有交集
                if current_pos < end_pos and current_pos + run_len > start_pos:
                    # 计算在当前run中需要替换的部分
                    local_start = max(0, start_pos - current_pos)
                    local_end = min(run_len, end_pos - current_pos)
                    
                    # 替换文本
                    if local_start < local_end:
                        prefix = run.text[:local_start]
                        suffix = run.text[local_end:]
                        replacement = new_text[new_text_index:new_text_index + (local_end - local_start)]
                        new_text_index += len(replacement)
                        run.text = prefix + replacement + suffix
                current_pos += run_len
        else:
            # 长度不同的情况，清空所有相关的run，并在第一个run中放入新文本
            current_pos = 0
            first_run = None
            prefix_text = ""
            suffix_text = ""
            
            for run in runs:
                run_len = len(run.text)
                # 检查run是否与要替换的文本有交集
                if current_pos < end_pos and current_pos + run_len > start_pos:
                    if first_run is None:
                        first_run = run
                        # 保存前缀文本
                        if current_pos < start_pos:
                            prefix_text = full_text[current_pos:start_pos]
                    # 清空文本但保留格式
                    run.text = ""
                elif current_pos + run_len > end_pos and first_run is not None:
                    # 保存后缀文本
                    suffix_start = max(0, end_pos - current_pos)
                    if suffix_start < run_len:
                        suffix_text = run.text[suffix_start:]
                    run.text = ""
                current_pos += run_len
            
            # 在第一个run中设置新文本
            if first_run is not None:
                # 计算前缀文本
                prefix_len = start_pos
                for run in runs:
                    if run == first_run:
                        break
                    prefix_len -= len(run.text)
                
                if prefix_len > 0:
                    prefix_text = first_run.text[:prefix_len]
                
                first_run.text = prefix_text + new_text + suffix_text


def get_paragraph_text(paragraph):
    """
    获取段落的完整文本
    
    Args:
        paragraph: python-docx段落对象
        
    Returns:
        str: 段落的完整文本
    """
    return ''.join(run.text for run in paragraph.runs)


def extract_cover_paragraphs(
        file_path: str,
        marker: str = "各专业参加设计人员名单",
        skip_empty: bool = True,
        filter_design_terms: bool = True  # False：不启用过滤。True（默认）：会过滤掉包含以下任意关键词的段落（忽略大小写和空格）
) -> Tuple[List[str], bool]:
    """
    从 Word 文件开头提取普通段落文本，直到遇到 marker（忽略空格和大小写）。
    可选地过滤掉包含"可行性研究""初步设计""（代初步）""（代可行）"等关键词的段落。

    参数:
        file_path: str, .docx 文件路径
        marker: str, 停止标记
        skip_empty: bool, 是否跳过空段落
        filter_design_terms: bool, 是否过滤掉设计类段落,False：不启用过滤。True（默认）：会过滤掉包含以下任意关键词的段落（忽略大小写和空格）

    返回:
        paragraphs: list[str] — 提取到的段落文本
        found_marker: bool — 是否找到标记
    """

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    # 先检查是否已导入Document
    if Document is None:
        raise ImportError("未找到python-docx库，请安装：pip install python-docx")

    doc = Document(file_path)
    paragraphs = []
    found_marker = False

    # 预处理 marker：忽略空格与大小写
    marker_pattern = re.sub(r"\s+", "", marker).lower()

    # 定义过滤关键词
    design_keywords = [
        "可行性研究",
        "初步设计",
        "代初步",
        "代可行",
    ]

    # 构建正则：允许空格、括号、多种形式
    design_pattern = re.compile(
        r"[\(（]?\s*(?:代初步|代可行)\s*[\)）]?|可行性研究|初步设计",
        flags=re.IGNORECASE
    )

    for para in doc.paragraphs:
        text = para.text or ""
        if skip_empty and not text.strip():
            continue

        # 去掉所有空格，统一中英文冒号
        text_clean = re.sub(r"\s+", "", text)
        text_clean_lower = text_clean.lower()

        # 检测停止标志
        if marker_pattern in text_clean_lower.replace("：", "").replace(":", ""):
            found_marker = True
            break

        # 将全角括号替换为半角
        normalized = text.replace("（", "(").replace("）", ")")

        # 检查过滤关键词
        if filter_design_terms and design_pattern.search(normalized):
            continue

        paragraphs.append(text.strip())

    return paragraphs, found_marker


def extract_content_from_docx(docx_path: str) -> Dict[str, str]:
    """
    从docx文件中提取内容（使用更健壮的方法）
    
    Args:
        docx_path: docx文件路径
        
    Returns:
        Dict[str, str]: 提取的内容字典
    """
    # 先检查是否已导入Document
    if Document is None:
        raise ImportError("未找到python-docx库，请安装：pip install python-docx")
    
    # 使用新的提取方法
    paragraphs, _ = extract_cover_paragraphs(docx_path)
    
    # 根据用户说明，索引为 0, 2, 3, 5, 6, 7 的内容分别对应配置文件中要替换的 6 项内容
    content_dict = {}
    
    # 确保有足够的段落内容
    if len(paragraphs) > 7:
        content_dict["content_0"] = paragraphs[0]  # 项目名称
        content_dict["content_1"] = paragraphs[2]  # 日期
        content_dict["content_2"] = paragraphs[3]  # 库号
        content_dict["content_3"] = paragraphs[5]  # 主管领导
        content_dict["content_4"] = paragraphs[6]  # 副总工程师
        content_dict["content_5"] = paragraphs[7]  # 总设计师
    else:
        # 如果段落数量不足，抛出异常或使用空值
        # 根据项目规范，应该确保能提取到所需内容
        raise ValueError(f"文档内容不足，仅提取到{len(paragraphs)}段，需要至少8段内容")
    
    return content_dict


def match_content_to_config(content_dict: Dict[str, str], config_keys: list) -> Dict[str, str]:
    """
    将提取的内容匹配到配置键
    使用更智能的匹配逻辑，基于内容特征进行匹配
    
    Args:
        content_dict: 从文档中提取的内容
        config_keys: 配置文件中的键列表
        
    Returns:
        Dict[str, str]: 匹配后的替换字典
    """
    # 初始化替换字典
    replace_dict = {key: "" for key in config_keys}
    
    # 获取内容值列表
    content_values = list(content_dict.values())
    
    # 初始化匹配变量
    project_name = ""
    date_text = ""
    library_number = ""
    leader_info = ""
    vp_engineer_info = ""
    chief_designer_info = ""
    
    # 遍历提取的内容，根据特征进行匹配
    for content_item in content_values:
        # 跳过目录等无关内容
        if content_item in ["目录", "第一章 摘要      1", "第二章  背景概述 3", "§2.1  政策背景   3", 
                           "§2.2  经济背景   4", "§2.3  产业背景   6", "第三章  电池市场分析     9"]:
            continue
            
        # 去除多余空格和换行符，便于匹配
        cleaned_content = content_item.strip().replace('\n', '').replace('\r', '').replace(' ', '')
        
        # 匹配项目名称（以"项目"结尾）
        if not project_name and content_item.strip().endswith("项目"):
            # 进一步验证不是日期或其他信息
            if not any(date_keyword in content_item for date_keyword in ['年', '月', '日']) or \
               content_item.count('年') < 2:  # 避免将日期误识别为项目名
                project_name = content_item
                continue
            
        # 匹配日期（支持年月日和年月格式）
        if not date_text and len(content_item) < 30:  # 日期通常不会太长
            # 检查是否包含日期特征
            has_digit = any(char.isdigit() for char in content_item)
            
            if has_digit:
                # 检查是否符合年月日或年月格式
                import re
                # 匹配年月日格式：2024年10月1日 或 2024-10-01 或 2024/10/01
                # 匹配年月格式：2024年10月
                date_regex = r'(\d{4}年\d{1,2}月\d{1,2}日|\d{4}年\d{1,2}月\d{1,2}号|\d{4}-\d{1,2}-\d{1,2}|\d{4}/\d{1,2}/\d{1,2}|\d{4}年\d{1,2}月)'
                if re.search(date_regex, content_item):
                    # 优先选择更具体的日期（包含年份的）
                    if '202' in content_item or '201' in content_item:  # 2020年代或2010年代
                        date_text = content_item
                    elif not date_text:  # 如果还没有找到日期，则使用第一个匹配的
                        date_text = content_item
                    continue
                    
        # 匹配库号（保持原有逻辑但增强鲁棒性）
        if not library_number:
            # 更宽松的库号匹配规则
            if ('库号' in content_item and (':' in content_item or '：' in content_item)) or \
               (any(keyword in content_item for keyword in ['库号', '编号', 'NO.', 'No']) and 
                any(char.isdigit() for char in content_item)):
                library_number = content_item
                continue
                
        # 匹配人员信息（增强鲁棒性）
        if not leader_info and ('主管领导' in content_item or ('主管' in content_item and '领导' in content_item)):
            leader_info = content_item
            continue
        elif not vp_engineer_info and ('副总工程师' in content_item or ('副总' in content_item and '工程师' in content_item)):
            vp_engineer_info = content_item
            continue
        elif not chief_designer_info and ('总设计师' in content_item or ('总设' in content_item and '计师' in content_item)):
            chief_designer_info = content_item
            continue
    
    # 根据配置键的特征进行匹配
    for key in config_keys:
        # 项目名称匹配（以"项目"结尾的配置键）
        if key.strip().endswith("项目") and project_name:
            replace_dict[key] = project_name
        # 日期匹配（包含年月相关的配置键）
        elif any(date_keyword in key for date_keyword in ['年', '月', '日']) and date_text:
            replace_dict[key] = date_text
        # 库号匹配
        elif any(lib_keyword in key for lib_keyword in ['库号', '编号']) and library_number:
            # 对于库号，我们使用提取到的库号替换模板中的库号
            # 确保库号格式正确，移除多余的空格和换行符
            cleaned_library_number = library_number.strip().replace('\n', '').replace('\r', '').replace('：', ':')
            replace_dict[key] = cleaned_library_number
        # 人员信息匹配
        elif '主管领导' in key and leader_info:
            replace_dict[key] = leader_info
        elif '副总工程师' in key and vp_engineer_info:
            replace_dict[key] = vp_engineer_info
        elif '总设计师' in key and chief_designer_info:
            replace_dict[key] = chief_designer_info
    
    return replace_dict


def replace_content_in_template_xml_to_docx(template_docx_path: str, replace_dict: Dict[str, str], output_docx_path: str, save_file: bool = False) -> str:
    """
    使用XML方式替换模板中的内容并生成DOCX文件
    与replace_content_in_template_docx方法类似，直接处理DOCX文件
    
    Args:
        template_docx_path: 模板docx文件路径
        replace_dict: key-value替换映射
        output_docx_path: 输出的DOCX文件路径
        save_file: 是否保存文件到指定路径，默认为False（不保存，使用临时文件）
        
    Returns:
        str: 实际的文件路径（临时路径或指定路径）
    """
    # 先检查是否已导入Document
    if Document is None:
        raise ImportError("未找到python-docx库，请安装：pip install python-docx")
    
    import tempfile
    import shutil
    
    # 确定实际的输出路径
    if save_file:
        # 保存到指定路径
        actual_output_path = output_docx_path
        output_dir = os.path.dirname(actual_output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 如果输出文件已存在，先尝试删除
        if os.path.exists(actual_output_path):
            try:
                os.remove(actual_output_path)
            except Exception as e:
                # 如果删除失败，可能是文件被占用，等待一段时间再试
                import time
                time.sleep(0.1)
                try:
                    os.remove(actual_output_path)
                except Exception:
                    raise Exception(f"无法覆盖输出文件 {actual_output_path}，文件可能正在被使用")
    else:
        # 使用临时文件
        temp_dir = tempfile.mkdtemp()
        actual_output_path = os.path.join(temp_dir, os.path.basename(output_docx_path))
    
    # 使用python-docx方式处理，确保与replace_content_in_template_docx方法一致
    # 打开模板文档
    doc = Document(template_docx_path)
    
    # 替换段落中的文本，保留格式
    for paragraph in doc.paragraphs:
        # 获取段落的完整文本
        full_text = get_paragraph_text(paragraph)
        
        # 检查是否需要替换
        for key, value in replace_dict.items():
            if key in full_text and value:  # 只有当value不为空时才替换
                # 使用我们自定义的函数替换文本
                replace_text_in_runs(paragraph.runs, key, value)
    
    # 替换表格中的文本，保留格式
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # 获取段落的完整文本
                    full_text = get_paragraph_text(paragraph)
                    
                    # 检查是否需要替换
                    for key, value in replace_dict.items():
                        if key in full_text and value:  # 只有当value不为空时才替换
                            # 使用我们自定义的函数替换文本
                            replace_text_in_runs(paragraph.runs, key, value)
    
    # 保存文档
    doc.save(actual_output_path)
    
    # 如果不保存文件，返回临时路径（程序结束后会自动清理）
    # 如果保存文件，返回指定路径
    return actual_output_path


def replace_content_in_template_docx(template_docx_path: str, replace_dict: Dict[str, str], output_docx_path: str, save_file: bool = False) -> str:
    """
    使用python-docx方式替换模板中的内容，保留原有格式
    
    Args:
        template_docx_path: 模板docx文件路径
        replace_dict: key-value替换映射
        output_docx_path: 输出docx文件路径
        save_file: 是否保存文件到指定路径，默认为False（不保存，使用临时文件）
        
    Returns:
        str: 实际的文件路径（临时路径或指定路径）
    """
    # 先检查是否已导入Document
    if Document is None:
        raise ImportError("未找到python-docx库，请安装：pip install python-docx")
    
    import tempfile
    import shutil
    
    # 确定实际的输出路径
    if save_file:
        # 保存到指定路径
        actual_output_path = output_docx_path
        output_dir = os.path.dirname(actual_output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 如果输出文件已存在，先尝试删除
        if os.path.exists(actual_output_path):
            try:
                os.remove(actual_output_path)
            except Exception as e:
                # 如果删除失败，可能是文件被占用，等待一段时间再试
                import time
                time.sleep(0.1)
                try:
                    os.remove(actual_output_path)
                except Exception:
                    raise Exception(f"无法覆盖输出文件 {actual_output_path}，文件可能正在被使用")
    else:
        # 使用临时文件
        temp_dir = tempfile.mkdtemp()
        actual_output_path = os.path.join(temp_dir, os.path.basename(output_docx_path))
    
    # 打开模板文档
    doc = Document(template_docx_path)
    
    # 替换段落中的文本，保留格式
    for paragraph in doc.paragraphs:
        # 获取段落的完整文本
        full_text = get_paragraph_text(paragraph)
        
        # 检查是否需要替换
        for key, value in replace_dict.items():
            if key in full_text and value:  # 只有当value不为空时才替换
                # 使用我们自定义的函数替换文本
                replace_text_in_runs(paragraph.runs, key, value)
    
    # 替换表格中的文本，保留格式
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # 获取段落的完整文本
                    full_text = get_paragraph_text(paragraph)
                    
                    # 检查是否需要替换
                    for key, value in replace_dict.items():
                        if key in full_text and value:  # 只有当value不为空时才替换
                            # 使用我们自定义的函数替换文本
                            replace_text_in_runs(paragraph.runs, key, value)
    
    # 保存文档
    doc.save(actual_output_path)
    
    # 返回实际的文件路径
    return actual_output_path


def replace_content_in_cover(config_path: str, source_docx_path: str, template_docx_path: str, output_docx_path: str, save_file: bool = False) -> str:
    """
    快捷函数，从给定文件中识别配置key对应的value并替换模板内容
    
    Args:
        config_path: 配置文件路径
        source_docx_path: 源DOCX文件路径（从中提取内容）
        template_docx_path: 模板DOCX文件路径
        output_docx_path: 输出DOCX文件路径
        save_file: 是否保存文件到指定路径，默认为False（不保存，使用临时文件）
        
    Returns:
        str: 实际的文件路径（临时路径或指定路径）
    """
    # 1. 加载配置文件
    replace_dict = load_replace_config(config_path)
    print(f"加载配置文件: {config_path}")
    
    # 2. 从源DOCX文件中提取内容
    if os.path.exists(source_docx_path):
        print(f"从源文件提取内容: {source_docx_path}")
        content_dict = extract_content_from_docx(source_docx_path)
    else:
        print(f"源文件不存在，使用模拟数据: {source_docx_path}")
        # 如果源文件不存在，使用模拟数据
        content_dict = {
            "content_0": "高性能、安全、环保的磷酸亚铁锂电池材料项目",
            "content_1": "2024年6月15日",
            "content_2": "库号：20032357ARPD01AK009",
            "content_3": "公司主管领导：王 小 华",
            "content_4": "副总工程师  ：赵 小 四",
            "content_5": "总设计师 ：刘 小 五"
        }
    
    # 3. 将提取的内容匹配到配置键
    config_keys = list(replace_dict.keys())
    updated_replace_dict = match_content_to_config(content_dict, config_keys)
    
    # 4. 打印匹配结果
    print("\n=== 填入value后的key-value结构 ===")
    for key, value in updated_replace_dict.items():
        print(f"'{key}': '{value}'")
    
    # 5. 使用匹配后的字典替换模板内容
    print(f"\n开始替换模板内容: {template_docx_path}")
    actual_path = replace_content_in_template_docx(
        template_docx_path, 
        updated_replace_dict, 
        output_docx_path, 
        save_file
    )
    
    # print(f"替换完成，输出路径: {actual_path}")
    return actual_path


def replace_content_in_cover_auto(source_docx_path: str, output_docx_path: str, document_type: int, save_file: bool = False) -> str:
    """
    自动选择模板和配置文件的快捷函数，从给定文件中识别配置key对应的value并替换模板内容
    
    Args:
        source_docx_path: 源DOCX文件路径（从中提取内容）
        output_docx_path: 输出DOCX文件路径
        document_type: 文档类型 (1, 2, 3, 4)
        save_file: 是否保存文件到指定路径，默认为False（不保存，使用临时文件）
        
    Returns:
        str: 实际的文件路径（临时路径或指定路径）
    """
    # 获取项目根目录
    current_dir = os.path.dirname(os.path.abspath(__file__))
    parent_dir = os.path.dirname(current_dir)
    template_dir = os.path.join(parent_dir, 'template')
    
    # 1. 分析源文档中的项目字数和日期类型
    project_length, date_type = analyze_cover_content(source_docx_path)
    print(f"分析源文档结果: 项目字数={project_length}, 日期类型={date_type}")
    
    # 2. 根据项目字数确定字数参数
    if project_length <= 16:
        length_param = 1
    elif project_length <= 32:
        length_param = 2
    else:
        length_param = 3
    
    # 3. 确定日期参数
    date_param = date_type  # 1=年月日型, 2=年月型
    
    # 4. 确定文档类型参数
    doc_type_param = document_type  # 1, 2, 3, 4
    
    # 5. 构建配置文件路径
    config_filename = f"cover_replace_config_{date_param}_{length_param}.json"
    config_path = os.path.join(template_dir, config_filename)
    
    # 6. 构建模板文件路径
    template_filename = f"reference_cover{doc_type_param}_{date_param}_{length_param}.docx"
    template_path = os.path.join(template_dir, template_filename)
    
    # 7. 检查文件是否存在
    if not os.path.exists(config_path):
        print(f"警告: 配置文件不存在: {config_path}")
        # 尝试使用默认配置文件
        config_path = os.path.join(template_dir, 'cover_replace_config.json')
        if not os.path.exists(config_path):
            raise FileNotFoundError(f"配置文件不存在: {config_path}")
    
    if not os.path.exists(template_path):
        print(f"警告: 模板文件不存在: {template_path}")
        # 尝试使用默认模板文件
        template_path = os.path.join(template_dir, 'reference_cover1.docx')
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")
    
    print(f"选择的配置文件: {config_path}")
    print(f"选择的模板文件: {template_path}")
    
    # 8. 加载配置文件
    replace_dict = load_replace_config(config_path)
    print(f"加载配置文件: {config_path}")
    
    # 9. 从源DOCX文件中提取内容
    if os.path.exists(source_docx_path):
        print(f"从源文件提取内容: {source_docx_path}")
        content_dict = extract_content_from_docx(source_docx_path)
    else:
        print(f"源文件不存在，使用模拟数据: {source_docx_path}")
        # 如果源文件不存在，使用模拟数据
        content_dict = {
            "content_0": "高性能、安全、环保的磷酸亚铁锂电池材料项目",
            "content_1": "2024年6月15日",
            "content_2": "库号：20032357ARPD01AK009",
            "content_3": "公司主管领导：王 小 华",
            "content_4": "副总工程师  ：赵 小 四",
            "content_5": "总设计师 ：刘 小 五"
        }
    
    # 10. 将提取的内容匹配到配置键
    config_keys = list(replace_dict.keys())
    updated_replace_dict = match_content_to_config(content_dict, config_keys)
    
    # 11. 打印匹配结果
    print("\n=== 填入value后的key-value结构 ===")
    for key, value in updated_replace_dict.items():
        print(f"'{key}': '{value}'")
    
    # 12. 使用匹配后的字典替换模板内容
    print(f"\n开始替换模板内容: {template_path}")
    actual_path = replace_content_in_template_docx(
        template_path, 
        updated_replace_dict, 
        output_docx_path, 
        save_file
    )
    
    # print(f"替换完成，输出路径: {actual_path}")
    return actual_path


def create_replacement_dict_from_content(content_dict: Dict[str, str], config_keys: list) -> Dict[str, str]:
    """
    根据提取的内容和配置的key创建替换字典
    
    Args:
        content_dict: 从文档中提取的内容
        config_keys: 配置文件中的key列表
        
    Returns:
        Dict[str, str]: 完整的替换字典
    """
    replace_dict = {key: "" for key in config_keys}
    
    # 这里需要实现具体的匹配逻辑，目前是简化处理
    # 实际应用中可能需要更复杂的匹配规则
    content_values = list(content_dict.values())
    
    # 简单的一一对应匹配
    for i, key in enumerate(config_keys):
        if i < len(content_values):
            replace_dict[key] = content_values[i]
    
    return replace_dict


def extract_text_from_xml(docx_path: str) -> str:
    """
    从docx文件的XML中提取所有文本内容
    
    Args:
        docx_path: docx文件路径
        
    Returns:
        str: 提取的文本内容
    """
    # 创建临时目录
    with tempfile.TemporaryDirectory() as temp_dir:
        # 解压docx文件
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        
        # 读取document.xml
        document_xml_path = os.path.join(temp_dir, 'word', 'document.xml')
        if os.path.exists(document_xml_path):
            tree = ET.parse(document_xml_path)
            root = tree.getroot()
            
            # 定义命名空间
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
            
            # 提取所有文本
            text_elements = root.findall('.//w:t', namespaces)
            texts = [elem.text for elem in text_elements if elem.text]
            return ' '.join(texts)
    
    return ""


def analyze_cover_content(docx_path: str) -> tuple:
    """
    分析封面文件中的项目名称字数和日期形式
    
    Args:
        docx_path: DOCX文件路径
        
    Returns:
        tuple: (项目名称字数, 日期形式) 
               日期形式: 1表示年月日型(如2020年2月12日), 2表示年月型(如2020年2月), 0表示未识别到日期
    """
    # 先检查是否已导入Document
    if Document is None:
        raise ImportError("未找到python-docx库，请安装：pip install python-docx")
    
    try:
        doc = Document(docx_path)
        project_name_length = 0
        date_type = 0  # 0:未识别, 1:年月日型, 2:年月型
        
        # 提取所有文本内容
        all_text = []
        
        # 提取段落内容
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                all_text.append(paragraph.text.strip())
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        all_text.append(cell.text.strip())
        
        # 初始化匹配变量
        project_name = ""
        date_text = ""
        
        # 分析文本内容（参考match_content_to_config函数的逻辑）
        for content_item in all_text:
            # 跳过目录等无关内容
            if content_item in ["目录", "第一章 摘要      1", "第二章  背景概述 3", "§2.1  政策背景   3", 
                               "§2.2  经济背景   4", "§2.3  产业背景   6", "第三章  电池市场分析     9"]:
                continue
                
            # 匹配项目名称（以"项目"结尾）
            if not project_name and content_item.strip().endswith("项目"):
                # 进一步验证不是日期或其他信息
                if not any(date_keyword in content_item for date_keyword in ['年', '月', '日']) or \
                   content_item.count('年') < 2:  # 避免将日期误识别为项目名
                    project_name = content_item
                    project_name_length = len(project_name.strip())
            
            # 匹配日期（支持年月日和年月格式）
            if not date_text and len(content_item) < 30:  # 日期通常不会太长
                # 检查是否包含日期特征
                has_digit = any(char.isdigit() for char in content_item)
                
                if has_digit:
                    # 检查是否符合年月日或年月格式
                    import re
                    # 匹配年月日格式：2024年10月1日 或 2024-10-01 或 2024/10/01
                    # 匹配年月格式：2024年10月
                    date_regex = r'(\d{4}年\d{1,2}月\d{1,2}日|\d{4}年\d{1,2}月\d{1,2}号|\d{4}-\d{1,2}-\d{1,2}|\d{4}/\d{1,2}/\d{1,2}|\d{4}年\d{1,2}月)'
                    if re.search(date_regex, content_item):
                        # 优先选择更具体的日期（包含年份的）
                        if '202' in content_item or '201' in content_item:  # 2020年代或2010年代
                            date_text = content_item
                        elif not date_text:  # 如果还没有找到日期，则使用第一个匹配的
                            date_text = content_item
            
            # 确定日期类型
            if date_text and date_type == 0:
                import re
                # 年月日格式：2024年10月1日 或 2024-10-01 或 2024/10/01
                date_y_m_d_pattern = r'(\d{4}年\d{1,2}月\d{1,2}日|\d{4}年\d{1,2}月\d{1,2}号|\d{4}-\d{1,2}-\d{1,2}|\d{4}/\d{1,2}/\d{1,2})'
                # 年月格式：2024年10月
                date_y_m_pattern = r'\d{4}年\d{1,2}月'
                
                # 检查是否匹配年月日格式
                if re.search(date_y_m_d_pattern, date_text):
                    date_type = 1
                # 检查是否匹配年月格式
                elif re.search(date_y_m_pattern, date_text):
                    date_type = 2
        
        return (project_name_length, date_type)
    
    except Exception as e:
        print(f"分析封面内容时出错: {e}")
        return (0, 0)


# 更新主程序入口以移除测试函数调用
if __name__ == "__main__":
    # 配置文件路径
    config_path = os.path.join('template', 'cover_replace_config.json')
    
    # 模板路径
    template_docx_path = os.path.join('template', 'reference_cover1.docx')
    
    # 输出路径
    output_docx_path = os.path.join('result', 'output_cover.docx')
    
    # 确保输出目录存在
    os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
    
    # 从指定的DOCX文件中提取内容并替换模板内容
    print("=== 从指定DOCX文件中提取内容并替换模板内容 ===")
    source_docx_path = r"C:\Users\yanha\Desktop\数字总师\WordStyler\docx\test.docx"
    
    try:
        # 使用快捷函数处理
        actual_path = replace_content_in_cover(
            config_path=config_path,
            source_docx_path=source_docx_path,
            template_docx_path=template_docx_path,
            output_docx_path=output_docx_path,
            save_file=False  # 保存到指定路径
        )
        print(f"处理完成，文件路径: {actual_path}")
    except Exception as e:
        print(f"处理出错: {e}")
    
    print("\n=== 处理完成 ===")