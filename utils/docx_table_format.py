import sys
import os
import copy
from pathlib import Path
from docx import Document
from docx.table import _Cell, Table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.document import Document as DocumentClass
from typing import Dict, List, Tuple

# 导入Excel就地转换功能
try:
    from enhanced_table_converter import convert_embedded_excels_inplace
    EXCEL_CONVERTER_AVAILABLE = True
except ImportError:
    print("⚠️ 警告: 无法导入Excel就地转换功能，将跳过Excel对象转换步骤")
    EXCEL_CONVERTER_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

# --- 配置 (保持不变) ---
TEMPLATE_DOC_PATH = "input/reference_content.docx"
SOURCE_DOC_PATH = "input/test_excel.docx"
OUTPUT_DOC_PATH = "output/formatted_document.docx"
CLEANED_DOC_PATH = "output/cleaned_document.docx"
# Excel就地转换配置
EXCEL_CONVERTED_DOC_PATH = "output/excel_converted_intermediate.docx"  # Excel转换后的中间文件


def convert_excel_objects_first(source_docx: str, output_docx: str) -> str:
    """
    第一步：执行Excel对象就地转换
    返回转换后的文档路径
    """
    print("\n" + "=" * 60)
    print("🔄 步骤0: Excel对象就地转换")
    print("=" * 60)

    if not EXCEL_CONVERTER_AVAILABLE:
        print("⚠️ Excel转换功能不可用，跳过此步骤")
        return source_docx

    if not Path(source_docx).exists():
        print(f"❌ 源文件不存在: {source_docx}")
        return source_docx

    print(f"📄 输入文件: {source_docx}")
    print(f"📄 输出文件: {output_docx}")

    try:
        # 创建输出目录
        Path(output_docx).parent.mkdir(parents=True, exist_ok=True)

        # 执行Excel就地转换
        converted_count = convert_embedded_excels_inplace(
            source_docx,
            output_docx,
            placeholder_when_no_pandas=True
        )

        if converted_count > 0:
            print(f"✅ Excel对象转换完成: {converted_count} 个对象已转换为Word表格")
            print(f"📁 转换后文件: {output_docx}")
            return output_docx
        else:
            print("ℹ️ 未发现Excel对象或转换失败，使用原文件")
            # 如果没有转换任何对象，删除可能生成的空文件
            if Path(output_docx).exists():
                os.remove(output_docx)
            return source_docx

    except Exception as e:
        print(f"❌ Excel对象转换失败: {e}")
        print("ℹ️ 将使用原文件继续处理")
        return source_docx


class TableStyleBrush:
    """
    一个用于将一个文档中表格的样式应用到另一个文档表格的工具类。
    支持根据页面方向（横向/纵向）进行匹配。
    """

    def __init__(self, template_document_path: str):
        """
        使用模板文档路径初始化样式刷。
        """
        try:
            self.template_doc = Document(template_document_path)
            if not self.template_doc.tables:
                raise ValueError("模板文件中未找到任何表格。")
            print(f"成功加载模板文件: '{template_document_path}'")

            print("正在扫描模板文件，按页面方向对表格进行分类...")
            categorized_templates = self._categorize_tables_by_orientation(self.template_doc)
            self.template_vertical_tables = categorized_templates['vertical']
            self.template_horizontal_tables = categorized_templates['horizontal']

            if not self.template_vertical_tables and not self.template_horizontal_tables:
                raise ValueError("在模板文件中未能找到任何可用的表格模板。")

            print("=== 成功建立模板库 ===")
            print(f"    - 找到 {len(self.template_vertical_tables)} 个纵向页面表格模板")
            print(f"    - 找到 {len(self.template_horizontal_tables)} 个横向页面表格模板")

            print("\n\n" + "=" * 20 + " 模板样式详细检查 " + "=" * 20)
            if self.template_vertical_tables:
                print(f"\n--- 检查 {len(self.template_vertical_tables)} 个纵向模板的格式 ---")
                self._inspect_and_print_template_styles(self.template_vertical_tables[0], f"第 1 个纵向模板")
            if self.template_horizontal_tables:
                print(f"\n--- 检查 {len(self.template_horizontal_tables)} 个横向模板的格式 ---")
                self._inspect_and_print_template_styles(self.template_horizontal_tables[0], f"第 1 个横向模板")
                self._inspect_and_print_template_content(self.template_horizontal_tables[0], f"第 1 个横向模板")
            print("\n" + "=" * 22 + " 模板检查结束 " + "=" * 22 + "\n")

            if not self.template_vertical_tables:
                print("    警告: 模板库中缺少纵向页面表格，将无法格式化任何纵向表格。")
            if not self.template_horizontal_tables:
                print("    警告: 模板库中缺少横向页面表格，将无法格式化任何横向表格。")

        except Exception as e:
            print(f"ERROR: 初始化样式刷失败: {e}")
            sys.exit(1)

    def _get_orientation_from_sectPr(self, sectPr) -> str:
        if sectPr is not None:
            pgSz = sectPr.find(qn('w:pgSz'))
            if pgSz is not None:
                orient = pgSz.get(qn('w:orient'))
                if orient == 'landscape':
                    return 'horizontal'
        return 'vertical'
    def _count_cols(self, tbl):
        maxc = 0
        for tr in tbl.findall(qn('w:tr')):
            maxc = max(maxc, len(tr.findall(qn('w:tc'))))
        return maxc

    def _get_styles_tree(self):
        styles_part = self.template_doc.styles.part
        return styles_part.element if styles_part is not None else None

    def _get_table_style_el(self, style_id):
        styles_el = self._get_styles_tree()
        if styles_el is None:
            return None
        for s in styles_el.findall(qn('w:style')):
            if s.get(qn('w:type')) == 'table' and s.get(qn('w:styleId')) == style_id:
                return s
        return None

    def _resolve_style_rpr_with_basedon(self, style_el):
        """沿 basedOn 递归，返回第一处出现的 rPr。"""
        seen = set()
        cur = style_el
        while cur is not None and id(cur) not in seen:
            seen.add(id(cur))
            rPr = cur.find(qn('w:rPr'))
            if rPr is not None:
                return rPr
            based = cur.find(qn('w:basedOn'))
            if based is None:
                break
            base_id = based.get(qn('w:val'))
            cur = self._get_table_style_el(base_id)
        return None

    def _get_docdefaults_rpr(self):
        styles_el = self._get_styles_tree()
        if styles_el is None:
            return None
        dd = styles_el.find(qn('w:docDefaults'))
        if dd is None:
            return None
        rpr_def = dd.find(qn('w:rPrDefault'))
        if rpr_def is None:
            return None
        return rpr_def.find(qn('w:rPr'))

    def _force_fixed_layout(self, table):
        try:
            table.autofit = False
        except Exception:
            pass
        tbl = table._tbl
        tblPr = tbl.tblPr or tbl.get_or_add_tblPr()
        layout = tblPr.find(qn('w:tblLayout'))
        if layout is None:
            layout = OxmlElement('w:tblLayout')
            tblPr.append(layout)
        layout.set(qn('w:type'), 'fixed')

    def _allow_autofit(self, table):
        try:
            table.autofit = True
        except Exception:
            pass
        tblPr = table._tbl.tblPr if table._tbl.tblPr is not None else table._tbl.get_or_add_tblPr()
        layout = tblPr.find(qn('w:tblLayout'))
        if layout is not None:
            tblPr.remove(layout)
        old_grid = table._tbl.find(qn('w:tblGrid'))
        if old_grid is not None:
            table._tbl.remove(old_grid)

    def _convert_width_to_absolute(self, width_elem, section_width_dxa=None):
        """
        将表格宽度从百分比或自动转换为绝对值(dxa)
        确保Excel转Word表格与普通Word表格宽度一致
        """
        if width_elem is None:
            return None

        width_type = width_elem.get(qn('w:type'))
        width_value = width_elem.get(qn('w:w'))

        if width_type == 'dxa':  # 已经是绝对值，直接返回
            return copy.deepcopy(width_elem)
        elif width_type == 'pct' and section_width_dxa:  # 百分比，需要转换
            # 计算绝对宽度：百分比 * 节宽度 / 5000 (Word的pct语义：5000=100%)
            absolute_width = int(width_value) * section_width_dxa // 5000
            new_width = copy.deepcopy(width_elem)
            new_width.set(qn('w:type'), 'dxa')
            new_width.set(qn('w:w'), str(absolute_width))
            return new_width
        elif width_type == 'auto':  # 自动宽度，设置为节宽度的80%
            if section_width_dxa:
                absolute_width = section_width_dxa * 8 // 10  # 80%
                new_width = copy.deepcopy(width_elem)
                new_width.set(qn('w:type'), 'dxa')
                new_width.set(qn('w:w'), str(absolute_width))
                return new_width
            else:
                return copy.deepcopy(width_elem)
        else:
            return copy.deepcopy(width_elem)

    def _get_template_text_width_twips(self):
        """
        取模板文档的版心宽度（twips）
        用于将模板的百分比宽度转换为准确的绝对值
        """
        try:
            sectPr = self.template_doc.element.body.find(qn('w:sectPr'))
            if sectPr is None:
                return None
            pgSz = sectPr.find(qn('w:pgSz'))
            pgMar = sectPr.find(qn('w:pgMar'))
            if pgSz is None or pgMar is None:
                return None

            w = int(pgSz.get(qn('w:w')))
            left = int(pgMar.get(qn('w:left')) or 0)
            right = int(pgMar.get(qn('w:right')) or 0)
            template_text_width = max(0, w - left - right)

            print(f"        📏 模板版心宽度: 页宽{w} - 左边距{left} - 右边距{right} = {template_text_width} twips")
            return template_text_width
        except Exception as e:
            print(f"        ⚠️ 获取模板版心宽度失败: {e}")
            return None

    def _get_section_width(self, table):
        """
        获取表格所在节的宽度(以dxa为单位)
        修复：使用前置分节而不是后置分节
        """
        try:
            # 查找表格所在的节设置
            table_element = table._tbl

            # 首先尝试查找最近的前置分节设置
            preceding_sect_prs = table_element.xpath('preceding::w:sectPr')
            if preceding_sect_prs:
                # 取最近的一个前置分节
                sect_pr = preceding_sect_prs[-1]
            else:
                # 如果没有前置分节，查找表格后的分节
                following_sect_prs = table_element.xpath('following::w:sectPr')
                if following_sect_prs:
                    sect_pr = following_sect_prs[0]
                else:
                    # 最后查找文档的默认节设置
                    sect_pr = table.document.element.body.find(qn('w:sectPr'))

            if sect_pr is not None:
                pgSz = sect_pr.find(qn('w:pgSz'))
                pgMar = sect_pr.find(qn('w:pgMar'))

                if pgSz is not None and pgMar is not None:
                    # 页面宽度 - 左边距 - 右边距
                    page_width = int(pgSz.get(qn('w:w'), '12240'))  # 默认A4宽度
                    left_margin = int(pgMar.get(qn('w:left'), '1440'))  # 默认左边距
                    right_margin = int(pgMar.get(qn('w:right'), '1440'))  # 默认右边距
                    section_width = page_width - left_margin - right_margin
                    print(f"        📏 节宽度计算: 页宽{page_width} - 左边距{left_margin} - 右边距{right_margin} = {section_width} dxa")
                    return section_width

            # 返回默认的节宽度(A4减去默认边距)
            print(f"        📏 使用默认节宽度: 9360 dxa (A4默认)")
            return 12240 - 1440 - 1440  # 9360 dxa
        except Exception as e:
            print(f"        ⚠️ 获取节宽度失败: {e}，使用默认值")
            return 9360  # 默认返回A4减去默认边距

    def _copy_table_box_metrics(self, src_table, tmpl_table):
        src_tblPr  = src_table._tbl.tblPr if src_table._tbl.tblPr is not None else src_table._tbl.get_or_add_tblPr()
        tmpl_tblPr = tmpl_table._tbl.tblPr if tmpl_table._tbl.tblPr is not None else tmpl_table._tbl.get_or_add_tblPr()

        # 1) 处理表格宽度（关键：即使模板没有tblW也要设置合理默认宽度）
        template_tblW = tmpl_tblPr.find(qn('w:tblW'))
        if template_tblW is not None:
            # 模板有宽度设置，用模板版心宽度转换并复制
            template_text_width = self._get_template_text_width_twips()
            if template_text_width is not None:
                absolute_width = self._convert_width_to_absolute(template_tblW, template_text_width)
            else:
                # 兜底：使用目标节宽度
                section_width = self._get_section_width(src_table)
                absolute_width = self._convert_width_to_absolute(template_tblW, section_width)
                print(f"        ⚠️ 模板版心宽度获取失败，使用目标节宽度")

            if absolute_width is not None:
                src_tblPr.append(absolute_width)
                width_type = absolute_width.get(qn('w:type'))
                width_value = absolute_width.get(qn('w:w'))
                print(f"        ✓ 表格宽度: {width_value} ({width_type}) - 从模板转换（基于模板版心）")
        else:
            # 模板没有宽度设置，使用模板版心宽度的90%作为默认值
            template_text_width = self._get_template_text_width_twips()
            if template_text_width is not None:
                default_width = template_text_width * 9 // 10  # 90%
            else:
                # 兜底：使用目标节宽度
                section_width = self._get_section_width(src_table)
                default_width = section_width * 9 // 10  # 90%
                print(f"        ⚠️ 模板版心宽度获取失败，使用目标节宽度")

            default_tblW = OxmlElement('w:tblW')
            default_tblW.set(qn('w:type'), 'dxa')
            default_tblW.set(qn('w:w'), str(default_width))
            src_tblPr.append(default_tblW)
            print(f"        ✓ 表格宽度: {default_width} (dxa) - 模板无宽度，使用90%模板版心")

        # 2) 处理单元格边距
        template_tblCellMar = tmpl_tblPr.find(qn('w:tblCellMar'))
        if template_tblCellMar is not None:
            src_tblPr.append(copy.deepcopy(template_tblCellMar))
            print(f"        ✓ 单元格边距设置 (从模板复制)")

    def _copy_tblGrid_exact(self, src_table, tmpl_table):
        src_tbl = src_table._tbl
        tmpl_tbl = tmpl_table._tbl
        old = src_tbl.find(qn('w:tblGrid'))
        if old is not None:
            src_tbl.remove(old)
        tmpl_grid = tmpl_tbl.find(qn('w:tblGrid'))
        if tmpl_grid is None:
            return

        # 复制并调整列网格宽度为绝对值
        new_grid = copy.deepcopy(tmpl_grid)
        section_width = self._get_section_width(src_table)

        # 处理每列的宽度
        grid_cols = new_grid.findall(qn('w:gridCol'))
        total_col_widths = 0
        for grid_col in grid_cols:
            col_width_elem = grid_col
            col_width_type = col_width_elem.get(qn('w:w'))
            if col_width_type:
                # 如果列宽也是百分比，转换为绝对值
                try:
                    col_width_value = int(col_width_type)
                    if col_width_value < 10000:  # 可能是百分比
                        absolute_col_width = col_width_value * section_width // 5000
                        col_width_elem.set(qn('w:w'), str(absolute_col_width))
                except:
                    pass
            total_col_widths += int(col_width_elem.get(qn('w:w'), '0'))

        # 插入新的列网格
        children = list(src_tbl)
        idx = 1 if children and children[0].tag == qn('w:tblPr') else 0
        src_tbl.insert(idx, new_grid)

        print(f"        ✓ 列网格: 复制了 {len(grid_cols)} 列，总宽度 {total_col_widths} dxa")
        for i, grid_col in enumerate(grid_cols[:5]):  # 只显示前5列避免输出过长
            col_width = grid_col.get(qn('w:w'))
            print(f"            - 列 {i+1} 宽度: {col_width} dxa")
        if len(grid_cols) > 5:
            print(f"            - ... 还有 {len(grid_cols)-5} 列")

    def _sync_table_dimensions_robust(self, target_table, template_table):
        """
        稳健同步尺寸：列数一致→精复制；列数不同→允许自适应避免竖排。
        修复Excel转Word表格宽度换算问题，确保表格大小一致。
        注意：此函数不再处理对齐，避免影响已设置的对齐方式。
        """
        print(f"        📐 开始尺寸同步...")
        self._copy_table_box_metrics(target_table, template_table)
        src_cols  = self._count_cols(target_table._tbl)
        tmpl_cols = self._count_cols(template_table._tbl)

        if src_cols > 0 and tmpl_cols > 0 and src_cols == tmpl_cols:
            self._copy_tblGrid_exact(target_table, template_table)
            self._force_fixed_layout(target_table)
            print(f"        ✓ 表格布局: fixed (强制固定布局)")
        else:
            self._allow_autofit(target_table)
            print(f"        ⚠️ 列数不匹配 (源:{src_cols} vs 模板:{tmpl_cols})，使用自适应布局")
        print(f"        📐 尺寸同步完成，对齐方式保持不变")
    def _categorize_tables_by_orientation(self, doc: DocumentClass) -> dict:
        categorized_tables = {'vertical': [], 'horizontal': []}
        doc_name = "模板文件" if doc == self.template_doc else "源文件"
        last_sect_pr = doc.element.body.find(qn('w:sectPr'))
        print(f"    -> 在'{doc_name}'中发现 {len(doc.tables)} 个表格，开始逐一分析...")
        for i, table in enumerate(doc.tables):
            table_element = table._tbl

            # 修复：使用前置分节而不是后置分节
            preceding_sect_prs = table_element.xpath('preceding::w:sectPr')
            if preceding_sect_prs:
                # 取最近的前置分节
                sectPr_for_table = preceding_sect_prs[-1]
                print(f"        - 表格 #{i + 1} 使用前置分节")
            else:
                # 如果没有前置分节，查找后置分节
                following_sect_prs = table_element.xpath('following::w:sectPr')
                sectPr_for_table = following_sect_prs[0] if following_sect_prs else last_sect_pr
                print(f"        - 表格 #{i + 1} 使用后置分节 (无前置分节)")

            orientation = self._get_orientation_from_sectPr(sectPr_for_table)
            print(f"        - 正在分析'{doc_name}'的表格 #{i + 1}... 识别页面方向为: {orientation}")
            if orientation == 'vertical':
                categorized_tables['vertical'].append(table)
            else:
                categorized_tables['horizontal'].append(table)
        print(
            f"    -> 分析完成。纵向: {len(categorized_tables['vertical'])} 个, 横向: {len(categorized_tables['horizontal'])} 个。")
        return categorized_tables

    def _inspect_and_print_template_styles(self, template_table: Table, template_name: str):
        """
        增强版模板样式检查，支持主题字体检测
        """
        print(f"\n🔍 正在检查模板 '{template_name}' 的样式 (基于第一个单元格):")
        if not (template_table.rows and template_table.columns):
            print("    - 模板表格为空，无法检查样式。")
            return
        try:
            cell = template_table.cell(0, 0)
            template_pPr, template_rPr, template_tcPr = None, None, cell._tc.tcPr
            for p_element in cell._tc.iterfind(qn('w:p')):
                if template_pPr is None: template_pPr = p_element.find(qn('w:pPr'))
                if template_rPr is None:
                    if template_pPr is not None: template_rPr = template_pPr.find(qn('w:rPr'))
                    if template_rPr is None:
                        for r_element in p_element.iterfind(qn('w:r')):
                            rPr = r_element.find(qn('w:rPr'))
                            if rPr is not None: template_rPr = rPr; break
                if template_pPr is not None and template_rPr is not None: break

            print("    🎨 字体样式:")
            if template_rPr is not None:
                font_names = template_rPr.find(qn('w:rFonts'))
                if font_names is not None:
                    # 检查直接指定的字体
                    ascii_font = font_names.get(qn('w:ascii'))
                    eastAsia_font = font_names.get(qn('w:eastAsia'))

                    # 检查主题字体
                    ascii_theme = font_names.get(qn('w:asciiTheme'))
                    eastAsia_theme = font_names.get(qn('w:eastAsiaTheme'))
                    hint = font_names.get(qn('w:hint'))

                    if ascii_font:
                        print(f"        - 西文字体: {ascii_font}")
                    elif ascii_theme:
                        resolved_ascii = self._resolve_theme_font(ascii_theme)
                        print(f"        - 西文字体: {resolved_ascii} (主题: {ascii_theme})")
                    else:
                        print(f"        - 西文字体: 未指定")

                    if eastAsia_font:
                        print(f"        - 中文字体: {eastAsia_font}")
                    elif eastAsia_theme:
                        resolved_eastAsia = self._resolve_theme_font(eastAsia_theme)
                        print(f"        - 中文字体: {resolved_eastAsia} (主题: {eastAsia_theme})")
                    else:
                        print(f"        - 中文字体: 未指定")

                    if hint:
                        print(f"        - 文字类型: {hint}")

                sz = template_rPr.find(qn('w:sz'))
                if sz is not None:
                    font_size = int(sz.get(qn('w:val'))) / 2
                    print(f"        - 字号: {font_size}")
                color = template_rPr.find(qn('w:color'))
                if color is not None: print(f"        - 颜色: {color.get(qn('w:val'))}")
                print(f"        - 加粗: {'是' if template_rPr.find(qn('w:b')) is not None else '否'}")
            else:
                print("        - 未找到明确的字体样式定义。")

            print("    📏 对齐方式:")
            # 检查表格级别的对齐
            table_alignment = getattr(template_table, 'alignment', None)
            print(f"        - 表格级别对齐: {table_alignment}")

            # 检查XML中的表格对齐设置
            template_tblPr = template_table._tbl.tblPr
            xml_table_align = None
            if template_tblPr is not None:
                jc = template_tblPr.find(qn('w:jc'))
                if jc is not None:
                    xml_table_align = jc.get(qn('w:val'))
                    print(f"        - XML表格对齐: {xml_table_align}")

            h_align = "未指定"
            if template_pPr is not None:
                jc = template_pPr.find(qn('w:jc'))
                if jc is not None: h_align = jc.get(qn('w:val'))
            print(f"        - 段落水平对齐: {h_align}")
            v_align = "未指定"
            if template_tcPr is not None:
                vAlign = template_tcPr.find(qn('w:vAlign'))
                if vAlign is not None: v_align = vAlign.get(qn('w:val'))
            print(f"        - 单元格垂直对齐: {v_align}")

            print("    🖼️ 边框样式:")
            if template_tcPr is not None and template_tcPr.find(qn('w:tcBorders')) is not None:
                print("        - 检测到明确的单元格边框定义。")
            else:
                print("        - 未检测到明确的单元格边框定义 (可能继承自表格样式)。")

        except Exception as e:
            print(f"    - 检查样式时出错: {e}")

    def _inspect_and_print_template_content(self, template_table: Table, template_name: str):
        print(f"\n📋 正在检查模板 '{template_name}' 的内容:")
        if not template_table.rows:
            print("    - 模板表格为空，无内容可显示。")
            return
        try:
            col_count = len(template_table.columns)
            border = "+-" + "-+-".join(["-" * 15 for _ in range(col_count)]) + "-+"
            print(border)
            for row in template_table.rows:
                row_content = [
                    cell.text.replace('\n', ' ').strip()[:12].ljust(15) + (
                        '...' if len(cell.text.replace('\n', ' ').strip()) > 12 else '   ')
                    for cell in row.cells
                ]
                print(f"| {' | '.join(row_content)} |")
            print(border)
        except Exception as e:
            print(f"    - 检查内容时出错: {e}")

    # 放到工具函数区
    def _get_or_add_tblPr(tbl):
        # 避免使用 “tbl.tblPr or tbl.get_or_add_tblPr()” 导致 FutureWarning
        return tbl.tblPr if tbl.tblPr is not None else tbl.get_or_add_tblPr()

    def _apply_table_borders_exact(self, target_table, template_table):
        """
        若你已有该函数，可以删掉这个内置版。
        作用：把模板表的 <w:tblBorders> 精确复制到目标表；若模板无边框则移除目标表的边框。
        """
        try:
            tgt_tblPr = target_table._tbl.tblPr if target_table._tbl.tblPr is not None else target_table._tbl.get_or_add_tblPr()
            tmpl_tblPr = template_table._tbl.tblPr

            old = tgt_tblPr.find(qn('w:tblBorders'))
            if old is not None:
                tgt_tblPr.remove(old)

            tmpl_b = tmpl_tblPr.find(qn('w:tblBorders')) if tmpl_tblPr is not None else None
            if tmpl_b is not None:
                tgt_tblPr.append(copy.deepcopy(tmpl_b))
        except Exception:
            pass
    def _clear_table_paragraph_formatting(self, table):
        """
        清除表格前后段落的格式，确保表格对齐不受段落影响
        """
        try:
            # 获取表格在文档中的位置
            tbl_element = table._tbl
            parent = tbl_element.getparent()

            if parent is not None:
                # 查找表格前后的段落
                table_index = list(parent).index(tbl_element)

                # 检查并清除表格前的段落格式
                if table_index > 0:
                    prev_element = parent[table_index - 1]
                    if prev_element.tag == qn('w:p'):
                        pPr = prev_element.find(qn('w:pPr'))
                        if pPr is not None:
                            jc = pPr.find(qn('w:jc'))
                            if jc is not None:
                                # 清除段落的居中对齐
                                pPr.remove(jc)
                                print(f"        ✓ 清除表格前段落的对齐设置")

                # 检查并清除表格后的段落格式
                if table_index < len(parent) - 1:
                    next_element = parent[table_index + 1]
                    if next_element.tag == qn('w:p'):
                        pPr = next_element.find(qn('w:pPr'))
                        if pPr is not None:
                            jc = pPr.find(qn('w:jc'))
                            if jc is not None:
                                # 清除段落的居中对齐
                                pPr.remove(jc)
                                print(f"        ✓ 清除表格后段落的对齐设置")
        except Exception as e:
            print(f"        ⚠️ 清除表格段落格式时出错: {e}")

    def _resolve_table_alignment_from_style(self, template_table):
        """
        返回模板表的对齐值 'left'/'center'/'right'（优先级：表本体 → 表格样式 → basedOn 递归）。
        找不到时返回 None。
        """
        # 1) 先看表本体的 tblPr
        tblPr = template_table._tbl.tblPr
        if tblPr is not None:
            jc = tblPr.find(qn('w:jc'))
            if jc is not None:
                return jc.get(qn('w:val'))

        # 2) 再看表格样式（含 basedOn 继承链）
        try:
            tblStyle = None
            if tblPr is not None:
                el = tblPr.find(qn('w:tblStyle'))
                if el is not None:
                    tblStyle = el.get(qn('w:val'))
            if not tblStyle:
                return None

            styles_part = self.template_doc.styles.part
            if styles_part is None:
                return None
            styles_el = styles_part.element

            # 收集所有表格样式为字典
            styles_by_id = {}
            for s in styles_el.findall(qn('w:style')):
                if s.get(qn('w:type')) == 'table':
                    styles_by_id[s.get(qn('w:styleId'))] = s

            def find_jc_in_style(style_el):
                if style_el is None:
                    return None
                tblPr_s = style_el.find(qn('w:tblPr'))
                if tblPr_s is not None:
                    jc_s = tblPr_s.find(qn('w:jc'))
                    if jc_s is not None:
                        return jc_s.get(qn('w:val'))
                # 递归 basedOn
                based = style_el.find(qn('w:basedOn'))
                if based is not None:
                    base_id = based.get(qn('w:val'))
                    return find_jc_in_style(styles_by_id.get(base_id))
                return None

            return find_jc_in_style(styles_by_id.get(tblStyle))
        except Exception:
            return None

    def _majority_cell_paragraph_alignment(self, tmpl_table):
        """
        统计模板表中所有单元格内段落的 w:jc（left/center/right），返回出现次数最多的一个；
        若都没写 w:jc，则返回 None
        """
        try:
            alignment_counts = {'left': 0, 'center': 0, 'right': 0}
            total_cells = 0

            # 遍历模板表的所有单元格
            for row in tmpl_table.rows:
                for cell in row.cells:
                    total_cells += 1
                    # 检查单元格内的所有段落
                    for p_element in cell._tc.iterfind(qn('w:p')):
                        pPr = p_element.find(qn('w:pPr'))
                        if pPr is not None:
                            jc = pPr.find(qn('w:jc'))
                            if jc is not None:
                                align_val = jc.get(qn('w:val'))
                                if align_val in alignment_counts:
                                    alignment_counts[align_val] += 1
                                break  # 每个单元格只统计第一个段落的对齐

            # 找到出现次数最多的对齐方式
            max_count = max(alignment_counts.values())
            if max_count == 0:
                print(f"        📊 单元格段落投票统计: {total_cells} 个单元格均未设置对齐，返回 None")
                return None

            # 找到票数最多的对齐方式
            for align, count in alignment_counts.items():
                if count == max_count:
                    print(f"        📊 单元格段落投票统计: {total_cells} 个单元格中 {align}={count} 票（最多）")
                    return align

            return None
        except Exception as e:
            print(f"        ⚠️ 单元格段落投票统计失败: {e}")
            return None

    def _apply_table_alignment_exact(self, target_table, template_table):
        """
        完全重写：当样式解析结果为None时，改用单元格段落投票；
        其他逻辑保持不变（清干扰、写回、同步缩进）。
        """
        try:
            # 清理样式展开属性，避免覆盖表级对齐
            tblPrEx = target_table._tbl.find(qn('w:tblPrEx'))
            if tblPrEx is not None:
                target_table._tbl.remove(tblPrEx)
                print("        ✓ 移除表格样式展开属性 (tblPrEx)")

            target_tblPr = target_table._tbl.tblPr if target_table._tbl.tblPr is not None else target_table._tbl.get_or_add_tblPr()
            template_tblPr = template_table._tbl.tblPr

            print(f"        🎯 开始精确对齐控制...")

            # ① 清干扰：tblpPr / 旧 tblInd / 旧 jc
            tblpPr = target_tblPr.find(qn('w:tblpPr'))
            if tblpPr is not None:
                target_tblPr.remove(tblpPr)
                print(f"        ✓ 移除表格浮动定位属性 (tblpPr)")

            old_ind = target_tblPr.find(qn('w:tblInd'))
            if old_ind is not None:
                target_tblPr.remove(old_ind)
                print(f"        ✓ 移除原有表格缩进 (tblInd)")

            old_jc = target_tblPr.find(qn('w:jc'))
            if old_jc is not None:
                target_tblPr.remove(old_jc)

            # ② 解析"模板对齐"（表本体→表格样式→basedOn）
            align_val = self._resolve_table_alignment_from_style(template_table)
            print(f"        🔍 样式解析结果: {align_val}")

            # ③ 如果样式解析返回None，改用单元格段落投票
            if align_val is None:
                align_val = self._majority_cell_paragraph_alignment(template_table)

            # ④ 写回目标表
            if align_val in ('left', 'center', 'right'):
                new_jc = OxmlElement('w:jc')
                new_jc.set(qn('w:val'), align_val)
                target_tblPr.append(new_jc)
                print(f"        ✓ 表格对齐: {align_val}（来自模板样式解析）")

                # 同时设置 python-docx 的对齐属性（双保险）
                try:
                    from docx.enum.table import WD_TABLE_ALIGNMENT
                    if align_val == 'center':
                        target_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    elif align_val == 'right':
                        target_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
                    elif align_val == 'left':
                        target_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    print(f"        ✓ 同步设置python-docx对齐属性: {align_val}")
                except Exception:
                    pass
            else:
                # 兜底：用 python-docx 的 alignment
                try:
                    from docx.enum.table import WD_TABLE_ALIGNMENT
                    al = getattr(template_table, 'alignment', None)
                    if al is not None:
                        target_table.alignment = al
                        print(f"        ✓ 表格对齐: {al} (来自python-docx alignment)")
                    else:
                        print(f"        ✓ 模板未显式声明对齐，保持默认")
                except Exception:
                    print(f"        ✓ 模板未显式声明对齐，保持默认")

            # ⑤ 同步模板的缩进（有就复制，没有就不加）
            if template_tblPr is not None:
                tmpl_ind = template_tblPr.find(qn('w:tblInd'))
                if tmpl_ind is not None:
                    target_tblPr.append(copy.deepcopy(tmpl_ind))
                    indent_value = tmpl_ind.get(qn('w:w'))
                    print(f"        ✓ 表格缩进: {indent_value} (复制模板)")

        except Exception as e:
            print(f"        ❌ 精确对齐控制失败: {e}")

    def _is_in_textbox(self, table):
        """
        检测表格是否在文本框/形状中
        只要祖先链中出现 w:txbxContent，就认为在文本框里
        """
        el = table._tbl
        for anc in el.iterancestors():
            if anc.tag == qn('w:txbxContent'):
                return True
        return False

    def _apply_template_fonts_table_fallback(self, source_table: Table, template_table: Table):
        """
        无条件字体刷法fallback：
        - 行列一致 → (r,c) 一一对应刷
        - 行列不一致 → 用模板(0,0)作为"基准样式"刷整表
        确保无论如何都能刷上字体！
        """
        try:
            # 取模板首格作"基准样式"
            base_spec = self._extract_effective_font_style_dict(template_table.cell(0, 0), template_table)

            # 判断是否同构
            same_shape = (len(source_table.rows) == len(template_table.rows) and
                          self._count_cols(source_table._tbl) == self._count_cols(template_table._tbl))

            total_cells = 0
            for r, row in enumerate(source_table.rows):
                for c, cell in enumerate(row.cells):
                    if same_shape:
                        # 行列一致：使用对应的模板单元格样式
                        spec = self._extract_effective_font_style_dict(template_table.cell(r, c), template_table)
                    else:
                        # 行列不一致：使用基准样式
                        spec = base_spec

                    # 应用字体样式
                    self._apply_font_style_dict_to_cell(cell, spec)
                    total_cells += 1

        except Exception as e:
            print(f"        ❌ 无条件字体刷法失败: {e}")

    def _apply_complete_template_cell_styles(self, source_table: Table, template_table: Table):
        """
        强制应用模板所有单元格样式：确保原文件表格完全按照模板格式
        包括字体、字号、加粗、对齐、边框、垂直对齐等所有属性
        """
        try:
            print(f"        🎯 强制应用模板所有单元格样式...")

            # 获取模板所有单元格的样式
            template_cells = []
            for row in template_table.rows:
                for cell in row.cells:
                    template_cells.append(cell)

            if not template_cells:
                print(f"        ⚠️ 模板表格没有单元格，跳过样式应用")
                return

            # 使用底层XML遍历源表格的所有单元格（包括合并单元格）
            source_table_element = source_table._tbl
            source_cells_flat = []

            for row_element in source_table_element.findall(qn('w:tr')):
                for cell_element in row_element.findall(qn('w:tc')):
                    from docx.table import _Cell
                    temp_cell = _Cell(cell_element, source_table)
                    source_cells_flat.append(temp_cell)

            # 一一对应应用样式（循环使用模板单元格）
            for i, source_cell in enumerate(source_cells_flat):
                template_cell = template_cells[i % len(template_cells)]

                # 1. 应用字体样式
                font_style_dict = self._extract_effective_font_style_dict(template_cell, template_table)
                self._apply_font_style_dict_to_cell(source_cell, font_style_dict)

                # 2. 应用段落样式（对齐等）
                self._apply_paragraph_alignment(source_cell, template_cell)

                # 3. 应用单元格属性（垂直对齐、边框等）
                self._apply_cell_properties(source_cell, template_cell)

            print(f"        ✅ 强制样式应用完成：处理了 {len(source_cells_flat)} 个单元格")

        except Exception as e:
            print(f"        ❌ 强制应用模板样式失败: {e}")
            import traceback
            traceback.print_exc()

    def _apply_paragraph_alignment(self, target_cell: _Cell, template_cell: _Cell):
        """应用模板单元格的段落对齐样式"""
        try:
            # 获取模板段落对齐设置
            template_alignment = None
            for p_element in template_cell._tc.iterfind(qn('w:p')):
                pPr = p_element.find(qn('w:pPr'))
                if pPr is not None:
                    jc = pPr.find(qn('w:jc'))
                    if jc is not None:
                        template_alignment = jc.get(qn('w:val'))
                        break

            # 应用到目标单元格的所有段落
            if template_alignment:
                for para in target_cell.paragraphs:
                    target_pPr = para._p.get_or_add_pPr()
                    # 清除原有对齐设置
                    old_jc = target_pPr.find(qn('w:jc'))
                    if old_jc is not None:
                        target_pPr.remove(old_jc)
                    # 设置新的对齐
                    new_jc = OxmlElement('w:jc')
                    new_jc.set(qn('w:val'), template_alignment)
                    target_pPr.append(new_jc)

        except Exception as e:
            print(f"        ⚠️ 应用段落对齐时出错: {e}")

    def _apply_cell_properties(self, target_cell: _Cell, template_cell: _Cell):
        """应用模板单元格的属性（垂直对齐、边框等）"""
        try:
            # 获取模板单元格属性
            template_tcPr = template_cell._tc.tcPr
            if template_tcPr is None:
                return

            target_tcPr = target_cell._tc.get_or_add_tcPr()

            # 应用垂直对齐
            template_vAlign = template_tcPr.find(qn('w:vAlign'))
            if template_vAlign is not None:
                old_vAlign = target_tcPr.find(qn('w:vAlign'))
                if old_vAlign is not None:
                    target_tcPr.remove(old_vAlign)
                target_tcPr.append(copy.deepcopy(template_vAlign))

            # 应用边框样式
            template_borders = template_tcPr.find(qn('w:tcBorders'))
            if template_borders is not None:
                old_borders = target_tcPr.find(qn('w:tcBorders'))
                if old_borders is not None:
                    target_tcPr.remove(old_borders)
                target_tcPr.append(copy.deepcopy(template_borders))

            # 应用其他单元格属性
            for child in list(template_tcPr):
                if child.tag not in [qn('w:vAlign'), qn('w:tcBorders')]:
                    # 跳过已处理的属性，复制其他属性
                    existing_child = target_tcPr.find(child.tag)
                    if existing_child is not None:
                        target_tcPr.remove(existing_child)
                    target_tcPr.append(copy.deepcopy(child))

        except Exception as e:
            print(f"        ⚠️ 应用单元格属性时出错: {e}")

    def _apply_table_style(self, table, template_table):
        """
        完整应用模板样式：确保原文件表格完全按照模板文件格式转换
        """
        print(f"    🔧 完整应用模板样式...")

        # 检测表格是否在文本框中
        if self._is_in_textbox(table):
            print("        ⚠️ 表格位于文本框/形状中：容器会限制对齐和宽度")

        # 1) 清除原有样式
        self._clear_table_paragraph_formatting(table)

        # 2) 应用表格样式名
        try:
            if template_table.style is not None:
                table.style = template_table.style
                print(f"        ✓ 表格样式: {template_table.style}")
        except Exception:
            pass

        # 3) 应用表级边框
        try:
            self._apply_table_borders_exact(table, template_table)
            print(f"        ✓ 表格级别边框")
        except Exception:
            pass

        # 4) 应用尺寸一致性
        self._sync_table_dimensions_robust(table, template_table)

        # 5) 【关键】应用模板字体到所有单元格
        self._apply_template_fonts_table_fallback(table, template_table)

        # 6) 应用表格和段落对齐
        self._apply_table_alignment_exact(table, template_table)

        # 7) 【新增】强制应用模板所有单元格样式
        self._apply_complete_template_cell_styles(table, template_table)


    # =================================================================================
    # |                           【核心修改区域 START】                                |
    # =================================================================================

    def _create_default_complete_borders(self):
        """
        从模板文档动态获取默认边框样式
        如果模板没有边框定义，使用Word标准默认值
        """
        tcBorders = OxmlElement('w:tcBorders')

        # 尝试从模板文档的表格样式获取默认边框
        try:
            # 查看第一个模板表格的边框样式
            if self.template_vertical_tables:
                template_table = self.template_vertical_tables[0]
                table_borders = self._extract_template_border_style(template_table)
                if table_borders is not None:
                    return copy.deepcopy(table_borders)
            elif self.template_horizontal_tables:
                template_table = self.template_horizontal_tables[0]
                table_borders = self._extract_template_border_style(template_table)
                if table_borders is not None:
                    return copy.deepcopy(table_borders)
        except Exception as e:
            print(f"        ⚠️ 从模板获取边框样式失败: {e}")

        # 如果无法从模板获取，使用Word标准默认边框
        border_details = {
            'w:top': {qn('w:val'): 'single', qn('w:sz'): '4', qn('w:color'): 'auto'},
            'w:left': {qn('w:val'): 'single', qn('w:sz'): '4', qn('w:color'): 'auto'},
            'w:bottom': {qn('w:val'): 'single', qn('w:sz'): '4', qn('w:color'): 'auto'},
            'w:right': {qn('w:val'): 'single', qn('w:sz'): '4', qn('w:color'): 'auto'},
        }
        for border_tag, attrs in border_details.items():
            border_el = OxmlElement(border_tag)
            for attr_key, attr_val in attrs.items():
                border_el.set(attr_key, attr_val)
            tcBorders.append(border_el)
        return tcBorders

    def _extract_template_border_style(self, template_table: Table):
        """
        从模板表格提取边框样式
        """
        try:
            # 1. 检查表格级别的边框
            tblPr = template_table._tbl.tblPr
            if tblPr is not None:
                tbl_borders = tblPr.find(qn('w:tblBorders'))
                if tbl_borders is not None:
                    return tbl_borders

            # 2. 检查第一个单元格的边框
            if template_table.rows and template_table.columns:
                cell = template_table.cell(0, 0)
                tcPr = cell._tc.tcPr
                if tcPr is not None:
                    tc_borders = tcPr.find(qn('w:tcBorders'))
                    if tc_borders is not None:
                        return tc_borders

            # 3. 检查表格样式中的边框
            table_style = self._find_table_style(template_table)
            if table_style is not None:
                tbl_borders = table_style.find(qn('w:tblBorders'))
                if tbl_borders is not None:
                    return tbl_borders

            return None
        except Exception as e:
            print(f"        ⚠️ 提取模板边框样式失败: {e}")
            return None

    def _debug_template_styles(self, template_cell: _Cell, cell_info: str = "模板单元格"):
        """
        调试方法：详细检查模板单元格中的样式信息，特别关注字体大小
        """
        print(f"\n🔍 调试 - {cell_info}样式检查:")
        try:
            template_tcPr = template_cell._tc.tcPr
            print(f"    单元格属性存在: {template_tcPr is not None}")

            # 检查段落和文本样式
            found_pPr = False
            found_rPr = False
            font_size_found = False

            for p_element in template_cell._tc.iterfind(qn('w:p')):
                pPr = p_element.find(qn('w:pPr'))
                if pPr is not None and not found_pPr:
                    print(f"    ✅ 找到段落样式 (pPr)")
                    found_pPr = True

                    # 检查对齐方式
                    jc = pPr.find(qn('w:jc'))
                    if jc is not None:
                        print(f"        - 水平对齐: {jc.get(qn('w:val'))}")

                for r_element in p_element.iterfind(qn('w:r')):
                    rPr = r_element.find(qn('w:rPr'))
                    if rPr is not None and not found_rPr:
                        print(f"    ✅ 找到文本样式 (rPr)")
                        found_rPr = True

                        # 检查字体属性
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            ascii_font = rFonts.get(qn('w:ascii'))
                            eastAsia_font = rFonts.get(qn('w:eastAsia'))
                            print(f"        - 西文字体: {ascii_font}")
                            print(f"        - 中文字体: {eastAsia_font}")

                        # 详细检查字号
                        sz = rPr.find(qn('w:sz'))
                        szCs = rPr.find(qn('w:szCs'))

                        if sz is not None:
                            font_size = int(sz.get(qn('w:val'))) / 2
                            print(f"        - 字号 (sz): {font_size}pt")
                            font_size_found = True

                        if szCs is not None:
                            font_size_cs = int(szCs.get(qn('w:val'))) / 2
                            print(f"        - 中文字号 (szCs): {font_size_cs}pt")
                            font_size_found = True

                        if not font_size_found:
                            print(f"        - ⚠️ 未找到字号定义")

                        # 检查颜色
                        color = rPr.find(qn('w:color'))
                        if color is not None:
                            text_color = color.get(qn('w:val'))
                            print(f"        - 颜色: {text_color}")

                        # 检查加粗
                        b = rPr.find(qn('w:b'))
                        if b is not None:
                            print(f"        - 加粗: 是")

                        # 检查斜体
                        i = rPr.find(qn('w:i'))
                        if i is not None:
                            print(f"        - 斜体: 是")

                        break

                if found_pPr and found_rPr:
                    break

            # 检查边框
            if template_tcPr is not None:
                borders = template_tcPr.find(qn('w:tcBorders'))
                if borders is not None:
                    print(f"    ✅ 找到边框样式")
                else:
                    print(f"    ⚠️ 未找到边框样式")

                # 检查垂直对齐
                vAlign = template_tcPr.find(qn('w:vAlign'))
                if vAlign is not None:
                    print(f"        - 垂直对齐: {vAlign.get(qn('w:val'))}")

            if not found_pPr and not found_rPr:
                print(f"    ❌ 未找到任何样式定义")
                print(f"    📝 单元格文本内容: '{template_cell.text}'")

        except Exception as e:
            print(f"    ❌ 调试检查出错: {e}")

    def _find_table_style(self, template_table: Table):
        """
        查找表格应用的样式定义，包括表格级别和样式级别
        """
        try:
            # 1. 查找表格级别的样式
            table_element = template_table._tbl
            tblPr = table_element.tblPr

            # 查找表格样式引用
            table_style = None
            if tblPr is not None:
                tblStyle = tblPr.find(qn('w:tblStyle'))
                if tblStyle is not None:
                    style_id = tblStyle.get(qn('w:val'))
                    # 在文档的样式部分查找对应的样式定义
                    try:
                        styles_part = self.template_doc.styles.part
                        if styles_part is not None:
                            styles_element = styles_part.element
                            for style in styles_element.findall(qn('w:style')):
                                if style.get(qn('w:styleId')) == style_id:
                                    table_style = style
                                    print(f"    ✅ 找到表格样式定义")
                                    break
                    except Exception:
                        pass

            return table_style

        except Exception as e:
            print(f"    ⚠️ 查找表格样式时出错: {e}")
            return None

    def _debug_table_style_font_size(self, table_style):
        """
        调试表格样式中的字体大小定义
        """
        try:
            print(f"        🔍 检查表格样式中的字体大小:")

            # 检查表格样式中的文本样式
            rPr = table_style.find(qn('w:rPr'))
            if rPr is not None:
                sz = rPr.find(qn('w:sz'))
                szCs = rPr.find(qn('w:szCs'))

                if sz is not None:
                    font_size = int(sz.get(qn('w:val'))) / 2
                    print(f"            - 表格样式字号 (sz): {font_size}pt")

                if szCs is not None:
                    font_size_cs = int(szCs.get(qn('w:val'))) / 2
                    print(f"            - 表格样式中文字号 (szCs): {font_size_cs}pt")

                if sz is None and szCs is None:
                    print(f"            - ⚠️ 表格样式中未定义字号")

            else:
                print(f"            - ⚠️ 表格样式中未找到文本样式定义")

            # 检查表格样式中的段落样式
            pPr = table_style.find(qn('w:pPr'))
            if pPr is not None:
                print(f"            - 找到表格样式段落定义")

        except Exception as e:
            print(f"            - ❌ 调试表格样式字体大小时出错: {e}")

    def _debug_font_size_inheritance(self, template_cell: _Cell, table_style):
        """
        调试字体大小的继承链：单元格 -> 表格样式 -> 默认样式
        """
        try:
            print(f"        📜 字体大小继承链分析:")

            # 1. 检查单元格级别的字体大小
            cell_font_size = None
            for p_element in template_cell._tc.iterfind(qn('w:p')):
                for r_element in p_element.iterfind(qn('w:r')):
                    rPr = r_element.find(qn('w:rPr'))
                    if rPr is not None:
                        sz = rPr.find(qn('w:sz'))
                        if sz is not None:
                            cell_font_size = int(sz.get(qn('w:val'))) / 2
                            print(f"            1️⃣ 单元格字号: {cell_font_size}pt")
                            break
                    if cell_font_size:
                        break
                if cell_font_size:
                    break

            # 2. 检查表格样式中的字体大小
            style_font_size = None
            if table_style is not None:
                rPr = table_style.find(qn('w:rPr'))
                if rPr is not None:
                    sz = rPr.find(qn('w:sz'))
                    if sz is not None:
                        style_font_size = int(sz.get(qn('w:val'))) / 2
                        print(f"            2️⃣ 表格样式字号: {style_font_size}pt")

            # 3. 检查默认样式
            try:
                normal_style = None
                styles_part = self.template_doc.styles.part
                if styles_part is not None:
                    styles_element = styles_part.element
                    for style in styles_element.findall(qn('w:style')):
                        if style.get(qn('w:styleId')) == 'Normal':
                            normal_style = style
                            break

                normal_font_size = None
                if normal_style is not None:
                    rPr = normal_style.find(qn('w:rPr'))
                    if rPr is not None:
                        sz = rPr.find(qn('w:sz'))
                        if sz is not None:
                            normal_font_size = int(sz.get(qn('w:val'))) / 2
                            print(f"            3️⃣ Normal样式字号: {normal_font_size}pt")
            except:
                print(f"            3️⃣ 无法获取Normal样式字号")

            # 4. 确定最终使用的字体大小
            final_font_size = cell_font_size or style_font_size or normal_font_size or 10.5
            print(f"            🎯 最终使用字号: {final_font_size}pt")

            return final_font_size

        except Exception as e:
            print(f"            - ❌ 字体大小继承分析出错: {e}")
            return 10.5

    def _pick_first_run_with_text(self, cell):
        """从单元格中选取第一个包含文本的run"""
        for p in cell.paragraphs:
            for r in p.runs:
                if r.text and r.text.strip():
                    return r
        # 全空时返回第一个run（给到结构）
        for p in cell.paragraphs:
            if p.runs:
                return p.runs[0]
        return None

    def _extract_effective_font_style_dict(self, template_cell: _Cell, template_table: Table):
        """
        增强版字体样式提取，支持主题字体解析
        """
        font_style_dict = {
            'ascii': None,      # 西文字体
            'eastAsia': None,   # 中文字体
            'size_pt': None,    # 字号（Pt对象）
            'bold': None,       # 加粗
            'italic': None,     # 斜体
            'color_hex': None,  # 颜色
        }

        try:
            # 1. 检查单元格/run级别的字体设置（支持主题字体）
            for p_element in template_cell._tc.iterfind(qn('w:p')):
                for r_element in p_element.iterfind(qn('w:r')):
                    rPr = r_element.find(qn('w:rPr'))
                    if rPr is not None:
                        # 字体 - 支持主题字体
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            # 优先使用直接指定的字体
                            font_style_dict['ascii'] = rFonts.get(qn('w:ascii'))
                            font_style_dict['eastAsia'] = rFonts.get(qn('w:eastAsia'))

                            # 如果没有直接指定字体，解析主题字体
                            if font_style_dict['ascii'] is None:
                                ascii_theme = rFonts.get(qn('w:asciiTheme'))
                                if ascii_theme:
                                    font_style_dict['ascii'] = self._resolve_theme_font(ascii_theme)

                            if font_style_dict['eastAsia'] is None:
                                eastAsia_theme = rFonts.get(qn('w:eastAsiaTheme'))
                                if eastAsia_theme:
                                    font_style_dict['eastAsia'] = self._resolve_theme_font(eastAsia_theme)

                            # 根据hint设置默认字体（从Normal样式获取）
                            if font_style_dict['ascii'] is None and font_style_dict['eastAsia'] is None:
                                hint = rFonts.get(qn('w:hint'))
                                normal_fonts = self._get_normal_style_fonts_for_extraction()
                                if hint == 'eastAsia':
                                    font_style_dict['eastAsia'] = normal_fonts.get('eastAsia', '宋体')
                                else:
                                    font_style_dict['ascii'] = normal_fonts.get('ascii', 'Times New Roman')

                        # 字号
                        sz = rPr.find(qn('w:sz'))
                        szCs = rPr.find(qn('w:szCs'))
                        if sz is not None or szCs is not None:
                            from docx.shared import Pt
                            font_size_val = int((sz or szCs).get(qn('w:val'))) / 2
                            font_style_dict['size_pt'] = Pt(font_size_val)

                        # 粗体
                        if rPr.find(qn('w:b')) is not None:
                            font_style_dict['bold'] = True

                        # 斜体
                        if rPr.find(qn('w:i')) is not None:
                            font_style_dict['italic'] = True

                        # 颜色
                        color = rPr.find(qn('w:color'))
                        if color is not None and color.get(qn('w:val')):
                            font_style_dict['color_hex'] = color.get(qn('w:val'))

                        # 如果找到字体设置，停止搜索
                        if font_style_dict['ascii'] or font_style_dict['eastAsia']:
                            break
                if font_style_dict['ascii'] or font_style_dict['eastAsia']:
                    break

            # 2. 如果没有找到字体，从表格样式获取
            if (font_style_dict['ascii'] is None and font_style_dict['eastAsia'] is None):
                table_style = self._find_table_style(template_table)
                if table_style is not None:
                    rPr = table_style.find(qn('w:rPr'))
                    if rPr is not None:
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            font_style_dict['ascii'] = rFonts.get(qn('w:ascii'))
                            font_style_dict['eastAsia'] = rFonts.get(qn('w:eastAsia'))

                            # 解析主题字体
                            if font_style_dict['ascii'] is None:
                                ascii_theme = rFonts.get(qn('w:asciiTheme'))
                                if ascii_theme:
                                    font_style_dict['ascii'] = self._resolve_theme_font(ascii_theme)

                            if font_style_dict['eastAsia'] is None:
                                eastAsia_theme = rFonts.get(qn('w:eastAsiaTheme'))
                                if eastAsia_theme:
                                    font_style_dict['eastAsia'] = self._resolve_theme_font(eastAsia_theme)

                        sz = rPr.find(qn('w:sz'))
                        if sz is not None and font_style_dict['size_pt'] is None:
                            from docx.shared import Pt
                            font_style_dict['size_pt'] = Pt(int(sz.get(qn('w:val'))) / 2)

                        if rPr.find(qn('w:b')) is not None:
                            font_style_dict['bold'] = True

            # 3. 如果仍然没有找到，使用Normal样式
            if (font_style_dict['ascii'] is None and font_style_dict['eastAsia'] is None):
                normal_fonts = self._get_normal_style_fonts_for_extraction()
                if normal_fonts.get('ascii'):
                    font_style_dict['ascii'] = normal_fonts['ascii']
                if normal_fonts.get('eastAsia'):
                    font_style_dict['eastAsia'] = normal_fonts['eastAsia']
                if normal_fonts.get('size_pt'):
                    font_style_dict['size_pt'] = normal_fonts['size_pt']

            # 4. 如果仍然没有找到，使用Normal样式的值
            if (font_style_dict['ascii'] is None and font_style_dict['eastAsia'] is None):
                normal_fonts = self._get_normal_style_fonts_for_extraction()
                if normal_fonts.get('ascii'):
                    font_style_dict['ascii'] = normal_fonts['ascii']
                if normal_fonts.get('eastAsia'):
                    font_style_dict['eastAsia'] = normal_fonts['eastAsia']
                if normal_fonts.get('size_pt'):
                    font_style_dict['size_pt'] = normal_fonts['size_pt']

            # 5. 最后兜底：从模板Normal样式获取默认值
            if font_style_dict['ascii'] is None or font_style_dict['eastAsia'] is None:
                normal_fonts = self._get_normal_style_fonts_for_extraction()
                if font_style_dict['ascii'] is None:
                    font_style_dict['ascii'] = normal_fonts.get('ascii', 'Times New Roman')
                if font_style_dict['eastAsia'] is None:
                    font_style_dict['eastAsia'] = normal_fonts.get('eastAsia', '宋体')
            if font_style_dict['size_pt'] is None:
                normal_fonts = self._get_normal_style_fonts_for_extraction()
                if normal_fonts.get('size_pt'):
                    font_style_dict['size_pt'] = normal_fonts['size_pt']
                else:
                    from docx.shared import Pt
                    font_style_dict['size_pt'] = Pt(10.5)  # Word默认值作为最后兜底

        except Exception as e:
            print(f"        ⚠️ 提取字体样式时出错: {e}")
            # 从模板Normal样式获取兜底值（完全动态）
            normal_fonts = self._get_normal_style_fonts_for_extraction()
            from docx.shared import Pt

            # 动态获取默认字体大小
            default_size = normal_fonts.get('size_pt')
            if default_size is None:
                # 尝试从docDefaults获取
                docdefaults_rpr = self._get_docdefaults_rpr()
                if docdefaults_rpr is not None:
                    sz = docdefaults_rpr.find(qn('w:sz'))
                    if sz is not None:
                        default_size = Pt(int(sz.get(qn('w:val'))) / 2)

            # 如果仍然没有获取到，使用Word标准默认值
            if default_size is None:
                default_size = Pt(10.5)

            font_style_dict = {
                'ascii': normal_fonts.get('ascii', 'Times New Roman'),
                'eastAsia': normal_fonts.get('eastAsia', '宋体'),
                'size_pt': default_size,
                'bold': None,  # 从模板中获取
                'italic': None,  # 从模板中获取
                'color_hex': None,  # 从模板中获取
            }

        return font_style_dict

    def _resolve_theme_font(self, theme_name):
        """
        解析主题字体名称，返回对应的实际字体名称
        从模板文档中动态获取主题字体映射
        """
        try:
            # 从模板文档的主题部分获取实际字体映射
            theme_part = self.template_doc.part.related_parts.get('/theme/theme1.xml')
            if theme_part is not None:
                theme_root = theme_part.element

                # 查找字体方案
                font_scheme = theme_root.find(qn('a:themeElements'))
                if font_scheme is not None:
                    # 检查minor和major字体方案
                    for scheme_type in ['minorFontScheme', 'majorFontScheme']:
                        font_scheme_element = font_scheme.find(qn(f'a:{scheme_type}'))
                        if font_scheme_element is not None:
                            # 查找各种字体类型
                            latin_font = font_scheme_element.find(qn('a:latin'))
                            ea_font = font_scheme_element.find(qn('a:ea'))
                            cs_font = font_scheme_element.find(qn('a:cs'))

                            # 根据主题名称返回对应字体
                            if theme_name == 'minorEastAsia' or theme_name == 'majorEastAsia':
                                if ea_font is not None:
                                    typeface = ea_font.get(qn('a:typeface'))
                                    if typeface:
                                        return typeface
                            elif theme_name == 'minorAscii' or theme_name == 'majorAscii':
                                if latin_font is not None:
                                    typeface = latin_font.get(qn('a:typeface'))
                                    if typeface:
                                        return typeface
                            elif theme_name.startswith('minor') or theme_name.startswith('major'):
                                # 其他字体类型的回退
                                if ea_font is not None:
                                    typeface = ea_font.get(qn('a:typeface'))
                                    if typeface:
                                        return typeface
                                if latin_font is not None:
                                    typeface = latin_font.get(qn('a:typeface'))
                                    if typeface:
                                        return typeface

            # 如果无法从主题获取，尝试从Normal样式获取
            normal_fonts = self._get_normal_style_fonts_for_extraction()
            if normal_fonts.get('eastAsia'):
                return normal_fonts['eastAsia']
            if normal_fonts.get('ascii'):
                return normal_fonts['ascii']

            # 最后的备用映射：基于常见的主题字体（仅在完全失败时使用）
            common_theme_fonts = {
                'minorEastAsia': '宋体',
                'minorEastAsian': '宋体',
                'minorAscii': 'Calibri',
                'minorHAnsi': 'Calibri',
                'majorEastAsia': '黑体',
                'majorEastAsian': '黑体',
                'majorAscii': 'Times New Roman',
                'majorHAnsi': 'Times New Roman',
            }

            return common_theme_fonts.get(theme_name, '宋体')

        except Exception as e:
            print(f"        ⚠️ 解析主题字体时出错: {e}")
            # 如果无法解析，从Normal样式获取默认字体
            normal_fonts = self._get_normal_style_fonts_for_extraction()
            return normal_fonts.get('eastAsia', '宋体')

    def _has_font_properties(self, font_dict):
        """检查字体字典是否已有基本的字体属性"""
        return font_dict.get('ascii') is not None or font_dict.get('eastAsia') is not None or font_dict.get('size_pt') is not None

    def _format_style_dict(self, font_dict):
        """格式化样式字典用于显示"""
        formatted = {}
        for key, value in font_dict.items():
            if value is not None:
                if hasattr(value, '__repr__'):  # Pt对象等
                    formatted[key] = repr(value)
                else:
                    formatted[key] = value
            else:
                formatted[key] = 'None'
        return formatted

    def _get_normal_style_fonts_for_extraction(self):
        """从Normal样式获取字体设置（用于字体提取）"""
        try:
            # 使用已加载的模板文档，避免重新创建文档对象
            doc = self.template_doc
            for style in doc.styles:
                if style.style_id == 'Normal' or style.name == 'Normal':
                    result = {}
                    if hasattr(style, 'font') and style.font:
                        if style.font.name:
                            result["ascii"] = style.font.name
                            result["eastAsia"] = style.font.name
                        if hasattr(style.font, 'size') and style.font.size:
                            # 检查字号是否合理，如果过大则使用默认值
                            font_size_pt = style.font.size.pt if hasattr(style.font.size, 'pt') else float(style.font.size)
                            if font_size_pt > 50:  # 如果字号大于50pt，可能是错误的，使用默认值
                                from docx.shared import Pt
                                result["size_pt"] = Pt(11)  # 使用11pt作为默认值
                            else:
                                result["size_pt"] = style.font.size

                    print(f"        从Normal样式获取字体: {result}")
                    return result
            return {}
        except Exception as e:
            print(f"        获取Normal样式字体失败: {e}")
            return {}

    def _apply_font_style_dict_to_cell(self, target_cell: _Cell, font_style_dict: dict):
        """
        将样式字典应用到目标单元格的每个文本片段(run)
        简化版本：只应用基本字体属性，减少调试输出
        """
        try:
            # 遍历目标单元格的每个段落
            for para in target_cell.paragraphs:
                # 遍历段落中的每个文本片段(run)
                for run in para.runs:
                    if font_style_dict.get('ascii'):
                        run.font.name = font_style_dict['ascii']

                    if font_style_dict.get('eastAsia') or font_style_dict.get('ascii'):
                        # 设置完整字体：ascii + eastAsia + hAnsi + cs
                        rPr = run._element.get_or_add_rPr()
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is None:
                            rFonts = OxmlElement('w:rFonts')
                            rPr.append(rFonts)

                        font_ascii = font_style_dict.get('ascii') or 'Times New Roman'
                        font_eastAsia = font_style_dict.get('eastAsia') or font_ascii

                        rFonts.set(qn('w:ascii'), font_ascii)
                        rFonts.set(qn('w:eastAsia'), font_eastAsia)
                        rFonts.set(qn('w:hAnsi'), font_ascii)
                        rFonts.set(qn('w:cs'), font_eastAsia)

                    if font_style_dict.get('size_pt'):
                        run.font.size = font_style_dict['size_pt']

                    if font_style_dict.get('bold') is not None:
                        run.font.bold = font_style_dict['bold']

                    if font_style_dict.get('italic') is not None:
                        run.font.italic = font_style_dict['italic']

                    if font_style_dict.get('color_hex'):
                        run.font.color.rgb = None
                        run.font.color.hex = font_style_dict['color_hex']

        except Exception as e:
            print(f"            ❌ 应用样式字典到目标单元格时出错: {e}")

    def _get_template_cell_fonts(self, template_cell: _Cell):
        """
        直接从模板单元格提取字体设置
        优先级: 单元格 > 段落 > 文本运行
        """
        try:
            font_settings = {}

            # 1. 检查单元格级别的字体设置
            tcPr = template_cell._tc.tcPr
            if tcPr is not None:
                cell_rFonts = tcPr.find(qn('w:rFonts'))
                if cell_rFonts is not None:
                    ascii_font = cell_rFonts.get(qn('w:ascii'))
                    eastAsia_font = cell_rFonts.get(qn('w:eastAsia'))
                    hAnsi_font = cell_rFonts.get(qn('w:hAnsi'))
                    cs_font = cell_rFonts.get(qn('w:cs'))

                    if ascii_font or eastAsia_font or hAnsi_font or cs_font:
                        if ascii_font:
                            font_settings['w:ascii'] = ascii_font
                        if eastAsia_font:
                            font_settings['w:eastAsia'] = eastAsia_font
                        if hAnsi_font:
                            font_settings['w:hAnsi'] = hAnsi_font
                        if cs_font:
                            font_settings['w:cs'] = cs_font
                        print(f"        ✅ 从单元格级别获取字体: {font_settings}")
                        return font_settings

            # 2. 检查段落级别的字体设置
            for p_element in template_cell._tc.iterfind(qn('w:p')):
                pPr = p_element.find(qn('w:pPr'))
                if pPr is not None:
                    p_rFonts = pPr.find(qn('w:rFonts'))
                    if p_rFonts is not None:
                        ascii_font = p_rFonts.get(qn('w:ascii'))
                        eastAsia_font = p_rFonts.get(qn('w:eastAsia'))
                        hAnsi_font = p_rFonts.get(qn('w:hAnsi'))
                        cs_font = p_rFonts.get(qn('w:cs'))

                        if ascii_font or eastAsia_font or hAnsi_font or cs_font:
                            if ascii_font:
                                font_settings['w:ascii'] = ascii_font
                            if eastAsia_font:
                                font_settings['w:eastAsia'] = eastAsia_font
                            if hAnsi_font:
                                font_settings['w:hAnsi'] = hAnsi_font
                            if cs_font:
                                font_settings['w:cs'] = cs_font
                            print(f"        ✅ 从段落级别获取字体: {font_settings}")
                            return font_settings

            # 3. 检查文本运行级别的字体设置
            for p_element in template_cell._tc.iterfind(qn('w:p')):
                for r_element in p_element.iterfind(qn('w:r')):
                    rPr = r_element.find(qn('w:rPr'))
                    if rPr is not None:
                        r_rFonts = rPr.find(qn('w:rFonts'))
                        if r_rFonts is not None:
                            ascii_font = r_rFonts.get(qn('w:ascii'))
                            eastAsia_font = r_rFonts.get(qn('w:eastAsia'))
                            hAnsi_font = r_rFonts.get(qn('w:hAnsi'))
                            cs_font = r_rFonts.get(qn('w:cs'))

                            if ascii_font or eastAsia_font or hAnsi_font or cs_font:
                                if ascii_font:
                                    font_settings['w:ascii'] = ascii_font
                                if eastAsia_font:
                                    font_settings['w:eastAsia'] = eastAsia_font
                                if hAnsi_font:
                                    font_settings['w:hAnsi'] = hAnsi_font
                                if cs_font:
                                    font_settings['w:cs'] = cs_font
                                print(f"        ✅ 从文本运行获取字体: {font_settings}")
                                return font_settings

            print(f"        ⚠️ 模板单元格未找到明确的字体设置")
            return None

        except Exception as e:
            print(f"        ❌ 提取模板单元格字体时出错: {e}")
            return None

    def _get_default_fonts(self):
        """
        获取合理的默认字体设置
        首先尝试从Normal样式获取字体，如果失败则使用专业的文档字体
        """
        try:
            # 尝试从Normal样式获取字体设置
            doc = self.template_doc

            for style in doc.styles:
                if style.style_id == 'Normal' or style.name == 'Normal':
                    if hasattr(style, 'font') and style.font and style.font.name:
                        font_name = style.font.name
                        normal_fonts = {
                            'w:ascii': font_name,
                            'w:hAnsi': font_name,
                            'w:eastAsia': font_name,
                            'w:cs': font_name
                        }
                        print(f"        ✅ 使用Normal样式字体: {font_name}")
                        return normal_fonts
                    break

        except Exception as e:
            print(f"        ⚠️ 获取Normal样式字体失败: {e}")

        # 备用方案：从模板文档的默认样式获取字体
        try:
            # 尝试从docDefaults获取字体设置
            docdefaults_rpr = self._get_docdefaults_rpr()
            if docdefaults_rpr is not None:
                rFonts = docdefaults_rpr.find(qn('w:rFonts'))
                if rFonts is not None:
                    ascii_font = rFonts.get(qn('w:ascii')) or 'Times New Roman'
                    eastAsia_font = rFonts.get(qn('w:eastAsia')) or ascii_font
                    hAnsi_font = rFonts.get(qn('w:hAnsi')) or ascii_font
                    cs_font = rFonts.get(qn('w:cs')) or eastAsia_font

                    default_fonts = {
                        'w:ascii': ascii_font,
                        'w:hAnsi': hAnsi_font,
                        'w:eastAsia': eastAsia_font,
                        'w:cs': cs_font
                    }
                    print(f"        ✅ 使用文档默认字体: {default_fonts}")
                    return default_fonts
        except Exception as e:
            print(f"        ⚠️ 获取docDefaults字体失败: {e}")

        # 最后的备用方案：使用常见的文档字体组合（仅在无法从模板获取时使用）
        default_fonts = {
            'w:ascii': 'Times New Roman',  # 西文常用字体
            'w:hAnsi': 'Times New Roman',  # 高ANSI字体
            'w:eastAsia': '宋体',          # 中文字体
            'w:cs': '宋体'                 # 复杂脚本字体
        }

        print(f"        ✅ 使用备用字体: {default_fonts}")
        return default_fonts

    def _get_correct_font_size(self, template_cell: _Cell, template_table: Table):
        """
        获取正确的字体大小（简化版本，不包含详细调试）
        """
        try:
            # 1. 检查单元格级别的字体大小
            for p_element in template_cell._tc.iterfind(qn('w:p')):
                for r_element in p_element.iterfind(qn('w:r')):
                    rPr = r_element.find(qn('w:rPr'))
                    if rPr is not None:
                        sz = rPr.find(qn('w:sz'))
                        if sz is not None:
                            return int(sz.get(qn('w:val'))) / 2

            # 2. 检查表格样式中的字体大小
            table_style = self._find_table_style(template_table)
            if table_style is not None:
                rPr = table_style.find(qn('w:rPr'))
                if rPr is not None:
                    sz = rPr.find(qn('w:sz'))
                    if sz is not None:
                        return int(sz.get(qn('w:val'))) / 2

            # 3. 检查Normal样式
            try:
                styles_part = self.template_doc.styles.part
                if styles_part is not None:
                    styles_element = styles_part.element
                    for style in styles_element.findall(qn('w:style')):
                        if style.get(qn('w:styleId')) == 'Normal':
                            rPr = style.find(qn('w:rPr'))
                            if rPr is not None:
                                sz = rPr.find(qn('w:sz'))
                                if sz is not None:
                                    return int(sz.get(qn('w:val'))) / 2
                            break
            except:
                pass

            # 4. 从模板Normal样式获取默认字体大小
            try:
                normal_fonts = self._get_normal_style_fonts_for_extraction()
                if normal_fonts.get('size_pt'):
                    if hasattr(normal_fonts['size_pt'], 'pt'):
                        return normal_fonts['size_pt'].pt
                    else:
                        return float(normal_fonts['size_pt'])
            except:
                pass

            # 5. 最后的兜底值
            return 10.5

        except Exception as e:
            # 异常情况下，尝试从Normal样式获取
            try:
                normal_fonts = self._get_normal_style_fonts_for_extraction()
                if normal_fonts.get('size_pt'):
                    if hasattr(normal_fonts['size_pt'], 'pt'):
                        return normal_fonts['size_pt'].pt
                    else:
                        return float(normal_fonts['size_pt'])
            except:
                pass
            return 10.5

    def _extract_comprehensive_template_styles(self, template_cell: _Cell, template_table: Table):
        """
        全面提取模板单元格样式，包括：
        1. 单元格级别的直接样式
        2. 表格级别的样式
        3. 表格样式中的样式定义
        4. 默认样式

        优化版本：增强样式检测和调试信息
        """
        extracted_styles = {
            'rPr': None,      # 文本样式
            'pPr': None,      # 段落样式
            'tcPr': None,     # 单元格样式
            'borders': None,  # 边框样式
            'vAlign': None    # 垂直对齐
        }

        try:
            # 1. 首先查找单元格级别的直接样式
            template_tcPr = template_cell._tc.tcPr

            # 2. 查找段落和文本样式 - 增强版本
            print(f"    🔍 开始提取模板单元格样式...")
            paragraph_count = 0
            run_count = 0

            for p_element in template_cell._tc.iterfind(qn('w:p')):
                paragraph_count += 1
                if extracted_styles['pPr'] is None:
                    pPr = p_element.find(qn('w:pPr'))
                    if pPr is not None:
                        extracted_styles['pPr'] = pPr
                        print(f"        ✅ 找到段落样式 (pPr) - 段落 {paragraph_count}")

                if extracted_styles['rPr'] is None:
                    for r_element in p_element.iterfind(qn('w:r')):
                        run_count += 1
                        rPr = r_element.find(qn('w:rPr'))
                        if rPr is not None:
                            extracted_styles['rPr'] = rPr
                            print(f"        ✅ 找到文本样式 (rPr) - 运行 {run_count}")
                            break

                if extracted_styles['pPr'] is not None and extracted_styles['rPr'] is not None:
                    break

            print(f"        📊 扫描了 {paragraph_count} 个段落，{run_count} 个文本运行")

            # 如果仍未找到样式，输出调试信息
            if extracted_styles['rPr'] is None:
                print(f"        ⚠️ 未找到文本样式 (rPr)")
                # 调用调试方法
                self._debug_template_styles(template_cell, "模板单元格")
            if extracted_styles['pPr'] is None:
                print(f"        ⚠️ 未找到段落样式 (pPr)")

            # 3. 查找表格样式
            table_style = self._find_table_style(template_table)

            # 4. 如果单元格级别没有找到样式，从表格样式中提取 - 增强版本
            if extracted_styles['rPr'] is None and table_style is not None:
                # 查找表格样式中的文本样式
                rPr = table_style.find(qn('w:rPr'))
                if rPr is not None:
                    extracted_styles['rPr'] = rPr
                    print(f"    ✅ 从表格样式中找到文本样式")

                    # 调试表格样式字体信息
                    sz = rPr.find(qn('w:sz'))
                    if sz is not None:
                        font_size = int(sz.get(qn('w:val'))) / 2
                        print(f"        📏 表格样式字号: {font_size}pt")
                else:
                    # 如果表格样式也没有文本样式，检查默认样式
                    print(f"    ⚠️ 表格样式中也未找到文本样式，检查默认样式")
                    self._debug_table_style_font_size(table_style)

            if extracted_styles['pPr'] is None and table_style is not None:
                # 查找表格样式中的段落样式
                pPr = table_style.find(qn('w:pPr'))
                if pPr is not None:
                    extracted_styles['pPr'] = pPr
                    print(f"    ✅ 从表格样式中找到段落样式")

            # 5. 查找边框样式（优先级：单元格 > 表格样式 > 默认）
            if template_tcPr is not None:
                borders = template_tcPr.find(qn('w:tcBorders'))
                if borders is not None:
                    extracted_styles['borders'] = borders
                elif table_style is not None:
                    # 从表格样式中查找边框
                    tblBorders = table_style.find(qn('w:tblBorders'))
                    if tblBorders is not None:
                        # 转换表格边框为单元格边框
                        extracted_styles['borders'] = tblBorders
                        print(f"    ✅ 从表格样式中找到边框样式")

            # 6. 查找垂直对齐
            if template_tcPr is not None:
                vAlign = template_tcPr.find(qn('w:vAlign'))
                if vAlign is not None:
                    extracted_styles['vAlign'] = vAlign

            return extracted_styles

        except Exception as e:
            print(f"    ❌ 提取样式时出错: {e}")
            return extracted_styles

    def _apply_cell_style(self, source_cell: _Cell, template_cell: _Cell):
        """
        简化版：提取并应用模板单元格样式到目标单元格
        """
        try:
            # 获取模板表格引用（通过单元格的父表格）
            template_table = None
            for table in self.template_doc.tables:
                try:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell._tc == template_cell._tc:
                                template_table = table
                                break
                        if template_table:
                            break
                except:
                    continue
                if template_table:
                    break

            # 提取有效字体样式字典
            font_style_dict = self._extract_effective_font_style_dict(template_cell, template_table)

            # 应用字体样式到目标单元格
            self._apply_font_style_dict_to_cell(source_cell, font_style_dict)

            # 提取和应用其他样式（边框、垂直对齐等）
            template_styles = self._extract_comprehensive_template_styles(template_cell, template_table)
            template_borders = template_styles['borders']
            template_v_align = template_styles['vAlign']

            # 应用边框样式
            target_tcPr = source_cell._tc.get_or_add_tcPr()
            applied_borders = template_borders
            if applied_borders is None:
                applied_borders = self._create_default_complete_borders()

            old_borders = target_tcPr.find(qn('w:tcBorders'))
            if old_borders is not None:
                target_tcPr.remove(old_borders)
            target_tcPr.append(copy.deepcopy(applied_borders))

            # 应用垂直对齐
            if template_v_align is not None:
                old_v_align = target_tcPr.find(qn('w:vAlign'))
                if old_v_align is not None:
                    target_tcPr.remove(old_v_align)
                target_tcPr.append(copy.deepcopy(template_v_align))

            # 清除段落缩进等属性
            for para in source_cell.paragraphs:
                target_pPr = para._p.get_or_add_pPr()
                old_indent = target_pPr.find(qn('w:ind'))
                if old_indent is not None:
                    target_pPr.remove(old_indent)
                old_jc = target_pPr.find(qn('w:jc'))
                if old_jc is not None:
                    target_pPr.remove(old_jc)

        except Exception as e:
            print(f"WARNING: 应用单元格样式时出错: {e}")

    # =================================================================================
    # |                           【核心修改区域 END】                                  |
    # =================================================================================

    def _clear_all_cell_styles(self, cell: _Cell, preserve_merge_structure=True):
        """
        清除单元格中的所有样式，包括段落缩进、字体样式、对齐方式等。
        修复版本：可选择是否保留合并单元格结构(gridSpan, vMerge)
        保留文本内容，移除所有格式。

        Args:
            preserve_merge_structure: 是否保留合并单元格结构属性
        """
        try:
            # 清除单元格级别的样式
            target_tcPr = cell._tc.get_or_add_tcPr()

            # 根据参数决定是否清除合并单元格属性
            if not preserve_merge_structure:
                # 清除合并单元格相关的属性
                grid_span = target_tcPr.find(qn('w:gridSpan'))
                if grid_span is not None:
                    target_tcPr.remove(grid_span)

                v_merge = target_tcPr.find(qn('w:vMerge'))
                if v_merge is not None:
                    target_tcPr.remove(v_merge)
            else:
                # 保留合并结构时，记录发现的合并属性（调试用）
                grid_span = target_tcPr.find(qn('w:gridSpan'))
                v_merge = target_tcPr.find(qn('w:vMerge'))
                if grid_span is not None or v_merge is not None:
                    pass  # 保留合并结构

            # 移除所有其他单元格属性（背景色等），但保留边框和合并结构
            elements_to_remove = []
            for child in list(target_tcPr):
                if preserve_merge_structure and child.tag in [qn('w:gridSpan'), qn('w:vMerge'), qn('w:tcBorders')]:
                    continue  # 保留合并属性和边框
                elements_to_remove.append(child)

            for child in elements_to_remove:
                target_tcPr.remove(child)

            # 清除段落级别的样式
            for para in cell.paragraphs:
                target_pPr = para._p.get_or_add_pPr()
                # 移除所有段落属性，包括缩进、对齐等
                for child in list(target_pPr):
                    target_pPr.remove(child)

                # 清除文本级别的样式
                for run in para.runs:
                    target_rPr = run._r.get_or_add_rPr()
                    # 移除所有文本属性，包括字体、字号、颜色等
                    for child in list(target_rPr):
                        target_rPr.remove(child)

        except Exception as e:
            print(f"⚠️ 清除单元格样式时发生错误: {e}")

    def _clear_all_table_styles(self, doc: Document):
        """
        清除文档中所有表格的样式，创建一个完全无格式的表格文档。
        """
        try:
            if not doc.tables:
                print("⚠️ 文档中未找到任何表格。")
                return

            print("🧹 开始清除所有表格样式...")
            total_cells = 0

            for table_idx, table in enumerate(doc.tables):
                print(f"  ➤ 正在清除第 {table_idx + 1} 个表格的样式...")

                # 清除表格级别的样式（完全清除包括宽度）
                table_pr = table._tbl.tblPr
                if table_pr is not None:
                    # 完全移除所有表格属性，包括宽度，确保后续完全按模板重建
                    for child in list(table_pr):
                        table_pr.remove(child)
                    print(f"        ✓ 完全清除表格级别属性 (包括tblW)")

                # 使用更安全的方法直接访问XML中的单元格，避免合并单元格导致的遍历错误
                table_element = table._tbl
                for row_element in table_element.findall(qn('w:tr')):
                    for cell_element in row_element.findall(qn('w:tc')):
                        # 创建一个临时的_Cell对象来处理样式清除
                        from docx.table import _Cell
                        temp_cell = _Cell(cell_element, table)
                        self._clear_all_cell_styles(temp_cell, preserve_merge_structure=True)
                        total_cells += 1

            print(f"    ✅ 已清除 {total_cells} 个单元格的所有样式。")

        except Exception as e:
            print(f"❌ 清除表格样式时发生错误: {e}")
            import traceback
            traceback.print_exc()

    def _analyze_table_merge_structure(self, table: Table):
        """
        分析表格的合并单元格结构，返回合并信息
        返回格式：{
            'grid_spans': [(row_idx, col_idx, span_count), ...],  # 水平合并
            'v_merges': [(row_idx, col_idx, merge_type), ...]      # 垂直合并
        }
        """
        try:
            merge_info = {
                'grid_spans': [],  # gridSpan: 水平合并
                'v_merges': []     # vMerge: 垂直合并
            }

            table_element = table._tbl
            for row_idx, row_element in enumerate(table_element.findall(qn('w:tr'))):
                for col_idx, cell_element in enumerate(row_element.findall(qn('w:tc'))):
                    tcPr = cell_element.find(qn('w:tcPr'))
                    if tcPr is not None:
                        # 检查水平合并 (gridSpan)
                        grid_span = tcPr.find(qn('w:gridSpan'))
                        if grid_span is not None:
                            span_val = int(grid_span.get(qn('w:val'), '1'))
                            if span_val > 1:
                                merge_info['grid_spans'].append((row_idx, col_idx, span_val))

                        # 检查垂直合并 (vMerge)
                        v_merge = tcPr.find(qn('w:vMerge'))
                        if v_merge is not None:
                            merge_attr = v_merge.get(qn('w:val'))
                            if merge_attr == 'restart':
                                merge_info['v_merges'].append((row_idx, col_idx, 'restart'))
                            elif merge_attr is None or merge_attr == 'continue':
                                merge_info['v_merges'].append((row_idx, col_idx, 'continue'))

            return merge_info

        except Exception as e:
            print(f"        ⚠️ 分析表格合并结构失败: {e}")
            return {'grid_spans': [], 'v_merges': []}

    def _apply_merge_attributes(self, source_cell: _Cell, template_merge_info, source_cell_idx, source_total_cells):
        """
        根据模板的合并信息，将合并属性应用到源单元格
        """
        try:
            source_tcPr = source_cell._tc.get_or_add_tcPr()

            # 计算源单元格在二维表格中的位置
            source_table_element = source_cell._tc.getparent().getparent()
            current_row_idx = 0
            current_col_idx = 0
            cell_count = 0

            for row_idx, row_element in enumerate(source_table_element.findall(qn('w:tr'))):
                row_cell_count = len(row_element.findall(qn('w:tc')))
                if cell_count + row_cell_count > source_cell_idx:
                    current_row_idx = row_idx
                    current_col_idx = source_cell_idx - cell_count
                    break
                cell_count += row_cell_count

            # 应用水平合并 (gridSpan)
            for template_grid_span in template_merge_info['grid_spans']:
                tmpl_row, tmpl_col, tmpl_span = template_grid_span
                # 使用模数运算循环应用模板的合并模式
                if current_row_idx % len(template_merge_info.get('template_rows', [0])) == tmpl_row and \
                        current_col_idx % tmpl_span == tmpl_col % tmpl_span:
                    # 移除现有的gridSpan
                    existing_grid_span = source_tcPr.find(qn('w:gridSpan'))
                    if existing_grid_span is not None:
                        source_tcPr.remove(existing_grid_span)
                    # 添加新的gridSpan
                    new_grid_span = OxmlElement('w:gridSpan')
                    new_grid_span.set(qn('w:val'), str(tmpl_span))
                    source_tcPr.append(new_grid_span)
                    break

            # 应用垂直合并 (vMerge)
            for template_v_merge in template_merge_info['v_merges']:
                tmpl_row, tmpl_col, tmpl_merge_type = template_v_merge
                if current_row_idx % len(template_merge_info.get('template_rows', [0])) == tmpl_row and \
                        current_col_idx == tmpl_col:
                    # 移除现有的vMerge
                    existing_v_merge = source_tcPr.find(qn('w:vMerge'))
                    if existing_v_merge is not None:
                        source_tcPr.remove(existing_v_merge)
                    # 添加新的vMerge
                    new_v_merge = OxmlElement('w:vMerge')
                    if tmpl_merge_type == 'restart':
                        new_v_merge.set(qn('w:val'), 'restart')
                    # 'continue' 类型不需要设置val属性
                    source_tcPr.append(new_v_merge)
                    break

        except Exception as e:
            print(f"        ⚠️ 应用合并属性失败: {e}")

    def _preserve_table_structure(self, source_table: Table, template_table: Table):
        """
        增强版本：支持合并单元格结构重建
        1. 分析模板表格的合并结构
        2. 使用底层XML遍历确保所有单元格被处理
        3. 精确复制合并属性(gridSpan, vMerge)
        """
        try:
            if not len(source_table.rows): return
            template_rows = template_table.rows
            if not len(template_rows):
                print(f"    ⚠️ 模板表格为空，无法应用样式。")
                return

            total_cells = 0

            print(f"    🔍 分析模板表格合并结构...")
            # 分析模板表格的合并结构
            template_merge_info = self._analyze_table_merge_structure(template_table)

            if template_merge_info['grid_spans'] or template_merge_info['v_merges']:
                print(f"        📋 发现模板合并结构:")
                print(f"            - 水平合并(gridSpan): {len(template_merge_info['grid_spans'])} 处")
                print(f"            - 垂直合并(vMerge): {len(template_merge_info['v_merges'])} 处")
            else:
                print(f"        📋 模板表格无合并结构，使用常规处理")

            # 获取模板单元格的扁平列表（循环使用）
            template_cells_flat = []
            for template_row in template_rows:
                for template_cell in template_row.cells:
                    template_cells_flat.append(template_cell)

            if not template_cells_flat:
                print(f"    ⚠️ 模板表格没有可用的单元格，跳过样式应用。")
                return

            print(f"    📋 使用底层XML遍历，重建合并单元格结构...")

            # 使用底层XML遍历源表格的所有单元格（包括被合并单元格"遮盖"的单元格）
            source_table_element = source_table._tbl
            source_cells_flat = []

            for row_element in source_table_element.findall(qn('w:tr')):
                for cell_element in row_element.findall(qn('w:tc')):
                    from docx.table import _Cell
                    temp_cell = _Cell(cell_element, source_table)
                    source_cells_flat.append(temp_cell)

            # 增强模板合并信息，包含模板行数用于循环计算
            template_merge_info['template_rows'] = len(template_rows)

            # 一对一应用样式：源单元格 → 模板单元格（循环使用模板）
            for i, source_cell in enumerate(source_cells_flat):
                template_cell = template_cells_flat[i % len(template_cells_flat)]

                # 应用基础样式
                self._apply_cell_style(source_cell, template_cell)

                # 应用合并属性
                self._apply_merge_attributes(source_cell, template_merge_info, i, len(source_cells_flat))

                total_cells += 1

            print(f"    ✅ 已应用样式到 {total_cells} 个单元格（包括合并结构重建）")

        except Exception as e:
            print(f"    ❌ 表格结构保持失败: {e}")
            import traceback
            traceback.print_exc()

    def format_document(self, source_document_path: str, output_document_path: str):
        try:
            source_doc = Document(source_document_path)
            if not source_doc.tables:
                print("⚠️ 源文件中未找到任何表格。")
                return

            print(f"📄 成功加载源文件: '{source_document_path}'")

            # 步骤1：清除所有表格样式，创建无格式中间文件
            print("\n" + "=" * 50)
            print("🔄 步骤1: 清除所有表格样式，创建无格式文档")
            print("=" * 50)

            # 创建文档副本用于清除样式
            cleaned_doc = Document(source_document_path)
            self._clear_all_table_styles(cleaned_doc)

            # 保存无格式的中间文件
            print(f"\n💾 正在保存无格式中间文档到: '{CLEANED_DOC_PATH}'...")
            Path(CLEANED_DOC_PATH).parent.mkdir(parents=True, exist_ok=True)
            cleaned_doc.save(CLEANED_DOC_PATH)
            print(f"✅ 无格式中间文档已保存至: '{CLEANED_DOC_PATH}'")

            # 步骤2：基于无格式文档应用模板样式
            print("\n" + "=" * 50)
            print("🎨 步骤2: 基于模板样式格式化表格")
            print("=" * 50)

            print(f"🕵️‍♀️ 正在扫描无格式文档，按页面方向对表格进行分类...")
            categorized_source = self._categorize_tables_by_orientation(cleaned_doc)
            source_vertical_tables = categorized_source['vertical']
            source_horizontal_tables = categorized_source['horizontal']
            print(
                f"📋 找到了 {len(source_vertical_tables)} 个纵向和 {len(source_horizontal_tables)} 个横向表格需要格式化...")

            formatted_count = 0
            if self.template_vertical_tables:
                print("\n--- 开始处理纵向页面表格 ---")
                # 统一使用第一个纵向模板
                first_vertical_template = self.template_vertical_tables[0]
                print(f"  📋 所有纵向表格将统一使用第一个纵向模板样式")
                for i, source_table in enumerate(source_vertical_tables):
                    print(f"  ➤ 正在处理第 {i + 1} 个纵向表格...")
                    # 先表级样式（包括无条件字体刷法）
                    self._apply_table_style(source_table, first_vertical_template)

                    # 只有同构时才复制合并结构（字体已经在上面的fallback中刷了）
                    if (len(source_table.rows) == len(first_vertical_template.rows) and
                            self._count_cols(source_table._tbl) == self._count_cols(first_vertical_template._tbl)):
                        self._preserve_table_structure(source_table, first_vertical_template)
                    else:
                        print("    ↩️ 行列不一致：跳过模板合并重建，保留源表原始合并")
                    formatted_count += 1

            if self.template_horizontal_tables:
                print("\n--- 开始处理横向页面表格 ---")
                # 统一使用第一个横向模板
                first_horizontal_template = self.template_horizontal_tables[0]
                print(f"  📋 所有横向表格将统一使用第一个横向模板样式")
                for i, source_table in enumerate(source_horizontal_tables):
                    print(f"  ➤ 正在处理第 {i + 1} 个横向表格...")
                    # 先表级样式（包括无条件字体刷法）
                    self._apply_table_style(source_table, first_horizontal_template)

                    # 只有同构时才复制合并结构（字体已经在上面的fallback中刷了）
                    if (len(source_table.rows) == len(first_horizontal_template.rows) and
                            self._count_cols(source_table._tbl) == self._count_cols(first_horizontal_template._tbl)):
                        self._preserve_table_structure(source_table, first_horizontal_template)
                    else:
                        print("    ↩️ 行列不一致：跳过模板合并重建，保留源表原始合并")
                    formatted_count += 1

            if formatted_count > 0:
                print(f"\n💾 正在保存最终格式化文档到: '{output_document_path}'...")
                Path(output_document_path).parent.mkdir(parents=True, exist_ok=True)
                cleaned_doc.save(output_document_path)
                print(f"\n🎉 成功！格式化后的文件已保存至: '{output_document_path}'")
                print(f"\n📁 文件处理流程:")
                print(f"   1️⃣ 原文件: {source_document_path}")
                print(f"   2️⃣ 无格式中间文件: {CLEANED_DOC_PATH}")
                print(f"   3️⃣ 最终格式化文件: {output_document_path}")
            else:
                print("\n🤷‍♀️ 本次运行没有成功格式化任何表格。")

        except Exception as e:
            print(f"❌ 处理文档时发生错误: {e}")
            import traceback
            traceback.print_exc()


def main():
    print("🚀 开始执行增强版表格样式刷脚本...");
    print("=" * 60)
    print("功能:")
    print("  1️⃣ Excel对象就地转换（如果有）")
    print("  2️⃣ 表格样式清理和格式化")
    print("=" * 60)

    # 检查必要文件
    if not Path(TEMPLATE_DOC_PATH).exists():
        print(f"❌ 错误：模板文件 '{TEMPLATE_DOC_PATH}' 不存在。")
        return

    if not Path(SOURCE_DOC_PATH).exists():
        print(f"❌ 错误：源文件 '{SOURCE_DOC_PATH}' 不存在。")
        return

    # 步骤0：Excel对象就地转换
    excel_converted_path = convert_excel_objects_first(SOURCE_DOC_PATH, EXCEL_CONVERTED_DOC_PATH)

    # 步骤1：表格样式格式化
    print(f"\n🎨 开始表格样式处理...")
    print(f"📄 输入文档: {excel_converted_path}")

    brush = TableStyleBrush(TEMPLATE_DOC_PATH)
    brush.format_document(excel_converted_path, OUTPUT_DOC_PATH)

    # 输出处理总结
    print("\n" + "=" * 60)
    print("📋 处理流程总结:")
    print(f"  📁 原始文档: {SOURCE_DOC_PATH}")
    if excel_converted_path != SOURCE_DOC_PATH:
        print(f"  🔄 Excel转换后: {excel_converted_path}")
    print(f"  🎨 最终格式化文档: {OUTPUT_DOC_PATH}")
    print("=" * 60)
    print("✅ 脚本执行完毕。")


if __name__ == "__main__":
    main()