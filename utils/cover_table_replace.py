# 添加lxml导入处理，与converter.py中的一致
try:
    from lxml import etree
    LXML_AVAILABLE = True
except ImportError:
    LXML_AVAILABLE = False
    etree = None

# 导入Document类（忽略静态分析工具的警告）
# pyright: reportAttributeAccessIssue=false
try:
    from docx import Document
except ImportError:
    Document = None

import copy
import re
from typing import Optional, Any

def _normalize_marker(s: str) -> str:
    """规范化比较用的字符串：去掉所有空白字符（包括全角空格）、转换小写。"""
    if s is None:
        return ""
    # \s 包括半角空白、换行、制表符；再去掉常见的全角空格 \u3000
    return re.sub(r'\s+', '', s).replace('\u3000', '').lower()

def _find_paragraph_index_with_marker(doc: Any, marker: str):
    """
    在 doc 中查找含有 marker 的段落（normalize 后包含关系），
    返回段落在 body 子元素中的索引（不是 paragraph list 的索引），
    若未找到返回 None。
    """
    norm_marker = _normalize_marker(marker)
    # doc.paragraphs 不能直接给出 body 元素索引（因为 body 包含表格等），
    # 所以遍历 body 的子元素并对 p 标签判断其文本是否包含 marker。
    body = doc.element.body
    for idx, child in enumerate(body):
        # 段落的标签通常以 'p' 结尾
        if child.tag.endswith('p'):
            # 将该段落的完整文本取出（类似 docx para.text）
            # 使用 python-docx 的 Paragraph wrapper 更方便获取 .text，
            # 但我们只有 element；构造 Paragraph 需要 parent doc. 简单方法：使用 child.itertext()
            para_text = ''.join(child.itertext())
            if _normalize_marker(para_text).find(norm_marker) != -1:
                return idx
    return None

def _find_first_table_after_index(doc: Any, start_body_index: int):
    """
    从 body 的 start_body_index 之后寻找第一个表格元素（tbl），
    返回对应的 table element（lxml element）和它在 body 中的索引。
    如果找不到，返回 (None, None)
    """
    body = doc.element.body
    for idx in range(start_body_index + 1, len(body)):
        child = body[idx]
        if child.tag.endswith('tbl'):
            return child, idx
    return None, None

def replace_table_after_marker(source_path: str, target_path: str,
                               marker: str = "各专业参加设计人员名单",
                               save_path: Optional[str] = None):
    """
    将 source_path 中 marker 后的第一个表格（保持结构/格式）复制到 target_path，
    替换 target_path 中 marker 后的第一个表格。保存为 save_path（不指定则在 target_path 同目录创建 *_replaced.docx）。
    """
    # 检查lxml是否可用
    if not LXML_AVAILABLE or etree is None:
        raise ImportError("lxml库不可用，请安装lxml: pip install lxml")
    
    # 检查Document是否可用
    if Document is None:
        raise ImportError("python-docx库不可用，请安装python-docx: pip install python-docx")
    
    # 读取文档
    src_doc = Document(source_path)
    tgt_doc = Document(target_path)

    # 在两个文档中查找段落索引
    src_para_idx = _find_paragraph_index_with_marker(src_doc, marker)
    if src_para_idx is None:
        raise ValueError(f"在源文档 '{source_path}' 中未找到标记（marker）：{marker!r}")

    tgt_para_idx = _find_paragraph_index_with_marker(tgt_doc, marker)
    if tgt_para_idx is None:
        raise ValueError(f"在目标文档 '{target_path}' 中未找到标记（marker）：{marker!r}")

    # 在两个文档中找到第一个表格（lxml element）位于 marker 后
    src_tbl_elem, src_tbl_body_idx = _find_first_table_after_index(src_doc, src_para_idx)
    if src_tbl_elem is None:
        raise ValueError(f"在源文档 '{source_path}' 中，标记后未找到表格。")

    tgt_tbl_elem, tgt_tbl_body_idx = _find_first_table_after_index(tgt_doc, tgt_para_idx)
    if tgt_tbl_elem is None:
        raise ValueError(f"在目标文档 '{target_path}' 中，标记后未找到表格。")

    # 深拷贝源表格的 XML 元素，避免从源文档移动它
    # lxml 的元素可以用 copy.deepcopy 复制
    src_tbl_copy = copy.deepcopy(src_tbl_elem)

    # 将复制的表格插入到目标表格之前或之后，然后删除旧表格
    # 我们使用 addnext 在旧表格后插入副本，然后移除旧表格
    tgt_parent = tgt_tbl_elem.getparent()
    # 插入复制（注意：要插入到目标文档的命名空间下）
    # 直接 addnext copy 会把 copy 的 nsmap 保留，这通常可以工作
    tgt_tbl_elem.addnext(src_tbl_copy)
    # 删除原表格
    tgt_parent.remove(tgt_tbl_elem)

    # 保存
    if not save_path:
        if target_path.lower().endswith('.docx'):
            save_path = target_path[:-5] + "_replaced.docx"
        else:
            save_path = target_path + "_replaced.docx"

    tgt_doc.save(save_path)
    return save_path

if __name__ == "__main__":
    # 示例调用（修改为你的实际路径）
    source_docx = r"C:\path\to\文档A.docx"
    target_docx = r"C:\path\to\文档B.docx"
    try:
        out = replace_table_after_marker(source_docx, target_docx,
                                         marker="各专业参加设计人员名单")
        print("替换成功，保存为：", out)
    except Exception as e:
        print("错误：", e)