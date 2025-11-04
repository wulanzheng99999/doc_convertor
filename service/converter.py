"""
完整文档格式化转换器

实现从源文档到格式化文档的完整转换流程:
1. 文档拆分 - 分离封面和正文
2. Pandoc转换 - 使用模板格式化正文(如果没有指定，默认使用template目录下的reference.docx)
3. 表格格式化 - 使用mcp服务格式化正文的表格（还没做到这一步，现在不管）
4. 文档合并 - 重新合并为完整文档
5. 标题修改 - 修改合并后的文档的目录标题
6. 图片格式化 - 图片居中，单倍行距
7. 补充处理 - 将库号右靠齐
8. 在目录之后插入分节符
9. 处理文档节的页码设置
10. 删除文件中中所有的突出显示

# 页眉替换 - 设置指定页眉内容 // 直接替换模板文件的，现在不管

"""

import os
import sys
import tempfile
import shutil
import time
import zipfile
import subprocess
from datetime import datetime
from typing import Optional, Tuple
from pathlib import Path

# 导入所需的工具模块
from utils.pandoc_converter import PandocConverter
from utils.docx_split import DocxSplitProcessor
from utils.docx_merge import copy_all_to_beginning
from utils.docx_update_toc_title import update_toc_title_xml
from docx.oxml.ns import qn

# COM 组件可用性检测（用于页面方向拆分等功能）
try:
    import pythoncom
    import win32com.client as win32
    from win32com.client import constants as C

    COM_AVAILABLE = True
except Exception:
    pythoncom = None
    win32 = None
    C = None
    COM_AVAILABLE = False

# 添加项目路径
current_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(current_dir)
sys.path.insert(0, parent_dir)

# 添加lxml导入用于去除突出显示
try:
    from lxml import etree

    LXML_AVAILABLE = True
except ImportError:
    LXML_AVAILABLE = False
    etree = None


def _kill_existing_winword_processes():
    if os.name != "nt":
        return
    cmd = ["taskkill", "/f", "/im", "WINWORD.EXE"]
    creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
    try:
        subprocess.run(
            cmd,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            check=False,
            creationflags=creationflags,
        )
    except Exception:
        pass


def _wait_for_file_release_for_step(path: str, timeout: float = 10.0, interval: float = 0.3) -> bool:
    """
    Wait until the file can be opened for read/write. Returns True on success.
    """
    if not path:
        return True
    target = Path(path)
    deadline = time.time() + timeout
    while True:
        try:
            if target.exists():
                with open(target, "a+b"):
                    return True
            else:
                return True
        except OSError:
            if time.time() >= deadline:
                return False
            time.sleep(interval)


class DocumentConverter:
    """文档格式化转换器"""

    def __init__(self, document_type: int = 1):
        """初始化转换器"""
        self.temp_dir = None
        self.pandoc_converter = None
        self.intermediate_files = {}  # 保存中间文件路径
        self.debug_output_dir = os.path.join(parent_dir, 'temp')  # 中间文件保存目录
        self.save_intermediate_files = False  # 是否保存中间文件的开关
        self.document_type = document_type  # 文档类型参数
        self.extracted_footer_content = None  # 存储提取的页脚内容
        self.step0_processed_source = None  # 记录步骤0输出文件

    def __enter__(self):
        """上下文管理器入口"""
        _kill_existing_winword_processes()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        """上下文管理器出口，清理临时文件"""
        self.cleanup()

    def cleanup(self):
        """清理临时文件"""
        if self.temp_dir and os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir, ignore_errors=True)

    def _append_document_with_formatting(
        self,
        base_file: str,
        file_to_append: str,
        output_file: str,
    ) -> bool:
        """将 file_to_append 追加到 base_file 末尾（保留页面方向）。"""
        try:
            from utils.merge_docs_preserve_orientation import (
                append_file_with_formatting as merge_append,
            )
        except ImportError as err:
            print(f"[warn] 无法导入 merge_docs_preserve_orientation: {err}")
            return False

        try:
            merge_append(base_file, file_to_append, output_file)
            return True
        except Exception as exc:
            print(f"[warn] 横版文档追加失败: {exc}")
            import traceback

            traceback.print_exc()
            return False


    def validate_input_files(self, source_file: str, template_file: str) -> bool:
        """
        验证输入文件的有效性

        Args:
            source_file: 源文档路径
            template_file: 模板文档路径

        Returns:
            bool: 文件是否有效
        """
        if os.name == 'nt':
            _kill_existing_winword_processes()

        if not os.path.exists(source_file):
            print(f"❌ 源文档不存在: {source_file}")
            return False

        if not os.path.exists(template_file):
            print(f"❌ 模板文档不存在: {template_file}")
            return False

        # 检查文件是否为DOCX格式
        if not source_file.lower().endswith('.docx'):
            print(f"❌ 源文档不是DOCX格式: {source_file}")
            return False

        if not template_file.lower().endswith('.docx'):
            print(f"❌ 模板文档不是DOCX格式: {template_file}")
            return False

        return True

    def _save_intermediate_file(self, source_path: str, step_name: str, file_description: str = "") -> None:
        """
        保存中间文件到指定目录便于查看和调试

        Args:
            source_path: 源文件路径
            step_name: 步骤名称（如 step1_split, step2_pandoc 等）
            file_description: 文件描述（如封面, 正文内容 等）
        """
        # 如果不保存中间文件，则直接返回
        if not self.save_intermediate_files:
            return

        try:
            # 确保输出目录存在
            if not os.path.exists(self.debug_output_dir):
                os.makedirs(self.debug_output_dir, exist_ok=True)
                print(f"📁 创建调试输出目录: {self.debug_output_dir}")

            # 生成带时间戳的文件名
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            base_name = os.path.splitext(os.path.basename(source_path))[0]
            if file_description:
                debug_filename = f"{step_name}_{file_description}_{timestamp}.docx"
            else:
                debug_filename = f"{step_name}_{base_name}_{timestamp}.docx"

            debug_path = os.path.join(self.debug_output_dir, debug_filename)

            # 复制文件
            shutil.copy2(source_path, debug_path)

            print(f"   💾 已保存调试文件: {debug_filename}")

        except Exception as e:
            print(f"   ⚠️ 保存调试文件失败: {str(e)}")

    def step0_preprocess_headings(self, source_file: str) -> str:
        """
        步骤-1: 预处理标题样式 - 根据编号规则自动设置标题样式

        Args:
            source_file: 源文档路径

        Returns:
            str: 预处理后的文档路径
        """
        print("-" * 50)
        print("📝 步骤0: 预处理标题样式")

        try:
            import docx

            # 确保临时目录存在
            if not self.temp_dir:
                raise ValueError("临时目录未初始化")

            # 生成预处理后的文件路径
            base_name = os.path.splitext(os.path.basename(source_file))[0]
            preprocessed_path = os.path.join(self.temp_dir, f"{base_name}_预处理标题.docx")

            # 复制源文件到临时文件
            shutil.copy2(source_file, preprocessed_path)

            # 加载文档并处理标题
            doc = docx.Document(preprocessed_path)

            for para in doc.paragraphs:
                # 获取段落文本
                text = para.text.strip()

                # 判断标题的层级（通过编号规则进行判断）
                if text and text[0].isdigit() and '.' in text:  # 以数字开始且包含.
                    if '，' not in text and '。' not in text and '；' not in text:  # 不包含，或者。
                        # 按点分割序号，判断层级
                        level = text.count('.') + 1

                        # 根据层级应用相应的样式
                        if level == 1:
                            para.style = 'Heading 1'  # 一级标题
                        elif level == 2:
                            para.style = 'Heading 2'  # 二级标题
                        elif level == 3:
                            para.style = 'Heading 3'  # 三级标题
                        elif level == 4:
                            para.style = 'Heading 4'  # 四级标题

            # 保存修改后的文档
            doc.save(preprocessed_path)
            print("✅ 标题预处理完成!")

            # 保存中间文件
            if self.save_intermediate_files:
                self._save_intermediate_file(preprocessed_path, "step0_preprocess", "标题预处理")

            return preprocessed_path

        except Exception as e:
            print(f"❌ 标题预处理过程中发生错误: {str(e)}")
            return source_file

    def step0_replace_header_footer(self, source_file: str, template_file: str) -> str:
        """
        步骤0: 页眉页脚替换 - 将源文档的页眉页脚内容替换到模板文档中

        Args:
            source_file: 源文档路径（提供页眉页脚内容）
            template_file: 模板文档路径（被替换页眉页脚内容）

        Returns:
            str: 替换页眉页脚后的模板文件路径
        """
        print("-" * 50)
        print("📑 步骤0: 页眉页脚替换")

        try:
            # 确保临时目录存在
            if not self.temp_dir:
                raise ValueError("临时目录未初始化")

            # 生成输出文件路径
            base_name = os.path.splitext(os.path.basename(template_file))[0]
            updated_template_path = os.path.join(self.temp_dir, f"{base_name}_页眉页脚替换后.docx")

            print(f"📄 源文档: {os.path.basename(source_file)}")
            print(f"📄 模板文档: {os.path.basename(template_file)}")
            print(f"📤 更新后模板: {os.path.basename(updated_template_path)}")

            # 使用docx_header_footer_replace.py中的方法进行页眉页脚替换
            try:
                from utils.docx_header_footer_replace import replace_header_footer_content, \
                    extract_header_footer_content
                # 提取源文档的页眉页脚内容
                source_content = extract_header_footer_content(source_file, 2)  # 从源文档第2节提取

                # 保存提取的页脚内容，用于后续步骤
                if "footers" in source_content and source_content["footers"]:
                    # 获取默认类型的页脚内容
                    default_footer = source_content["footers"].get("default", "")
                    even_footer = source_content["footers"].get("even", "")
                    first_footer = source_content["footers"].get("first", "")

                    # 优先使用默认页脚，如果没有则尝试其他类型
                    self.extracted_footer_content = default_footer or even_footer or first_footer
                    print(f"📄 提取的页脚内容: {self.extracted_footer_content}")
                else:
                    print("⚠️ 未提取到页脚内容，将使用默认页脚格式")
                    self.extracted_footer_content = None
                # 执行页眉页脚替换
                success = replace_header_footer_content(
                    source_docx_path=source_file,
                    target_docx_path=template_file,
                    source_section_index=2,  # 从源文档第1节提取
                    target_section_index=1,  # 替换到模板文档第1节
                    save_path=updated_template_path
                )

                if success and os.path.exists(updated_template_path):
                    print("✅ 页眉页脚替换成功!")

                    # 保存中间文件到指定目录便于查看调试
                    if self.save_intermediate_files:
                        print(f"   更新后模板: {os.path.basename(updated_template_path)}")
                        print(f"📁 正在保存step0中间文件到: {self.debug_output_dir}")
                        self._save_intermediate_file(updated_template_path, "step0_header_footer", "替换后模板")

                    return updated_template_path
                else:
                    print("❌ 页眉页脚替换失败，使用原始模板文件")
                    return template_file

            except Exception as replace_error:
                print(f"❌ 页眉页脚替换过程中发生错误: {str(replace_error)}")
                print("   继续使用原始模板文件")
                return template_file

        except Exception as e:
            print(f"❌ 页眉页脚替换过程中发生错误: {str(e)}")
            return template_file

    def step1_split_document(self, source_file: str) -> Tuple[Optional[str], Optional[str]]:
        """
        步骤1: 文档拆分 - 将源文档分离为封面和正文

        Args:
            source_file: 源文档路径

        Returns:
            Tuple[str, str]: (封面文件路径, 不含目录的正文内容文件路径)

        说明:
            - 封面文档：只包含目录之前的内容（如封面、标题等），不包含目录本身
            - 正文内容文档：包含从目录之后开始的所有内容（不包括目录）
        """
        print("-" * 50)
        print("📑 步骤1: 文档拆分")

        try:
            # 确保临时目录存在
            if not self.temp_dir:
                raise ValueError("临时目录未初始化")

            # 生成输出文件路径
            base_name = os.path.splitext(os.path.basename(source_file))[0]
            cover_toc_path = os.path.join(self.temp_dir, f"{base_name}_封面.docx")
            content_no_toc_path = os.path.join(self.temp_dir, f"{base_name}_正文内容.docx")

            print(f"📄 源文档: {os.path.basename(source_file)}")
            print(f"📤 封面输出: {os.path.basename(cover_toc_path)}")
            print(f"📤 正文内容输出: {os.path.basename(content_no_toc_path)}")

            # 使用整合的拆分方法
            processor = DocxSplitProcessor()
            success = processor.split_document_for_conversion(
                source_path=source_file,
                cover_toc_output=cover_toc_path,
                content_no_toc_output=content_no_toc_path,
                toc_keywords=['目录', '目 录', '目  录', '目   录', '目    录', '目     录', '目      录',
                              '目       录', '目        录', '目         录', 'Contents', 'TABLE OF CONTENTS']
            )

            if success and os.path.exists(cover_toc_path) and os.path.exists(content_no_toc_path):
                print("✅ 文档拆分成功!")

                # 使用cover_replace.py中的便捷函数处理封面文档
                try:
                    # 导入cover_replace模块
                    from utils.cover_replace import replace_content_in_cover_auto
                    import json
                    from datetime import datetime  # 导入datetime模块

                    # 配置文件路径
                    parent_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

                    # 生成处理后的封面路径
                    processed_cover_path = os.path.join(self.temp_dir, f"{base_name}_封面_处理后.docx")

                    # 使用自动选择模板和配置文件的函数处理封面
                    print("🔧 使用cover_replace_auto处理封面文档...")
                    actual_path = replace_content_in_cover_auto(
                        source_docx_path=cover_toc_path,  # 使用拆分后的封面作为源
                        output_docx_path=processed_cover_path,
                        document_type=self.document_type,  # 使用文档类型参数
                        save_file=self.save_intermediate_files  # 与convert_document的save_intermediate参数关联
                    )

                    # 如果处理成功，更新cover_toc_path指向处理后的文件
                    if os.path.exists(actual_path):
                        cover_toc_path = actual_path
                        print(f"✅ 封面文档处理成功，使用处理后的文件: {os.path.basename(cover_toc_path)}")

                        # 如果需要保存中间文件，也将处理后的封面文件复制到调试目录
                        if self.save_intermediate_files:
                            processed_cover_filename = f"step1_split_封面处理后_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                            processed_cover_debug_path = os.path.join(self.debug_output_dir, processed_cover_filename)
                            shutil.copy2(cover_toc_path, processed_cover_debug_path)
                            print(f"   💾 已保存处理后的封面文件: {processed_cover_filename}")
                    else:
                        print("⚠️ 封面文档处理失败，使用原始拆分文件")
                except Exception as cover_error:
                    print(f"⚠️ 封面文档处理过程中发生错误: {str(cover_error)}")
                    print("   继续使用原始拆分文件")

                # 使用cover_table_replace.py中的函数替换处理后封面中的表格
                try:
                    from utils.cover_table_replace import replace_table_after_marker
                    from datetime import datetime  # 导入datetime模块

                    # 生成表格替换后的封面路径
                    table_replaced_cover_path = os.path.join(self.temp_dir, f"{base_name}_封面_表格替换后.docx")

                    # 使用源文档作为表格来源，处理后的封面作为目标进行表格替换
                    print("🔧 使用cover_table_replace替换处理后封面中的表格...")
                    replaced_path = replace_table_after_marker(
                        source_path=source_file,  # 使用原始源文档作为表格来源
                        target_path=cover_toc_path,  # 使用处理后的封面作为替换目标
                        marker="各专业参加设计人员名单",  # 使用默认标记
                        save_path=table_replaced_cover_path
                    )

                    # 如果替换成功，更新cover_toc_path指向表格替换后的文件
                    if os.path.exists(replaced_path):
                        cover_toc_path = replaced_path
                        print(f"✅ 封面表格替换成功，使用表格替换后的文件: {os.path.basename(cover_toc_path)}")

                        # 如果需要保存中间文件，也将表格替换后的封面文件复制到调试目录
                        if self.save_intermediate_files:
                            table_replaced_cover_filename = f"step1_split_封面_表格替换后_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                            table_replaced_cover_debug_path = os.path.join(self.debug_output_dir,
                                                                           table_replaced_cover_filename)
                            shutil.copy2(cover_toc_path, table_replaced_cover_debug_path)
                            print(f"   💾 已保存表格替换后的封面文件: {table_replaced_cover_filename}")
                    else:
                        print("⚠️ 封面表格替换失败，使用处理后的封面文件")
                except Exception as table_error:
                    print(f"⚠️ 封面表格替换过程中发生错误: {str(table_error)}")
                    print("   继续使用处理后的封面文件")

                # 对正文内容文档中的Excel表格进行转换处理
                try:
                    from utils.docx_table_excel import convert_embedded_excels_inplace
                    from datetime import datetime  # 导入datetime模块

                    # 生成处理后的正文内容路径
                    processed_content_path = os.path.join(self.temp_dir, f"{base_name}_正文内容_表格处理后.docx")

                    # 使用docx_table_excel处理正文中的Excel表格
                    print("🔧 使用docx_table_excel处理正文内容中的Excel表格...")
                    try:
                        converted_count = convert_embedded_excels_inplace(
                            source_docx=content_no_toc_path,
                            output_docx=processed_content_path,
                            placeholder_when_no_pandas=True
                        )

                        # 如果处理成功，更新content_no_toc_path指向处理后的文件
                        if os.path.exists(processed_content_path) and converted_count >= 0:
                            content_no_toc_path = processed_content_path
                            print(f"✅ 正文内容中的Excel表格处理成功，共转换 {converted_count} 个表格")

                            # 如果需要保存中间文件，也将处理后的正文文件复制到调试目录
                            if self.save_intermediate_files:
                                processed_content_filename = f"step1_split_正文内容_表格处理后_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                                processed_content_debug_path = os.path.join(self.debug_output_dir,
                                                                            processed_content_filename)
                                shutil.copy2(content_no_toc_path, processed_content_debug_path)
                                print(f"   💾 已保存处理后的正文内容文件: {processed_content_filename}")
                        else:
                            print("⚠️ 正文内容中的Excel表格处理失败，使用原始正文文件")
                    except Exception as convert_error:
                        print(f"⚠️ 正文内容中的Excel表格转换失败: {str(convert_error)}")
                        print("   继续使用原始正文文件")
                except Exception as content_error:
                    print(f"⚠️ 正文内容中的Excel表格处理过程中发生错误: {str(content_error)}")
                    print("   继续使用原始正文文件")

                # 保存中间文件到指定目录便于查看调试
                if self.save_intermediate_files:
                    print(f"   封面: {os.path.basename(cover_toc_path)}")
                    print(f"   正文内容: {os.path.basename(content_no_toc_path)}")
                    print(f"📁 正在保存step1中间文件到: {self.debug_output_dir}")
                    self._save_intermediate_file(cover_toc_path, "step1_split", "封面")
                    self._save_intermediate_file(content_no_toc_path, "step1_split", "正文内容")

                # 保存中间文件路径
                self.intermediate_files['cover_toc'] = cover_toc_path
                self.intermediate_files['original_content'] = content_no_toc_path

                return cover_toc_path, content_no_toc_path
            else:
                print("❌ 文档拆分失败")
                return None, None

        except Exception as e:
            print(f"❌ 文档拆分过程中发生错误: {str(e)}")
            return None, None

    def _create_orientation_specific_copy(
        self,
        source_content_path: str,
        output_path: str,
        target_orientation: int,
    ) -> bool:
        """使用 Word COM 仅保留指定方向的节并另存为新文档。"""
        if not COM_AVAILABLE or not win32 or not pythoncom or C is None:
            return False

        word_app = None
        cloned_doc = None
        com_initialized = False
        kept_sections = 0

        source_content_path = os.path.abspath(source_content_path)
        output_path = os.path.abspath(output_path)

        if not os.path.exists(source_content_path):
            return False

        try:
            shutil.copy2(source_content_path, output_path)
        except Exception as copy_error:
            print(f"⚠️ 正文克隆失败: {copy_error}")
            return False

        try:
            pythoncom.CoInitialize()
            com_initialized = True

            word_app = win32.DispatchEx("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = False

            cloned_doc = word_app.Documents.Open(output_path, ReadOnly=False)
            sections_count = cloned_doc.Sections.Count if cloned_doc else 0

            for index in range(sections_count, 0, -1):
                section = cloned_doc.Sections(index)
                orientation = section.PageSetup.Orientation
                if orientation != target_orientation:
                    section.Range.Delete()
                else:
                    kept_sections += 1

            remaining_sections = cloned_doc.Sections.Count if cloned_doc else 0

            if kept_sections > 0 and remaining_sections > 0:
                header_types = [
                    getattr(C, "wdHeaderFooterPrimary", 1),
                    getattr(C, "wdHeaderFooterFirstPage", 2),
                    getattr(C, "wdHeaderFooterEvenPages", 3),
                ]
                footer_types = [
                    getattr(C, "wdHeaderFooterPrimary", 1),
                    getattr(C, "wdHeaderFooterFirstPage", 2),
                    getattr(C, "wdHeaderFooterEvenPages", 3),
                ]
                try:
                    for sec_index in range(1, remaining_sections + 1):
                        section = cloned_doc.Sections(sec_index)
                        for header_type in header_types:
                            try:
                                header = section.Headers(header_type)
                                if header is not None and getattr(header, "Exists", True):
                                    header.LinkToPrevious = False
                            except Exception:
                                continue
                        for footer_type in footer_types:
                            try:
                                footer = section.Footers(footer_type)
                                if footer is not None and getattr(footer, "Exists", True):
                                    footer.LinkToPrevious = False
                            except Exception:
                                continue
                except Exception as link_error:
                    print(f"ℹ️ 取消页眉页脚链接时出现问题: {link_error}")

            content_length = cloned_doc.Content.StoryLength if cloned_doc else 0

            if kept_sections > 0 and content_length > 1:
                cloned_doc.Save()
                return True

            if os.path.exists(output_path):
                try:
                    os.remove(output_path)
                except Exception:
                    pass

            return False
        except Exception as exc:
            print(f"⚠️ 页面方向内容裁剪失败: {exc}")
            return False
        finally:
            if cloned_doc is not None:
                try:
                    cloned_doc.Close(SaveChanges=0)
                except Exception:
                    pass
            if word_app is not None:
                try:
                    word_app.Quit()
                except Exception:
                    pass
            if COM_AVAILABLE and com_initialized:
                try:
                    pythoncom.CoUninitialize()
                except Exception:
                    pass

            if (kept_sections <= 0 or not os.path.exists(output_path) or os.path.getsize(output_path) == 0):
                if os.path.exists(output_path):
                    try:
                        os.remove(output_path)
                    except Exception:
                        pass

    def _split_content_by_orientation_com(
        self,
        source_content_path: str,
        portrait_output_path: str,
        landscape_output_path: str,
    ) -> Tuple[bool, bool]:
        """基于 Word COM 按页面方向拆分正文内容。"""
        print("🔧 开始基于 COM API 按页面方向拆分正文内容...")

        if not COM_AVAILABLE:
            print("❌ COM 库不可用，无法进行页面方向拆分")
            return False, False

        if not source_content_path or not os.path.exists(source_content_path):
            print("⚠️ 正文内容文件不存在，跳过页面方向拆分")
            return False, False

        portrait_success = self._create_orientation_specific_copy(
            source_content_path=source_content_path,
            output_path=portrait_output_path,
            target_orientation=getattr(C, "wdOrientPortrait", 0),
        )

        landscape_success = self._create_orientation_specific_copy(
            source_content_path=source_content_path,
            output_path=landscape_output_path,
            target_orientation=getattr(C, "wdOrientLandscape", 1),
        )

        if not portrait_success and os.path.exists(portrait_output_path):
            try:
                os.remove(portrait_output_path)
            except Exception:
                pass

        if not landscape_success and os.path.exists(landscape_output_path):
            try:
                os.remove(landscape_output_path)
            except Exception:
                pass

        if portrait_success:
            print(f"   ✅ 已保存竖版文件: {os.path.basename(portrait_output_path)}")
        else:
            print("ℹ️ 竖版节不存在或处理失败，继续使用原始正文文件")

        if landscape_success:
            print(f"   ✅ 已保存横版文件: {os.path.basename(landscape_output_path)}")
        else:
            print("ℹ️ 横版节不存在或处理失败，不生成横版文件")

        return portrait_success, landscape_success

    def step1_5_split_by_orientation(self, content_file: str) -> Tuple[str, Optional[str]]:
        """
        步骤1.5: 基于页面方向拆分正文内容。

        Args:
            content_file: 步骤1输出的正文内容文件路径

        Returns:
            Tuple[str, Optional[str]]: (竖版内容文件, 横版内容文件或 None)
        """
        print("-" * 50)
        print("🧭 步骤1.5: 按页面方向拆分正文内容")

        portrait_result = content_file
        landscape_result: Optional[str] = None

        try:
            if not content_file or not os.path.exists(content_file):
                print("⚠️ 正文内容文件不存在，跳过页面方向拆分")
                return portrait_result, landscape_result

            if not self.temp_dir:
                raise ValueError("临时目录未初始化")

            base_name = os.path.splitext(os.path.basename(content_file))[0]
            portrait_path = os.path.join(self.temp_dir, f"{base_name}_竖版内容.docx")
            landscape_path = os.path.join(self.temp_dir, f"{base_name}_横版内容.docx")

            portrait_success, landscape_success = self._split_content_by_orientation_com(
                source_content_path=content_file,
                portrait_output_path=portrait_path,
                landscape_output_path=landscape_path,
            )

            if portrait_success and os.path.exists(portrait_path) and os.path.getsize(portrait_path) > 0:
                portrait_result = portrait_path
                print(f"✅ 将竖版文件用于后续处理: {os.path.basename(portrait_result)}")
                if self.save_intermediate_files:
                    self._save_intermediate_file(portrait_result, "step1_5_split", "竖版内容")
            else:
                if os.path.exists(portrait_path) and os.path.getsize(portrait_path) == 0:
                    try:
                        os.remove(portrait_path)
                    except Exception:
                        pass
                print("⚠️ 竖版拆分失败或无竖版内容，继续使用原始正文文件")

            if landscape_success and os.path.exists(landscape_path) and os.path.getsize(landscape_path) > 0:
                landscape_result = landscape_path
                print(f"✅ 横版文件已保存: {os.path.basename(landscape_result)}")
                if self.save_intermediate_files:
                    self._save_intermediate_file(landscape_result, "step1_5_split", "横版内容")
            else:
                if os.path.exists(landscape_path) and os.path.getsize(landscape_path) == 0:
                    try:
                        os.remove(landscape_path)
                    except Exception:
                        pass
                print("ℹ️ 横版拆分失败或无横版内容")

        except Exception as exc:
            print(f"⚠️ 步骤1.5 处理过程中发生错误: {exc}")
            portrait_result = content_file
            landscape_result = None
        finally:
            self.intermediate_files['portrait_content'] = portrait_result
            if landscape_result:
                self.intermediate_files['landscape_content'] = landscape_result
            else:
                self.intermediate_files.pop('landscape_content', None)

        return portrait_result, landscape_result

    def step2_pandoc_convert(self, content_file: str, template_file: str) -> Optional[str]:
        """
        步骤2: Pandoc转换 - 使用模板文件格式化正�?

        Args:
            content_file: 正文内容文件路径
            template_file: 模板文件路径

        Returns:
            str: Pandoc处理后的文件路径
        """
        print("-" * 50)
        print("🔄 步骤2: Pandoc转换")

        try:
            # 初始化Pandoc转换器
            if self.pandoc_converter is None:
                # 查找pandoc可执行文件
                pandoc_path = self._find_pandoc_executable()
                if not pandoc_path:
                    print("⚠️ Pandoc不可用，跳过Pandoc转换步骤")

                    # 生成一个标记后的文件，表示跳过了Pandoc转换
                    base_name = os.path.splitext(os.path.basename(content_file))[0]
                    if not self.temp_dir:
                        raise ValueError("临时目录未初始化")
                    skipped_output = os.path.join(self.temp_dir, f"{base_name}_跳过Pandoc转换.docx")
                    shutil.copy2(content_file, skipped_output)

                    # 保存中间文件
                    self.intermediate_files['pandoc_converted'] = skipped_output

                    # 保存调试文件
                    if self.save_intermediate_files:
                        print(f"📁 正在保存step2中间文件到: {self.debug_output_dir}")
                        self._save_intermediate_file(skipped_output, "step2_pandoc", "跳过转换")

                    return skipped_output

                try:
                    self.pandoc_converter = PandocConverter(pandoc_path)
                except Exception as init_error:
                    print(f"⚠️ Pandoc初始化失败: {init_error}")
                    print("跳过Pandoc转换步骤")

                    # 生成一个标记后的文件，表示跳过了Pandoc转换
                    base_name = os.path.splitext(os.path.basename(content_file))[0]
                    if not self.temp_dir:
                        raise ValueError("临时目录未初始化")
                    init_failed_output = os.path.join(self.temp_dir, f"{base_name}_Pandoc初始化失败.docx")
                    shutil.copy2(content_file, init_failed_output)

                    # 保存中间文件
                    self.intermediate_files['pandoc_converted'] = init_failed_output

                    # 保存调试文件
                    if self.save_intermediate_files:
                        print(f"📁 正在保存step2中间文件到: {self.debug_output_dir}")
                        self._save_intermediate_file(init_failed_output, "step2_pandoc", "初始化失败")

                    return init_failed_output

            # 生成Pandoc输出文件路径
            base_name = os.path.splitext(os.path.basename(content_file))[0]
            if not self.temp_dir:
                raise ValueError("临时目录未初始化")
            pandoc_output = os.path.join(self.temp_dir, f"{base_name}_pandoc转换.docx")

            if self.save_intermediate_files:
                print(f"📄 正在使用模板转换: {os.path.basename(template_file)}")
                print(f"📤 输出文件: {os.path.basename(pandoc_output)}")

            # 使用模板进行转换，保持表格结构
            success = self.pandoc_converter.convert_with_template(
                input_file=content_file,
                output_file=pandoc_output,
                template_file=template_file,
                additional_args=[
                    "--preserve-tabs",  # 保持制表符
                    "--wrap=none",  # 不自动换行
                    "--reference-links",  # 使用引用链接
                    "--columns=80",  # 设置合适的列宽
                    "--table-of-contents",  # 保持目录结构
                    "--standalone",  # 独立文档模式
                ]
            )

            if success and os.path.exists(pandoc_output):
                print("✅ Pandoc转换成功!")
                # 保存中间文件
                self.intermediate_files['pandoc_converted'] = pandoc_output

                # 保存调试文件
                if self.save_intermediate_files:
                    print(f"   转换后文件: {os.path.basename(pandoc_output)}")
                    print(f"📁 正在保存step2中间文件到: {self.debug_output_dir}")
                    self._save_intermediate_file(pandoc_output, "step2_pandoc", "转换成功")

                return pandoc_output
            else:
                print("❌ Pandoc转换失败，使用原始文件继续")

                # 在转换失败时，复制原文件作为备用
                base_name = os.path.splitext(os.path.basename(content_file))[0]
                fallback_output = os.path.join(self.temp_dir, f"{base_name}_Pandoc失败备用.docx")
                shutil.copy2(content_file, fallback_output)

                # 保存中间文件
                self.intermediate_files['pandoc_converted'] = fallback_output

                # 保存调试文件
                if self.save_intermediate_files:
                    print(f"📁 正在保存step2中间文件到: {self.debug_output_dir}")
                    self._save_intermediate_file(fallback_output, "step2_pandoc", "失败备用")

                return fallback_output

        except Exception as e:
            print(f"❌ Pandoc转换过程中发生错误: {str(e)}")
            print("使用原始文件继续后续处理")

            # 在发生异常时，复制原文件作为备用
            base_name = os.path.splitext(os.path.basename(content_file))[0]
            if not self.temp_dir:
                raise ValueError("临时目录未初始化")
            error_output = os.path.join(self.temp_dir, f"{base_name}_Pandoc异常备用.docx")
            shutil.copy2(content_file, error_output)

            # 保存中间文件
            self.intermediate_files['pandoc_converted'] = error_output

            # 保存调试文件
            if self.save_intermediate_files:
                print(f"📁 正在保存step2中间文件到: {self.debug_output_dir}")
                self._save_intermediate_file(error_output, "step2_pandoc", "异常备用")

            return error_output

    def step3_format_tables(self, content_file: str, template_file: str, original_content_file: Optional[str] = None) -> \
    Optional[str]:
        """
        步骤3: 表格格式化 - 使用原始内容文件中的表格替换处理后的文件中的表格

        Args:
            content_file: 正文内容文件路径（被替换表格内容）
            template_file: 模板文件路径
            original_content_file: 原始正文内容文件路径（提供表格内容），如果提供则使用表格替换功能

        Returns:
            str: 表格格式化后的文件路径
        """
        print("-" * 50)
        print("📊 步骤3: 表格格式化")

        try:
            # 确保临时目录存在
            if not self.temp_dir:
                raise ValueError("临时目录未初始化")

            # 生成表格格式化输出文件路径
            base_name = os.path.splitext(os.path.basename(content_file))[0]

            # 如果提供了原始内容文件，则使用表格替换功能
            if original_content_file and os.path.exists(original_content_file):
                formatted_output = os.path.join(self.temp_dir, f"{base_name}_表格替换.docx")
                print("🔧 使用表格替换功能处理表格...")

                # 使用docx_table_replace.py中的方法进行表格替换
                try:
                    from utils.docx_table_replace import replace_tables_by_index

                    # 使用原始正文内容文件中的表格替换Pandoc转换后的文件中的表格
                    success = replace_tables_by_index(
                        original_path=original_content_file,  # 原始正文内容文件（提供表格）
                        edited_path=content_file,  # Pandoc转换后的文件（被替换表格）
                        output_path=formatted_output  # 输出文件
                    )

                    if success and os.path.exists(formatted_output):
                        print("✅ 表格替换成功!")

                        # 保存中间文件
                        self.intermediate_files['table_replaced'] = formatted_output

                        # 保存调试文件
                        if self.save_intermediate_files:
                            print(f"   表格替换后文件: {os.path.basename(formatted_output)}")
                            print(f"📁 正在保存step3中间文件到: {self.debug_output_dir}")
                            self._save_intermediate_file(formatted_output, "step3_table", "替换完成")

                        return formatted_output
                    else:
                        print("❌ 表格替换失败，使用原始文件继续")
                        return content_file

                except Exception as replace_error:
                    print(f"❌ 表格替换过程中发生错误: {str(replace_error)}")
                    print("使用原始文件继续后续处理")
                    return content_file
            else:
                print("⚠️ 未提供原始内容文件，跳过表格替换步骤")
                return content_file

        except Exception as e:
            print(f"❌ 表格格式化过程中发生错误: {str(e)}")
            return content_file

    def step4_merge_documents(self, cover_toc_file: str, processed_content_file: str, output_file: str) -> bool:
        """
        步骤4: 文档合并 - 将封面添加到正文开始

        Args:
            cover_toc_file: 封面文件路径
            processed_content_file: 处理后的正文文件路径
            output_file: 最终输出文件路径

        Returns:
            bool: 合并是否成功
        """
        print("-" * 50)
        print("📚 步骤4: 文档合并")

        try:
            if self.save_intermediate_files:
                print(f"📄 封面: {os.path.basename(cover_toc_file)}")
                print(f"📄 正文内容: {os.path.basename(processed_content_file)}")
                print(f"📤 最终输出: {os.path.basename(output_file)}")

            # 确保输出目录存在
            output_dir = os.path.dirname(output_file)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir, exist_ok=True)
                print(f"📁 创建输出目录: {output_dir}")

            # 使用docx_merge.py中的方法进行文档合并
            try:
                copy_all_to_beginning(
                    file_a=cover_toc_file,
                    file_b=processed_content_file,
                    output_file=output_file
                )

                if os.path.exists(output_file):
                    print("✅ 文档合并成功!")

                    # 验证输出文件的有效性
                    if self._validate_output_file(output_file):
                        print("✅ 输出文件验证通过")
                        # 保存最终文件
                        self.intermediate_files['final_document'] = output_file

                        # 保存调试文件
                        if self.save_intermediate_files:
                            print(f"   最终文档: {os.path.basename(output_file)}")
                            print(f"📁 正在保存step4最终文件到: {self.debug_output_dir}")
                            self._save_intermediate_file(output_file, "step4_final", "最终文档")
                        return True
                    else:
                        print("❌ 输出文件验证失败")
                        return False

                else:
                    print("❌ 文档合并失败，输出文件不存在")
                    return False

            except Exception as merge_error:
                print(f"❌ 使用docx_merge方法合并文档时发生错误: {str(merge_error)}")
                return False

        except Exception as e:
            print(f"❌ 文档合并过程中发生错误: {str(e)}")
            return False

    def step5_update_toc_title(self, docx_file: str, new_title: str = "目录") -> bool:
        """
        步骤5: 更新目录标题 - 修改文档中的目录标题

        Args:
            docx_file: 需要修改的文档路径
            new_title: 新的目录标题

        Returns:
            bool: 更新是否成功
        """
        print("-" * 50)
        print("🏷️ 步骤5: 更新目录标题")

        try:
            if self.save_intermediate_files:
                print(f"📄 目标文档: {os.path.basename(docx_file)}")
                print(f"🔤 新标题: '{new_title}'")

            # 检查文件是否存在
            if not os.path.exists(docx_file):
                print(f"❌ 文件不存在: {docx_file}")
                return False

            # 使用docx_update_toc_title中的方法更新目录标题
            try:
                # 使用XML方式更新目录标题，保留原有格式
                update_toc_title_xml(docx_file, new_title)
                print("✅ 目录标题更新成功!")
                return True
            except Exception as xml_error:
                print(f"⚠️ XML方式更新目录标题失败: {str(xml_error)}")
                print("尝试使用COM方式更新目录标题...")

                try:
                    # 使用COM方式更新目录标题
                    from utils.docx_update_toc_title import update_toc_title
                    update_toc_title(docx_file, new_title)
                    print("✅ 目录标题更新成功!")
                    return True
                except Exception as com_error:
                    print(f"❌ COM方式更新目录标题也失败: {str(com_error)}")
                    return False

        except Exception as e:
            print(f"❌ 更新目录标题过程中发生错误: {str(e)}")
            return False

    def step6_format_pictures(self, docx_file: str) -> bool:
        """
        步骤6: 图片格式化 - 图片居中，单倍行距

        Args:
            docx_file: 需要处理的文档路径

        Returns:
            bool: 处理是否成功
        """
        print("-" * 50)
        print("🖼️ 步骤6: 图片格式化")

        try:
            if self.save_intermediate_files:
                print(f"📄 目标文档: {os.path.basename(docx_file)}")

            # 检查文件是否存在
            if not os.path.exists(docx_file):
                print(f"❌ 文件不存在: {docx_file}")
                return False

            # 生成处理后的文件路径
            base_name = os.path.splitext(os.path.basename(docx_file))[0]
            if not self.temp_dir:
                raise ValueError("临时目录未初始化")
            formatted_output = os.path.join(self.temp_dir, f"{base_name}_图片格式化.docx")

            # 使用docx_picture.py中的高级处理方式
            try:
                from utils.docx_picture import format_pictures_with_advanced_settings

                # 调用高级图片格式化函数
                success = format_pictures_with_advanced_settings(
                    doc_path=docx_file,
                    save_path=formatted_output
                )

                if success and os.path.exists(formatted_output):
                    print("✅ 图片格式化成功!")

                    # 保存中间文件
                    self.intermediate_files['picture_formatted'] = formatted_output

                    # 保存调试文件
                    if self.save_intermediate_files:
                        print(f"   图片格式化后文件: {os.path.basename(formatted_output)}")
                        print(f"📁 正在保存step6中间文件到: {self.debug_output_dir}")
                        self._save_intermediate_file(formatted_output, "step6_picture", "格式化完成")

                    # 将处理后的文件复制回原文件路径，以便后续步骤使用
                    shutil.copy2(formatted_output, docx_file)
                    return True
                else:
                    print("❌ 图片格式化失败")
                    return False

            except Exception as format_error:
                print(f"❌ 图片格式化过程中发生错误: {str(format_error)}")
                import traceback
                traceback.print_exc()
                return False

        except Exception as e:
            print(f"❌ 图片格式化过程中发生错误: {str(e)}")
            return False

    def step7_format_library_number(self, docx_file: str) -> bool:
        """
        步骤7: 库号信息格式化 - 将库号右靠齐

        Args:
            docx_file: 需要处理的文档路径

        Returns:
            bool: 处理是否成功
        """
        print("-" * 50)
        print("🔢 步骤7: 库号信息格式化")

        try:
            if self.save_intermediate_files:
                print(f"📄 目标文档: {os.path.basename(docx_file)}")

            # 检查文件是否存在
            if not os.path.exists(docx_file):
                print(f"❌ 文件不存在: {docx_file}")
                return False

            # 生成处理后的文件路径
            base_name = os.path.splitext(os.path.basename(docx_file))[0]
            if not self.temp_dir:
                raise ValueError("临时目录未初始化")
            formatted_output = os.path.join(self.temp_dir, f"{base_name}_库号格式化.docx")

            # 使用docx_supplement.py中的高级处理方式
            try:
                from utils.docx_supplement import format_library_number_advanced

                # 调用高级库号格式化函数
                success = format_library_number_advanced(
                    doc_path=docx_file,
                    save_path=formatted_output
                )

                if success and os.path.exists(formatted_output):
                    print("✅ 库号信息格式化成功!")

                    # 保存中间文件
                    self.intermediate_files['library_number_formatted'] = formatted_output

                    # 保存调试文件
                    if self.save_intermediate_files:
                        print(f"   库号格式化后文件: {os.path.basename(formatted_output)}")
                        print(f"📁 正在保存step7中间文件到: {self.debug_output_dir}")
                        self._save_intermediate_file(formatted_output, "step7_library_number", "格式化完成")

                    # 将处理后的文件复制回原文件路径，以便后续步骤使用
                    shutil.copy2(formatted_output, docx_file)
                    return True
                else:
                    print("❌ 库号信息格式化失败")
                    return False

            except Exception as format_error:
                print(f"❌ 库号信息格式化过程中发生错误: {str(format_error)}")
                import traceback
                traceback.print_exc()
                return False

        except Exception as e:
            print(f"❌ 库号信息格式化过程中发生错误: {str(e)}")
            return False

    def step8_insert_section_break(self, docx_file: str) -> bool:
        """
        步骤8: 在目录后插入分节符

        Args:
            docx_file: 需要处理的文档路径

        Returns:
            bool: 处理是否成功
        """
        print("-" * 50)
        print("📑 步骤8: 在目录后插入分节符")

        try:
            if self.save_intermediate_files:
                print(f"📄 目标文档: {os.path.basename(docx_file)}")

            # 检查文件是否存在
            if not os.path.exists(docx_file):
                print(f"❌ 文件不存在: {docx_file}")
                return False

            # 添加延迟，确保之前的COM操作完全释放资源
            print("⏳ 等待COM资源释放...")
            import time
            time.sleep(3)

            # 生成处理后的文件路径
            base_name = os.path.splitext(os.path.basename(docx_file))[0]
            if not self.temp_dir:
                raise ValueError("临时目录未初始化")
            formatted_output = os.path.join(self.temp_dir, f"{base_name}_插入分节符.docx")

            # 使用docx_supplement.py中的方法
            try:
                from utils.docx_supplement import insert_section_break_after_toc

                # 调用插入分节符函数
                success = insert_section_break_after_toc(
                    doc_path=docx_file,
                    save_path=formatted_output
                )

                if success and os.path.exists(formatted_output):
                    print("✅ 在目录后插入分节符成功!")

                    # 保存中间文件
                    self.intermediate_files['section_break_inserted'] = formatted_output

                    # 保存调试文件
                    if self.save_intermediate_files:
                        print(f"   插入分节符后文件: {os.path.basename(formatted_output)}")
                        print(f"📁 正在保存step8中间文件到: {self.debug_output_dir}")
                        self._save_intermediate_file(formatted_output, "step8_section_break", "插入完成")

                    # 将处理后的文件复制回原文件路径，以便后续步骤使用
                    shutil.copy2(formatted_output, docx_file)
                    return True
                else:
                    print("❌ 在目录后插入分节符失败")
                    # 尝试使用备选方法
                    print("🔄 尝试使用备选方法...")
                    return self._insert_section_break_fallback(docx_file, formatted_output)

            except Exception as format_error:
                print(f"❌ 在目录后插入分节符过程中发生错误: {str(format_error)}")
                # 尝试使用备选方法
                print("🔄 尝试使用备选方法...")
                return self._insert_section_break_fallback(docx_file, formatted_output)

        except Exception as e:
            print(f"❌ 在目录后插入分节符过程中发生错误: {str(e)}")
            return False

    def step9_process_sections(self, docx_file: str) -> bool:
        """
        步骤9: 处理文档节的页码设置
        - 取消第三节与第二节的链接
        - 处理第二节的页码（删除PAGE域）
        - 处理第三节的页码（重置为1）

        Args:
            docx_file: 需要处理的文档路径

        Returns:
            bool: 处理是否成功
        """
        print("-" * 50)
        print("📑 步骤9: 处理文档节的页码设置")

        try:
            if self.save_intermediate_files:
                print(f"📄 目标文档: {os.path.basename(docx_file)}")

            # 检查文件是否存在
            if not os.path.exists(docx_file):
                print(f"❌ 文件不存在: {docx_file}")
                return False

            # 添加延迟，确保之前的COM操作完全释放资源
            print("⏳ 等待COM资源释放...")
            import time
            time.sleep(3)

            # 生成临时文件路径
            base_name = os.path.splitext(os.path.basename(docx_file))[0]
            if not self.temp_dir:
                raise ValueError("临时目录未初始化")

            temp_file1 = os.path.join(self.temp_dir, f"{base_name}_取消节链接.docx")
            temp_file2 = os.path.join(self.temp_dir, f"{base_name}_处理第二节页码.docx")
            formatted_output = os.path.join(self.temp_dir, f"{base_name}_处理第三节页码.docx")

            # 步骤9.1: 使用docx_supplement.py中的方法取消第三节与第二节的链接
            print("\n步骤9.1: 取消第三节与第二节的链接...")
            try:
                from utils.docx_supplement import cancel_section_link_com

                success = cancel_section_link_com(
                    doc_path=docx_file,
                    save_path=temp_file1,
                    section_number=3  # 第三节
                )

                if not success or not os.path.exists(temp_file1):
                    print("❌ 步骤9.1失败，无法继续执行后续步骤")
                    return False
                else:
                    print("✅ 步骤9.1完成")

                    # 保存调试文件
                    if self.save_intermediate_files:
                        print(f"   取消节链接后文件: {os.path.basename(temp_file1)}")
                        print(f"📁 正在保存step9.1中间文件到: {self.debug_output_dir}")
                        self._save_intermediate_file(temp_file1, "step9_section_link", "取消链接完成")
            except Exception as e:
                print(f"❌ 步骤9.1失败: {e}")
                return False

            # 步骤9.2: 使用docx_supplement.py中的方法处理第二节的页码
            print("\n步骤9.2: 处理第二节的页码...")
            try:
                from utils.docx_supplement import process_section2_docx

                process_section2_docx(temp_file1, temp_file2, section_index=2)
                print("✅ 步骤9.2完成")

                # 保存调试文件
                if self.save_intermediate_files:
                    print(f"   处理第二节页码后文件: {os.path.basename(temp_file2)}")
                    print(f"📁 正在保存step9.2中间文件到: {self.debug_output_dir}")
                    self._save_intermediate_file(temp_file2, "step9_section2_page", "处理完成")
            except Exception as e:
                print(f"❌ 步骤9.2失败: {e}")
                return False

            # 步骤9.3: 使用docx_supplement.py中的方法处理第三节的页码
            print("\n步骤9.3: 处理第三节的页码...")
            try:
                from utils.docx_supplement import process_section3_docx

                process_section3_docx(temp_file2, formatted_output)
                print("✅ 步骤9.3完成")

                # 保存中间文件
                self.intermediate_files['section_page_processed'] = formatted_output

                # 保存调试文件
                if self.save_intermediate_files:
                    print(f"   处理第三节页码后文件: {os.path.basename(formatted_output)}")
                    print(f"📁 正在保存step9.3中间文件到: {self.debug_output_dir}")
                    self._save_intermediate_file(formatted_output, "step9_section3_page", "处理完成")

                # 将处理后的文件复制回原文件路径，以便后续步骤使用
                shutil.copy2(formatted_output, docx_file)
                return True
            except Exception as e:
                print(f"❌ 步骤9.3失败: {e}")
                return False

        except Exception as e:
            print(f"❌ 处理文档节的页码设置过程中发生错误: {str(e)}")
            return False

    def step11_format_library_number_advanced(self, docx_file: str) -> bool:
        """
        步骤11: 高级库号信息格式化 - 将"库号："设置为三号黑体加粗，数字字母设置为三号Times New Roman加粗

        Args:
            docx_file: 需要处理的文档路径

        Returns:
            bool: 处理是否成功
        """
        print("-" * 50)
        print("🔢 步骤11: 高级库号信息格式化")

        try:
            if self.save_intermediate_files:
                print(f"📄 目标文档: {os.path.basename(docx_file)}")

            # 检查文件是否存在
            if not os.path.exists(docx_file):
                print(f"❌ 文件不存在: {docx_file}")
                return False

            # 生成处理后的文件路径
            base_name = os.path.splitext(os.path.basename(docx_file))[0]
            if not self.temp_dir:
                raise ValueError("临时目录未初始化")
            formatted_output = os.path.join(self.temp_dir, f"{base_name}_库号高级格式化.docx")

            # 使用cover_number_format.py中的处理方式
            try:
                from utils.cover_number_format import process_document

                # 调用库号格式化函数
                process_document(docx_file, formatted_output)

                if os.path.exists(formatted_output):
                    print("✅ 库号信息高级格式化成功!")

                    # 保存中间文件
                    self.intermediate_files['library_number_advanced_formatted'] = formatted_output

                    # 保存调试文件
                    if self.save_intermediate_files:
                        print(f"   库号高级格式化后文件: {os.path.basename(formatted_output)}")
                        print(f"📁 正在保存step11中间文件到: {self.debug_output_dir}")
                        self._save_intermediate_file(formatted_output, "step11_library_number", "高级格式化完成")

                    # 将处理后的文件复制回原文件路径，以便后续步骤使用
                    shutil.copy2(formatted_output, docx_file)
                    return True
                else:
                    print("❌ 库号信息高级格式化失败")
                    return False

            except Exception as format_error:
                print(f"❌ 库号信息高级格式化过程中发生错误: {str(format_error)}")
                import traceback
                traceback.print_exc()
                return False

        except Exception as e:
            print(f"❌ 库号信息高级格式化过程中发生错误: {str(e)}")
            return False

    def step12_post_process_headers_footers(self, docx_path: str) -> bool:
        """
        步骤12: 后处理页眉页脚 - 在所有转换步骤完成后处理页眉页脚

        Args:
            docx_path: 文档路径

        Returns:
            bool: 处理是否成功
        """
        print("-" * 50)
        print("📑 步骤13: 后处理页眉页脚")

        try:
            if self.save_intermediate_files:
                print(f"📄 目标文档: {os.path.basename(docx_path)}")

            # 检查文件是否存在
            if not os.path.exists(docx_path):
                print(f"❌ 文件不存在: {docx_path}")
                return False

            # 生成处理后的文件路径
            base_name = os.path.splitext(os.path.basename(docx_path))[0]
            if not self.temp_dir:
                raise ValueError("临时目录未初始化")
            formatted_output = os.path.join(self.temp_dir, f"{base_name}_step13_headers_footers.docx")

            # 确定页脚文本内容
            if self.extracted_footer_content and self.extracted_footer_content.strip():
                footer_text = self.extracted_footer_content.strip()
                print(f"📄 使用提取的页脚内容: {footer_text}")
            else:
                footer_text = "PAGE \\* MERGEFORMAT / NUMPAGES \\* MERGEFORMAT"
                print(f"📄 使用默认页脚内容: {footer_text}")

            from docx import Document
            from docx.shared import Pt, Cm

            # 执行页眉页脚处理逻辑
            doc = Document(docx_path)

            # 遍历文档的所有sections
            for i, section in enumerate(doc.sections):
                # 获取当前节的页眉
                header = section.header

                # 检查页眉是否包含文本
                header_contains_text = any(paragraph.text.strip() for paragraph in header.paragraphs)

                if header_contains_text:
                    # 如果当前节的页眉包含文本，则为该节及后续节添加页脚
                    footer = section.footer

                    # 遍历页脚段落并查找是否已经有页码
                    existing_page_numbers = []
                    for paragraph in footer.paragraphs:
                        for run in paragraph.runs:
                            if "PAGE" in run.text or "NUMPAGES" in run.text:
                                existing_page_numbers.append(run.text)

                    # 删除多余的制表符等符号
                    if footer.paragraphs:
                        para = footer.paragraphs[0]
                        tabs_deleted = 0  # 已删除的制表符数量
                        runs = para.runs  # 段落的文本片段列表

                        # 遍历所有run，从第一个开始处理
                        for run in runs:
                            run_text = run.text  # 当前run的文本

                            if tabs_deleted >= 3:
                                break  # 已删除三个，停止处理

                            # 查找当前run中开头的制表符
                            if not run_text:
                                continue  # 空run跳过

                            # 统计当前run开头的制表符数量
                            start_tabs = 0
                            while start_tabs < len(run_text) and (run_text[start_tabs] == '\t'):
                                start_tabs += 1

                            if start_tabs == 0:
                                continue  # 该run开头没有制表符，跳过

                            # 计算需要从当前run删除的制表符数量
                            need_delete = min(start_tabs, 3 - tabs_deleted)

                            # 删除开头的need_delete个制表符
                            run.text = run_text[need_delete:]  # 截断文本，保留剩余内容
                            tabs_deleted += need_delete

                        # 在footer.paragraphs[0]中添加footer_text
                        para.add_run(footer_text)

                        # 设置字号五号字
                        for run in para.runs:
                            run.font.size = Pt(10.5)

            # 保存修改后的文档
            doc.save(formatted_output)
            print("✅ 页眉页脚处理成功!")

            # 保存中间文件
            self.intermediate_files['headers_footers_processed'] = formatted_output

            # 保存调试文件
            if self.save_intermediate_files:
                print(f"   页眉页脚处理后文件: {os.path.basename(formatted_output)}")
                print(f"📁 正在保存step12表格中间文件到: {self.debug_output_dir}")
                self._save_intermediate_file(formatted_output, "step13_headers_footers", "处理完成")

            # 将处理后的文件复制回原文件路径，以便后续步骤使用
            shutil.copy2(formatted_output, docx_path)
            return True

        except Exception as e:
            print(f"❌ 页眉页脚处理过程中发生错误: {str(e)}")
            return False

    def step13_word_table_replace(self, edited_docx: str, original_docx: str) -> bool:
        """Step12: Table replacement using the Word COM helper."""
        try:
            from utils.word_table_replace import replace_tables
        except ImportError as import_error:
            print(f"[warn] Unable to import word_table_replace: {import_error}")
            return False

        print("-" * 50)
        print("步骤12: 表格替换")

        if not original_docx:
            print("[warn] Missing step0 output file, skip table replacement")
            return False

        try:
            if not os.path.exists(edited_docx):
                print(f"[warn] Edited document not found: {edited_docx}")
                return False
            if not os.path.exists(original_docx):
                print(f"[warn] Original document not found: {original_docx}")
                return False

            if not self.temp_dir:
                raise ValueError("临时目录未初始化")

            cached_sections = None
            try:
                from docx import Document
                pre_doc = Document(edited_docx)
                cached_sections = []
                for section in pre_doc.sections:
                    section_snapshot = {}
                    ps = getattr(section, "page_setup", None)
                    if ps is not None:
                        section_snapshot = {
                            "footer_distance": getattr(ps, "footer_distance", None),
                            "header_distance": getattr(ps, "header_distance", None),
                            "top_margin": getattr(ps, "top_margin", None),
                            "bottom_margin": getattr(ps, "bottom_margin", None),
                            "left_margin": getattr(ps, "left_margin", None),
                            "right_margin": getattr(ps, "right_margin", None),
                            "gutter": getattr(ps, "gutter", None),
                        }
                    else:
                        section_snapshot = {
                            "footer_distance": getattr(section, "footer_distance", None),
                            "header_distance": getattr(section, "header_distance", None),
                            "top_margin": getattr(section, "top_margin", None),
                            "bottom_margin": getattr(section, "bottom_margin", None),
                            "left_margin": getattr(section, "left_margin", None),
                            "right_margin": getattr(section, "right_margin", None),
                            "gutter": getattr(section, "gutter", None),
                        }
                    cached_sections.append(section_snapshot)
            except Exception as snapshot_err:
                print(f"[warn] Unable to snapshot section settings: {snapshot_err}")
                cached_sections = None

            base_name = os.path.splitext(os.path.basename(edited_docx))[0]
            replaced_output = os.path.join(self.temp_dir, f"{base_name}_step12_table_replace.docx")

            if self.save_intermediate_files:
                print(f"源文件(步骤0): {os.path.basename(original_docx)}")
                print(f"当前待替换文件: {os.path.basename(edited_docx)}")
                print(f"表格替换输出: {os.path.basename(replaced_output)}")

            _kill_existing_winword_processes()
            replace_tables(src_path=original_docx, dst_path=edited_docx, out_path=replaced_output)

            if os.path.exists(replaced_output):
                if cached_sections is not None:
                    try:
                        from docx import Document
                        _wait_for_file_release_for_step(replaced_output)
                        post_doc = Document(replaced_output)
                        for section, saved in zip(post_doc.sections, cached_sections):
                            ps = getattr(section, "page_setup", None)
                            if ps is not None:
                                if saved.get("footer_distance") is not None:
                                    ps.footer_distance = saved["footer_distance"]
                                if saved.get("header_distance") is not None:
                                    ps.header_distance = saved["header_distance"]
                                if saved.get("top_margin") is not None:
                                    ps.top_margin = saved["top_margin"]
                                if saved.get("bottom_margin") is not None:
                                    ps.bottom_margin = saved["bottom_margin"]
                                if saved.get("left_margin") is not None:
                                    ps.left_margin = saved["left_margin"]
                                if saved.get("right_margin") is not None:
                                    ps.right_margin = saved["right_margin"]
                                if saved.get("gutter") is not None:
                                    ps.gutter = saved["gutter"]
                            else:
                                for key, value in saved.items():
                                    if value is None:
                                        continue
                                    if hasattr(section, key):
                                        setattr(section, key, value)
                        post_doc.save(replaced_output)
                    except Exception as restore_err:
                        print(f"[warn] Unable to restore section settings: {restore_err}")

                _wait_for_file_release_for_step(edited_docx)
                _wait_for_file_release_for_step(replaced_output)
                shutil.copy2(replaced_output, edited_docx)
                self.intermediate_files['table_replaced_before_headers'] = replaced_output

                if self.save_intermediate_files:
                    print(f"   表格替换后文件: {os.path.basename(replaced_output)}")
                    print(f"📁 正在保存step12表格中间文件到: {self.debug_output_dir}")
                    self._save_intermediate_file(replaced_output, "step12_table", "表格替换完成")

                print("表格替换完成")
                return True

            print("表格替换失败，未生成输出文件")
            return False

        except Exception as exc:
            print(f"表格替换过程中发生错误: {exc}")
            import traceback
            traceback.print_exc()
            return False




    def _insert_section_break_fallback(self, docx_file: str, formatted_output: str) -> bool:
        """
        备选方法：在目录后插入分节符的降级处理

        Args:
            docx_file: 需要处理的文档路径
            formatted_output: 输出文件路径

        Returns:
            bool: 处理是否成功
        """
        try:
            print("🔧 尝试使用XML方法插入分节符...")

            # 使用XML方法作为备选
            from utils.docx_section_break import insert_section_break_after_toc_xml

            success = insert_section_break_after_toc_xml(
                doc_path=docx_file,
                save_path=formatted_output
            )

            if success and os.path.exists(formatted_output):
                print("✅ 使用XML方法插入分节符成功!")
                return True
            else:
                print("❌ 使用XML方法插入分节符失败")
                return False

        except Exception as xml_error:
            print(f"❌ 使用XML方法插入分节符时发生错误: {str(xml_error)}")
            print("⚠️ 无法在目录后插入分节符，继续使用原有格式")
            return False

    def _validate_output_file(self, file_path: str) -> bool:
        """
        验证输出文件的有效性

        Args:
            file_path: 文件路径

        Returns:
            bool: 文件是否有效
        """
        try:
            if not os.path.exists(file_path):
                return False

            # 检查文件大小
            file_size = os.path.getsize(file_path)
            if file_size < 1000:  # 小于1KB可能有问题
                print(f"⚠️ 文件大小异常: {file_size} bytes")
                return False

            # 检查是否为有效的ZIP文件（DOCX本质上ZIP文件）
            import zipfile
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_file:
                    # 检查必要的文件
                    required_files = ['[Content_Types].xml', '_rels/.rels', 'word/document.xml']
                    file_list = zip_file.namelist()

                    for req_file in required_files:
                        if req_file not in file_list:
                            print(f"⚠️ 缺少必要文件: {req_file}")
                            return False

                    return True
            except zipfile.BadZipFile:
                print("⚠️ 文件不是有效的ZIP格式")
                return False

        except Exception as e:
            print(f"⚠️ 文件验证过程中发生错误: {str(e)}")
            return False

    def step10_remove_highlights(self, source_file: str) -> str:
        """
        步骤10: 删除文档中所有的突出显示（高亮、底纹、颜色）

        Args:
            source_file: 源文档路径

        Returns:
            str: 去除突出显示后的文档路径
        """
        print("-" * 50)
        print("📑 步骤10: 删除文档中所有的突出显示")

        try:
            # 确保临时目录存在
            if not self.temp_dir:
                raise ValueError("临时目录未初始化")

            # 检查lxml是否可用
            if not LXML_AVAILABLE:
                print("⚠️ lxml库不可用，跳过突出显示删除步骤")
                return source_file

            # 生成输出文件路径
            base_name = os.path.splitext(os.path.basename(source_file))[0]
            no_highlight_file = os.path.join(self.temp_dir, f"{base_name}_无突出显示.docx")

            print(f"📄 输入文档: {os.path.basename(source_file)}")
            print(f"📤 输出文档: {os.path.basename(no_highlight_file)}")

            # 使用lxml删除Word文件中的所有高亮
            try:
                # 定义命名空间和常量
                W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                NSMAP = {"w": W_NS}
                XML_PARSER = etree.XMLParser(ns_clean=True, recover=True, remove_blank_text=False) if (
                            LXML_AVAILABLE and etree is not None) else None
                REMOVE_COLOR_NODE = True  # 彻底移除颜色节点

                def process_xml_bytes(data: bytes, remove_color_node=False) -> bytes:
                    """删除 w:highlight、w:shd，并处理 w:color"""
                    if not LXML_AVAILABLE or etree is None:
                        return data

                    try:
                        root = etree.fromstring(data, parser=XML_PARSER) if etree is not None else None
                        if root is None:
                            return data
                    except Exception:
                        return data

                    changed = False
                    # 删除 highlight
                    for node in list(root.iter(qn('w:highlight'))):
                        parent = node.getparent()
                        if parent is not None:
                            parent.remove(node)
                            changed = True

                    # 删除底纹
                    for node in list(root.iter(qn('w:shd'))):
                        parent = node.getparent()
                        if parent is not None:
                            parent.remove(node)
                            changed = True

                    # 处理颜色
                    for color in list(root.iter(qn('w:color'))):
                        val = color.get("val")
                        if val is not None and val.lower() != "auto":
                            if remove_color_node:
                                parent = color.getparent()
                                if parent is not None:
                                    parent.remove(color)
                                    changed = True
                            else:
                                color.set("val", "auto")
                                changed = True

                    if not changed:
                        return data
                    return etree.tostring(root, encoding="utf-8", xml_declaration=True) if (
                                LXML_AVAILABLE and etree is not None) else data

                # 处理DOCX文件
                src = Path(source_file)
                if not src.exists():
                    raise FileNotFoundError(f"文件不存在: {src}")

                dest = Path(no_highlight_file)

                with zipfile.ZipFile(src, 'r') as zin:
                    with zipfile.ZipFile(dest, 'w', compression=zipfile.ZIP_DEFLATED) as zout:
                        for name in zin.namelist():
                            data = zin.read(name)
                            if name.startswith("word/") and name.endswith(".xml"):
                                try:
                                    new_data = process_xml_bytes(data, remove_color_node=REMOVE_COLOR_NODE)
                                    zout.writestr(name, new_data)
                                except Exception as e:
                                    print(f"⚠ 处理 {name} 出错，保留原文件。错误：{e}")
                                    zout.writestr(name, data)
                            else:
                                zout.writestr(name, data)

                print("✅ 突出显示删除成功!")

                # 保存中间文件到指定目录便于查看调试
                if self.save_intermediate_files:
                    print(f"   无突出显示文档: {os.path.basename(no_highlight_file)}")
                    print(f"📁 正在保存step10中间文件到: {self.debug_output_dir}")
                    self._save_intermediate_file(no_highlight_file, "step10_highlights", "无突出显示")

                # 保存中间文件路径
                self.intermediate_files['no_highlights'] = no_highlight_file

                return no_highlight_file

            except Exception as remove_error:
                print(f"❌ 突出显示删除过程中发生错误: {str(remove_error)}")
                print("   继续使用原始文档")
                return source_file

        except Exception as e:
            print(f"❌ 突出显示删除过程中发生错误: {str(e)}")
            return source_file

    def _find_pandoc_executable(self) -> Optional[str]:
        """
        查找Pandoc可执行文件

        Returns:
            str: Pandoc可执行文件路径
        """
        import subprocess

        # 可能的Pandoc位置 - 优先utils目录
        current_dir = os.path.dirname(os.path.abspath(__file__))
        parent_dir = os.path.dirname(current_dir)
        utils_pandoc = os.path.join(parent_dir, 'utils', 'pandoc.exe')

        possible_paths = [
            # 优先使用utils目录中的pandoc.exe
            utils_pandoc,
            # 系统PATH中的pandoc
            "pandoc",
            "pandoc.exe",
            # 当前目录下的pandoc.exe
            os.path.join(os.path.dirname(__file__), "pandoc.exe"),
            os.path.join(os.path.dirname(os.path.dirname(__file__)), "pandoc.exe"),
            # 常见安装位置
            r"C:\Program Files\Pandoc\pandoc.exe",
            r"C:\Program Files (x86)\Pandoc\pandoc.exe",
            # Conda环境
            os.path.join(os.environ.get('CONDA_PREFIX', ''), 'Scripts', 'pandoc.exe'),
            os.path.join(os.environ.get('CONDA_PREFIX', ''), 'bin', 'pandoc')
        ]

        for path in possible_paths:
            if not path:  # 跳过空路径
                continue

            try:
                # 尝试执行pandoc --version
                result = subprocess.run(
                    [path, "--version"],
                    capture_output=True,
                    text=True,
                    timeout=10,
                    creationflags=subprocess.CREATE_NO_WINDOW if os.name == 'nt' else 0
                )
                if result.returncode == 0:
                    return path
            except (subprocess.TimeoutExpired, FileNotFoundError, PermissionError, Exception):
                continue

        print("⚠️ 未找到Pandoc可执行文件")
        print("请通过以下方式安装Pandoc:")
        print("1. 下载安装: https://pandoc.org/installing.html")
        print("2. 使用conda: conda install pandoc")
        print("3. 使用choco: choco install pandoc")
        print(f"4. utils目录路径: {utils_pandoc}")
        return None

    def convert_document(
            self,
            source_file: str,
            output_file: str,
            template_file: Optional[str] = None,
            header_text: str = "格式化文档",
            toc_title: str = "目 录",
            save_intermediate: bool = False,
            intermediate_dir: Optional[str] = None,
            document_type: int = 1
    ) -> bool:
        """
        完整的文档格式化转换流程

        Args:
            source_file: 源文档路径
            output_file: 输出文档路径
            template_file: 模板文档路径（可选，默认使用template/reference.docx）
            header_text: 页眉文本
            toc_title: 目录标题（可选，默认为"目 录"）
            save_intermediate: 是否保存中间文件（默认为False）
            intermediate_dir: 中间文件保存目录（仅在save_intermediate为True时有效）
            document_type: 文档类型 (1, 2, 3, 4)

        Returns:
            bool: 转换是否成功
        """

        # 设置是否保存中间文件
        self.save_intermediate_files = save_intermediate

        # 设置文档类型
        self.document_type = document_type
        print(f"   文档类型: {document_type}")
        # 如果指定了中间文件目录，则使用该目录
        if save_intermediate and intermediate_dir:
            self.debug_output_dir = intermediate_dir

        # 如果未指定模板文件，则使用默认模板
        if template_file is None:
            # 获取项目根目录
            current_dir = os.path.dirname(os.path.abspath(__file__))
            parent_dir = os.path.dirname(current_dir)
            template_file = os.path.join(parent_dir, 'template', 'reference_content.docx')
            if self.save_intermediate_files:
                print(f"信息: 未指定模板文件，使用默认模板: {template_file}")

        start_time = time.time()

        print("🚀 开始文档格式化转换")
        print("=" * 80)
        print(f"📁 源文档: {source_file}")
        print(f"📄 模板文档: {template_file}")
        print(f"📤 输出文档: {output_file}")
        print(f"📋 页眉文本: '{header_text}'")
        print(f"📋 目录标题: '{toc_title}'")
        if save_intermediate:
            print(f"💾 保存中间文件: 是")
            print(f"📂 中间文件目录: {self.debug_output_dir}")
        else:
            print(f"💾 保存中间文件: 否")
        print(f"⏰ 开始时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print("=" * 80)
        print()

        # 验证输入文件
        if not self.validate_input_files(source_file, template_file):
            return False

        # 创建临时目录
        self.temp_dir = tempfile.mkdtemp(prefix="doc_converter_")
        if self.save_intermediate_files:
            print(f"📁 临时目录: {self.temp_dir}")

        try:
            # 步骤0: 页眉页脚替换
            updated_template_file = self.step0_replace_header_footer(source_file, template_file)
            # 使用更新后的模板文件进行后续处理
            template_file = updated_template_file

            # 10.22新增
            preprocessed_file = self.step0_preprocess_headings(source_file)
            source_file = preprocessed_file  # 使用预处理后的文件进行后续操作
            self.step0_processed_source = preprocessed_file

            # 步骤1: 文档拆分
            cover_toc_file, content_file = self.step1_split_document(source_file)
            if not cover_toc_file or not content_file:
                print("❌ 转换失败: 文档拆分失败")
                return False

            original_content_for_tables = content_file

            # 步骤1.5: 按页面方向拆分正文
            portrait_content_file, landscape_content_file = self.step1_5_split_by_orientation(content_file)
            content_file = portrait_content_file

            # 步骤2: Pandoc转换
            pandoc_file = self.step2_pandoc_convert(content_file, template_file)
            if not pandoc_file:
                print("⚠️ Pandoc转换失败，使用原始正文文件继续")
                pandoc_file = content_file

            # 步骤3: 表格格式化
            # 使用步骤1拆分后的正文文件中的表格替换步骤2 Pandoc转换后的文件中的表格
            table_formatted_file = self.step3_format_tables(
                content_file=pandoc_file,
                template_file=template_file,
                original_content_file=original_content_for_tables  # 传入原始正文内容文件
            )
            if not table_formatted_file:
                print("⚠️ 表格格式化失败，使用Pandoc转换后的文件继续")
                table_formatted_file = pandoc_file

            # 步骤4: 文档合并
            success = self.step4_merge_documents(cover_toc_file, table_formatted_file, output_file)

            if success:
                # 步骤4.5: 在竖版结果末尾追加横版节
                if landscape_content_file and os.path.exists(landscape_content_file):
                    try:
                        has_landscape = os.path.getsize(landscape_content_file) > 0
                    except Exception:
                        has_landscape = False

                    if has_landscape:
                        print("-" * 50)
                        print("📚 步骤4.5: 在合并文档尾部追加横版节")
                        if not self.temp_dir:
                            raise ValueError("临时目录未初始化")
                        base_name = os.path.splitext(os.path.basename(output_file))[0]
                        step45_output = os.path.join(self.temp_dir, f"{base_name}_step45_append_landscape.docx")
                        append_success = self._append_document_with_formatting(
                            base_file=output_file,
                            file_to_append=landscape_content_file,
                            output_file=step45_output,
                        )
                        if append_success and os.path.exists(step45_output):
                            shutil.copy2(step45_output, output_file)
                            self.intermediate_files['landscape_appended'] = step45_output
                            if self.save_intermediate_files:
                                print(f"   横版追加后文档: {os.path.basename(step45_output)}")
                                print(f"📁 正在保存step4_5中间文件到: {self.debug_output_dir}")
                                self._save_intermediate_file(step45_output, "step4_5_append", "竖横合并后")
                        else:
                            print("⚠️ 横版文档追加失败，继续使用竖版结果")
                    else:
                        print("ℹ️ 横版文档为空，跳过步骤4.5")
                else:
                    print("ℹ️ 无横版内容，跳过步骤4.5")

                # 步骤5: 更新目录标题
                toc_update_success = self.step5_update_toc_title(output_file, toc_title)
                # if toc_update_success:
                #     print("✅ 目录标题更新完成!")
                # else:
                #     print("⚠️ 目录标题更新失败，继续使用原有标题")
                if not toc_update_success:
                    print("⚠️ 目录标题更新失败，继续使用原有标题")

                # 步骤6: 图片格式化
                picture_format_success = self.step6_format_pictures(output_file)
                if not picture_format_success:
                    print("⚠️ 图片格式化失败，继续使用原有格式")

                # 步骤7: 库号信息格式化
                library_number_format_success = self.step7_format_library_number(output_file)
                if not library_number_format_success:
                    print("⚠️ 库号信息格式化失败，继续使用原有格式")

                # 步骤8: 在目录后插入分节符
                section_break_insert_success = self.step8_insert_section_break(output_file)
                if not section_break_insert_success:
                    print("⚠️ 在目录后插入分节符失败，继续使用原有格式")

                # 步骤9: 处理文档节的页码设置
                section_page_process_success = self.step9_process_sections(output_file)
                if not section_page_process_success:
                    print("⚠️ 处理文档节的页码设置失败，继续使用原有格式")

                # 步骤10: 删除文档中所有的突出显示
                no_highlights_file = self.step10_remove_highlights(output_file)
                # 如果成功去除突出显示，将结果复制回输出文件
                if no_highlights_file != output_file and os.path.exists(no_highlights_file):
                    shutil.copy2(no_highlights_file, output_file)
                    print("✅ 突出显示删除成功!")
                else:
                    print("⚠️ 删除突出显示失败，继续使用原有格式")

                # 步骤11: 高级库号信息格式化
                library_number_advanced_format_success = self.step11_format_library_number_advanced(output_file)
                if not library_number_advanced_format_success:
                    print("⚠️ 库号信息高级格式化失败，继续使用原有格式")

                # 步骤12: 页眉页脚后处理
                if self.step0_processed_source:
                    table_replace_success = self.step13_word_table_replace(
                        edited_docx=output_file,
                        original_docx=self.step0_processed_source
                    )
                    if not table_replace_success:
                        print('⚠️ 表格替换失败，继续使用当前文档')
                else:
                    print('⚠️ 跳过表格替换: 未获取步骤0的输出文件')

                headers_footers_process_success = self.step12_post_process_headers_footers(output_file)
                if not headers_footers_process_success:
                    print("⚠️ 页眉页脚处理失败，继续使用原有格式")

                end_time = time.time()
                duration = end_time - start_time

                print("\n" + "=" * 80)
                print("✅ 文档转换成功!")
                print(f"⏱️ 总耗时: {duration:.2f} 秒")
                print(f"📤 最终文档: {output_file}")

                # 显示中间文件信息
                if save_intermediate:
                    print("\n📋 中间文件保存在临时目录:")
                    for key, path in self.intermediate_files.items():
                        if os.path.exists(path):
                            print(f"   {key}: {os.path.basename(path)}")

                    print(f"\n📁 所有中间文件已同步保存到: {self.debug_output_dir}")
                    print("🔍 您可以在该目录中查看每个步骤的处理结果，便于调试和优化")
                else:
                    print("\n📋 中间文件未保存（根据设置）")

                print("\n💡 提示:")
                print("   - 在Word中打开文档，右键目录选择'更新域'来刷新页码")
                print("   - 检查文档格式是否符合要求")
                print("=" * 80)

                return True
            else:
                print("❌ 转换失败: 文档合并失败")
                return False

        except Exception as e:
            print(f"\n❌ 转换过程中发生错误: {str(e)}")
            return False

        finally:
            # 不立即清理临时文件，保留中间结果供调试
            if self.save_intermediate_files:
                print(f"\n📁 临时文件保留在: {self.temp_dir}")
                print("   您可以手动删除该目录，或重启程序时自动清理")


def quick_convert_document(
        source_file: str,
        output_file: str,
        template_file: Optional[str] = None,
        header_text: str = "格式化文档",
        toc_title: str = "目 录",
        save_intermediate: bool = False,
        intermediate_dir: Optional[str] = None,
        document_type: int = 1
) -> bool:
    """
    便捷函数: 快速进行文档格式化转换

    Args:
        source_file: 源文档路径
        output_file: 输出文档路径
        template_file: 模板文档路径（可选，默认使用template/reference.docx）
        header_text: 页眉文本
        toc_title: 目录标题（可选，默认为"目 录"）
        save_intermediate: 是否保存中间文件（默认为False）
        intermediate_dir: 中间文件保存目录（仅在save_intermediate为True时有效）
        document_type: 文档类型 (1, 2, 3, 4)

    Returns:
        bool: 转换是否成功
    """
    with DocumentConverter(document_type=document_type) as converter:
        return converter.convert_document(
            source_file=source_file,
            output_file=output_file,
            template_file=template_file,
            header_text=header_text,
            toc_title=toc_title,
            save_intermediate=save_intermediate,
            intermediate_dir=intermediate_dir,
            document_type=document_type
        )


if __name__ == "__main__":
    source_file = r"C:\Users\yanha\Desktop\数字总师\文档\可行性报告（test）.docx"

    # 创建result目录
    result_dir = os.path.join(current_dir, "result")
    os.makedirs(result_dir, exist_ok=True)

    # 输出文件路径
    output_file = os.path.join(result_dir, "formatted_document.docx")

    # 使用默认模板（template/reference.docx）

    # 执行转换
    with DocumentConverter(document_type=1) as converter:  # 添加文档类型参数
        success = converter.convert_document(
            source_file=source_file,
            output_file=output_file,
            header_text="数字总师可行性报告",
            toc_title="目      录",
            save_intermediate=False,
            document_type=1  # 添加文档类型参数
        )

